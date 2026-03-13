# noinspection PyUnresolvedReferences
# noinspection PyStringFormat
import numpy as np
import matplotlib.pyplot as plt
import os
import time
import hashlib
from functools import lru_cache
import locale
import json

# 尝试导入 streamlit，如果失败则创建一个模拟对象
try:
    import streamlit as st
except ImportError:
    # 创建一个模拟的 st 模块
    class MockSt:
        def error(self, msg):
            print(f"ERROR: {msg}")
        def warning(self, msg):
            print(f"WARNING: {msg}")
        def success(self, msg):
            print(f"SUCCESS: {msg}")
        def info(self, msg):
            print(f"INFO: {msg}")
        def markdown(self, msg, **kwargs):
            pass
        def write(self, msg):
            pass
        def subheader(self, msg):
            pass
        def expander(self, msg, **kwargs):
            class MockExpander:
                def __enter__(self):
                    pass
                def __exit__(self, *args):
                    pass
            return MockExpander()
        def divider(self):
            pass
        def caption(self, msg):
            pass
        def latex(self, msg):
            pass
        def empty(self):
            class MockEmpty:
                def info(self, msg):
                    pass
                def success(self, msg):
                    pass
            return MockEmpty()
        def columns(self, n):
            return [MockSt() for _ in range(n)]
        def slider(self, *args, **kwargs):
            return 5
        def text_input(self, *args, **kwargs):
            return "1"
        def number_input(self, *args, **kwargs):
            return 1.0
        def selectbox(self, *args, **kwargs):
            return "根轨迹分析"
        def button(self, *args, **kwargs):
            return False
        def download_button(self, *args, **kwargs):
            pass
        def session_state(self):
            return {}
    
    st = MockSt()
    st.session_state = {}


# ---------------------- 全局 Matplotlib 配置（确保中文在本地和部署环境都能正常显示） ----------------------
def configure_matplotlib_for_chinese():
    """配置Matplotlib以支持中文显示，兼容Windows、Linux和部署环境"""
    import platform
    
    # 不设置locale，让系统保持默认设置
    plt.rcParams['axes.formatter.use_locale'] = False
    
    # 配置中文支持 - 尝试多种字体，确保在不同环境下都能工作
    system_name = platform.system()
    if system_name == 'Windows':
        # Windows环境
        plt.rcParams['font.family'] = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei', 'sans-serif']
    elif system_name == 'Darwin':
        # macOS环境
        plt.rcParams['font.family'] = ['PingFang SC', 'Heiti TC', 'Arial Unicode MS', 'sans-serif']
    else:
        # Linux/部署环境
        plt.rcParams['font.family'] = ['WenQuanYi Micro Hei', 'DejaVu Sans', 'Arial', 'sans-serif']
    
    # 确保负号正常显示
    plt.rcParams['axes.unicode_minus'] = False
    
    # 禁用Matplotlib的TeX渲染，避免字体问题
    plt.rcParams['text.usetex'] = False

# 初始化Matplotlib中文配置
configure_matplotlib_for_chinese()

# ---------------------- 兼容处理：先导入所有依赖，避免IDE报错 ----------------------
# 修复火山方舟SDK导入逻辑（只检查新版SDK）
try:
    from volcenginesdkarkruntime import Ark

    VOLC_AVAILABLE = True
except ImportError:
    VOLC_AVAILABLE = False
    st.warning("未安装volcengine-python-sdk[ark]库！AI功能将不可用：pip install --upgrade \"volcengine-python-sdk[ark]\"")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    import docx
    from docx.shared import Inches
    # 尝试导入 Spire.Doc 库
    try:
        from spire.doc import Document, BreakType, FileFormat, HorizontalAlignment, OfficeMath
        SPIRE_AVAILABLE = True
    except ImportError:
        SPIRE_AVAILABLE = False
        st.warning("未安装spire.doc.free库！Word公式功能将不可用：pip install spire.doc.free")

    EXPORT_AVAILABLE = True
except ImportError:
    EXPORT_AVAILABLE = False
    SPIRE_AVAILABLE = False
    st.warning("未安装reportlab/python-docx库！导出功能将不可用：pip install reportlab python-docx")

try:
    import control as ctrl
    from control.matlab import pole, zero, step, feedback, evalfr

    CONTROL_AVAILABLE = True
except ImportError:
    CONTROL_AVAILABLE = False
    st.warning("未安装control库！请先运行：pip install control==0.10.2（兼容版）")

# ---------------------- 完整知识点库（新设计） ----------------------
KNOWLEDGE_CONTENT = {
    "一、基础概念模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 一、基础概念模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">入门必懂 · 所有分析的前提</p>

---

### 1. 控制系统的基本概念 ⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>定义</strong>：控制系统是由若干相互联系、相互作用的部件组成，能对被控对象的工作状态进行控制，使被控量按预定规律变化的系统。</p>
</div>

**分类**：
- 按控制方式：开环控制系统、闭环控制系统（负反馈为主）
- 按系统特性：线性系统（叠加原理适用）、非线性系统（叠加原理不适用）；定常系统（参数不随时间变化）、时变系统（参数随时间变化）
- 按信号形式：连续控制系统、离散控制系统（采样控制）

<div style="background: #fff0f0; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>核心区别</strong>：开环无反馈（精度低、抗干扰差），闭环有负反馈（精度高、抗干扰强，重点）。
</div>

**通俗解读**：开环像"盲操作"，只给指令不检查结果；闭环像"带眼睛操作"，会根据结果调整指令。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：闭环系统不一定稳定，需通过判据验证；线性系统必须满足叠加原理（齐次性 + 可加性）。
</div>

**考点提示**：选择题判断系统类型；填空题区分开环/闭环的优缺点。

---

### 2. 传递函数 ⭐⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>专业定义</strong>：线性定常系统在零初始条件下，系统输出量的拉普拉斯变换与输入量的拉普拉斯变换之比，记为 $G(s)$。</p>
</div>

**核心公式**：
$$G(s) = \\frac{C(s)}{R(s)} = \\frac{b_m s^m + b_{m-1} s^{m-1} + \\dots + b_1 s + b_0}{a_n s^n + a_{n-1} s^{n-1} + \\dots + a_1 s + a_0}$$
（分子：输出多项式；分母：输入多项式，且 $n\\geq m$，否则物理不可实现）

**基本性质**：
1. 只与系统的结构、参数有关，与输入信号、初始条件无关；
2. 仅适用于线性定常系统，非线性、时变系统不适用；
3. 传递函数是复频域的数学模型，不反映系统的物理结构；
4. 若系统有多个输入、多个输出，需用传递函数矩阵描述。

**通俗解读**：传递函数就是系统的"输入-输出说明书"，告诉你"输入一个信号，系统会输出什么信号"，和具体输入无关，只看系统本身的"脾气"。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 必须满足"零初始条件"（初始位移、初始速度均为 0），否则传递函数不成立；
2. 不能用传递函数描述非线性系统（比如含有饱和、死区的系统）；
3. 分母多项式的次数 $n$ 必须大于等于分子次数 $m$，否则系统物理不可实现。
</div>

**考点提示**：
1. 填空题：根据微分方程求传递函数；
2. 大题：传递函数化简、方框图等效变换；
3. 选择题：判断传递函数的基本性质。

---

### 3. 极点与零点 ⭐⭐⭐⭐⭐

**专业定义**：
- 零点：传递函数分子多项式等于 0 的根，记为 $z_1, z_2, \\dots, z_m$；
- 极点：传递函数分母多项式等于 0 的根，记为 $p_1, p_2, \\dots, p_n$；
- 特征方程：分母多项式等于 0 的方程，即 $a_n s^n + a_{n-1} s^{n-1} + \\dots + a_1 s + a_0 = 0$。

<div style="background: #fff0f0; padding: 20px; border-radius: 12px; margin: 20px 0;">
<strong>核心作用</strong>：
1. <strong>极点</strong>：决定系统的稳定性、动态响应速度和振荡特性（核心中的核心）；
   - 极点在 $s$ 平面左半平面（$\\text{实部} < 0$）$\\Rightarrow$ 系统稳定；
   - 极点在 $s$ 平面右半平面（$\\text{实部} > 0$）$\\Rightarrow$ 系统不稳定；
   - 极点在 $s$ 平面虚轴（$\\text{实部} = 0$）$\\Rightarrow$ 系统临界稳定（等幅振荡）；
   - 极点越靠近虚轴，系统响应越慢；极点实部绝对值越大，响应越快。
2. <strong>零点</strong>：不影响系统的稳定性，仅影响系统的动态性能（超调量、上升时间）和相位特性；
   - 零点靠近虚轴，会增大系统超调量；零点远离虚轴，对系统性能影响较小。
</div>

**通俗解读**：
- 极点：系统的"固有性格"，决定系统"稳不稳、反应快不快"；
- 零点：系统的"调节按钮"，不改变系统的"性格"，但能微调系统的"反应方式"。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 零点不影响系统稳定性，只有极点能决定系统是否稳定；
2. 多重极点（比如二重极点）会导致系统响应变慢，且超调量减小；
3. 若传递函数分子、分母有相同的因子（零极点抵消），会丢失系统的部分动态特性，需注意避免零极点抵消（除非题目明确要求）。
</div>

**考点提示**：
1. 选择题：根据极点位置判断系统稳定性；
2. 大题：结合根轨迹分析极点随增益 $K$ 的变化；
3. 填空题：写出系统的极点、零点。

---

### 4. 方框图等效变换 ⭐⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>核心规则（必须背）</strong></p>
</div>

1. **串联连接**：两个系统串联，总传递函数等于两个传递函数的乘积，$G(s) = G_1(s) \\cdot G_2(s)$；
   - 特点：前一个系统的输出是后一个系统的输入，无交叉反馈。
2. **并联连接**：两个系统并联，总传递函数等于两个传递函数的和，$G(s) = G_1(s) + G_2(s)$；
   - 特点：两个系统有相同的输入，输出相加（或相减）。
3. **反馈连接（核心中的核心）**：
   - 负反馈（重点）：总传递函数 $G(s) = \\frac{G_1(s)}{1 + G_1(s)H(s)}$（分母为 $1 + $ 开环传递函数）；
   - 正反馈（较少考）：总传递函数 $G(s) = \\frac{G_1(s)}{1 - G_1(s)H(s)}$（分母为 $1 - $ 开环传递函数）；
   - 开环传递函数：$G_1(s)H(s)$（反馈通道与前向通道传递函数的乘积）。
4. **分支点前移 / 后移**：
   - 分支点前移：从系统输出端前移到输入端，需在分支通道上乘以对应系统的传递函数；
   - 分支点后移：从系统输入端后移到输出端，需在分支通道上除以对应系统的传递函数。
5. **相加点前移 / 后移**：
   - 相加点前移：从系统输出端前移到输入端，需在被移动的通道上乘以对应系统的传递函数；
   - 相加点后移：从系统输入端后移到输出端，需在被移动的通道上除以对应系统的传递函数。

**化简步骤**：
1. 消除交叉反馈（若有），将复杂框图转化为串联、并联、反馈的基本形式；
2. 按照"先串联 / 并联，后反馈"的顺序逐步化简；
3. 最终化简为单回路闭环传递函数（或单传递函数）。

**通俗解读**：方框图变换就像"数学运算的变形"，核心是"保持输入输出关系不变"，把复杂的"网状框图"化简成简单的"单条线路"，方便后续分析。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 反馈连接的符号：负反馈是"$1 + G_1H$"，正反馈是"$1 - G_1H$"，极易记反；
2. 分支点、相加点不能随意交换位置，必须补乘 / 除对应传递函数，否则会改变输入输出关系；
3. 化简过程中，避免零极点抵消（除非题目明确要求）。
</div>

**考点提示**：
1. 大题：方框图化简求闭环传递函数；
2. 填空题：根据框图写出开环 / 闭环传递函数。

---

### 5. 信号流图与梅森公式 ⭐⭐⭐

**信号流图定义**：用节点（表示信号）和支路（表示信号传递关系）组成的图形，用于描述系统各变量之间的传递关系，比方框图更简洁。

**核心术语**：
1. 节点：表示系统中的信号（如输入信号 $R(s)$、输出信号 $C(s)$、中间信号）；
2. 支路：连接两个节点的线段，支路增益即两个信号的传递函数；
3. 前向通路：从输入节点到输出节点，不重复经过任何节点的通路；
4. 回路：从某个节点出发，经过若干支路后回到原节点，不重复经过任何节点的闭合通路；
5. 不接触回路：两个回路没有公共节点。

**梅森公式（核心公式）**：
$$G(s) = \\frac{1}{\\Delta} \\sum_{k=1}^{n} P_k \\Delta_k$$
其中：
- $\\Delta$：信号流图的特征式，$\\Delta = 1 - \\sum L_i + \\sum L_iL_j - \\sum L_iL_jL_k + \\dots$（$L_i$为单个回路增益，$L_iL_j$为两个不接触回路增益的乘积，以此类推）；
- $P_k$：第 $k$ 条前向通路的增益；
- $\\Delta_k$：第 $k$ 条前向通路的余因子式，即去掉与第 $k$ 条前向通路接触的所有回路后，剩余信号流图的特征式。

**通俗解读**：梅森公式是"直接求闭环传递函数的捷径"，无需逐步化简框图，适合复杂系统（多回路、多前向通路）。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 前向通路必须是"从输入到输出，不重复节点"，否则不算有效前向通路；
2. 回路必须是"闭合通路，不重复节点"，且不接触回路是"无公共节点"，不是无公共支路；
3. 计算 $\\Delta$ 时，需注意符号（减单个回路，加两个不接触回路，减三个不接触回路，以此类推）。
</div>

**考点提示**：大题中复杂系统的传递函数求解（可选，比框图化简更快捷）。
</div>
"""
    },
    "二、二阶系统模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 二、二阶系统模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">绝对核心 · 每年必考 · 占比 20%+</p>

---

### 1. 二阶系统的标准形式与参数 ⭐⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>专业定义</strong>：传递函数可化为如下标准形式的系统，称为二阶系统：</p>
</div>

**核心公式**：
$$G(s) = \\frac{\\omega_n^2}{s^2 + 2\\zeta\\omega_n s + \\omega_n^2}$$

**其中**：
1. $\\omega_n$：无阻尼自然频率（单位：$\\text{rad/s}$），反映系统的响应速度；
2. $\\zeta$：阻尼比（无量纲），反映系统的振荡程度；
3. 阻尼振荡频率 $\\omega_d$：$\\omega_d = \\omega_n\\sqrt{1 - \\zeta^2}$（仅 $0 < \\zeta < 1$ 时存在）。

<div style="background: #fff0f0; padding: 20px; border-radius: 12px; margin: 20px 0;">
<strong>核心参数解读</strong>：
1. <strong>$\\omega_n$（自然频率）</strong>：
   - $\\omega_n$ 越大，系统响应越快（上升时间、调节时间越短）；
   - $\\omega_n$ 越小，系统响应越慢。
2. <strong>$\\zeta$（阻尼比）</strong>：
   - $\\zeta$ 是"实际阻尼系数"与"临界阻尼系数"的比值，决定系统的振荡特性；
   - $\\zeta$ 的取值不同，系统的响应形式完全不同。
</div>

**阻尼特性分类（必须背）**：

| 阻尼比 $\\zeta$ | 系统特性 | 响应形式 | 核心特点 | 重要程度 |
|:---:|:---:|:---:|:---|:---:|
| $\\zeta = 0$ | 无阻尼 | 等幅振荡 | 输出一直振荡，不会稳定，实际中不存在 | ⭐⭐ |
| $0 < \\zeta < 1$ | 欠阻尼 | 衰减振荡 | 输出先冲过稳态值（超调），然后衰减到稳态值，是工程中最常用的形式 | ⭐⭐⭐⭐⭐ |
| $\\zeta = 1$ | 临界阻尼 | 无振荡最快收敛 | 输出无超调，且收敛速度最快，适用于不允许振荡的场景 | ⭐⭐⭐ |
| $\\zeta > 1$ | 过阻尼 | 单调收敛 | 输出无超调，但收敛速度慢，适用于对响应速度要求不高的场景 | ⭐⭐ |

**通俗解读**：
- $\\omega_n$：系统的"跑步速度"，$\\omega_n$ 越大，跑得越快（响应越快）；
- $\\zeta$：系统的"刹车力度"，$\\zeta$ 越大，刹车越狠（振荡越小）；
- 欠阻尼：刹车力度适中，跑得快且轻微振荡；临界阻尼：刹车力度刚好，不振荡且最快停稳；过阻尼：刹车太狠，跑得慢且不振荡。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 90% 的题目考查"欠阻尼（$0 < \\zeta < 1$）"，其他阻尼形式仅作概念考察；
2. 阻尼振荡频率 $\\omega_d$ 仅在 $0 < \\zeta < 1$ 时存在，$\\zeta \\geq 1$ 时无振荡，$\\omega_d$ 无意义；
3. 二阶系统的标准形式必须是"分母为 $s^2 + 2\\zeta\\omega_n s + \\omega_n^2$"，若不是标准形式，需先化简为标准形式，再确定 $\\zeta$ 和 $\\omega_n$。
</div>

**考点提示**：
1. 填空题：根据传递函数求 $\\zeta$ 和 $\\omega_n$；
2. 大题：已知 $\\zeta$ 和 $\\omega_n$，分析系统响应特性、计算动态性能指标；
3. 选择题：判断二阶系统的阻尼类型。

---

### 2. 二阶系统的动态性能指标（欠阻尼系统）⭐⭐⭐⭐⭐

**动态性能指标定义**：描述系统过渡过程（从输入作用到系统稳定）的性能，核心指标有 5 个（前 4 个为重点）。

**核心公式（必须背，欠阻尼系统 $0 < \\zeta < 1$）**：
1. **超调量（$\\sigma\\%$）**：系统输出冲过稳态值的百分比，反映系统的振荡程度，越小越稳定。
   $$\\sigma\\% = e^{-\\frac{\\zeta\\pi}{\\sqrt{1-\\zeta^2}}} \\times 100\\%$$
2. **峰值时间（$t_p$）**：系统输出达到第一个峰值（最大超调点）的时间，反映系统的振荡频率。
   $$t_p = \\frac{\\pi}{\\omega_n\\sqrt{1-\\zeta^2}} = \\frac{\\pi}{\\omega_d}$$
3. **调节时间（$t_s$）**：系统输出衰减到稳态值的 $\\pm 2\\%$（或 $\\pm 5\\%$）误差带内，且不再超出的时间，反映系统的收敛速度，越小越好。
   - $2\\%$ 误差带：$t_s = \\frac{4}{\\zeta\\omega_n}$
   - $5\\%$ 误差带：$t_s = \\frac{3}{\\zeta\\omega_n}$
4. **上升时间（$t_r$）**：系统输出从稳态值的 $10\\%$ 上升到 $90\\%$（欠阻尼系统）的时间，反映系统的响应速度，越小越快。
   $$t_r = \\frac{\\pi - \\theta}{\\omega_n\\sqrt{1-\\zeta^2}} \\quad (\\theta = \\arccos\\zeta)$$
5. **延迟时间（$t_d$）**：系统输出从 0 上升到稳态值的 $50\\%$ 的时间，反映系统的初始响应速度（次重点）。

<div style="background: #fff0f0; padding: 20px; border-radius: 12px; margin: 20px 0;">
<strong>指标之间的关系（分析题重点）</strong>：
1. $\\zeta$ 增大：超调量 $\\sigma\\%$ 减小，上升时间 $t_r$ 增大，调节时间 $t_s$ 先减小后增大（$\\zeta$=0.4~0.8 时，$t_s$ 最小）；
2. $\\omega_n$ 增大：上升时间 $t_r$、峰值时间 $t_p$、调节时间 $t_s$ 均减小，超调量 $\\sigma\\%$ 不变（仅与 $\\zeta$ 有关）。
</div>

**通俗解读**：
- 超调量：系统"冲过目标"的幅度，比如目标是 10，输出冲到 12，超调量就是 20%；
- 调节时间：系统"稳定下来"的时间，比如调节时间 5s，就是 5s 后系统输出稳定在目标值附近；
- 上升时间：系统"从起步到接近目标"的时间，比如上升时间 1s，就是 1s 后输出达到目标值的 90%。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 超调量的分母必须是"实际稳态值 $y_{ss}$"，不是 1，若系统稳态值不是 1，用 $\\frac{y_{max} - y_{ss}}{y_{ss}} \\times 100\\%$ 计算；
2. 调节时间的误差带：默认 $2\\%$ 误差带，若题目明确要求 $5\\%$，则用 $\\frac{3}{\\zeta\\omega_n}$；
3. 只有欠阻尼系统（$0 < \\zeta < 1$）有超调量、峰值时间，临界阻尼（$\\zeta$=1）、过阻尼（$\\zeta$>1）无超调，无峰值时间；
4. 超调量仅与 $\\zeta$ 有关，与 $\\omega_n$ 无关；调节时间、上升时间、峰值时间与 $\\zeta$ 和 $\\omega_n$ 都有关。
</div>

**考点提示**：
1. 计算题：已知 $\\zeta$ 和 $\\omega_n$，求 $\\sigma\\%$、$t_p$、$t_s$、$t_r$（每年必考，分值 5-10 分）；
2. 分析题：改变 $\\zeta$ 和 $\\omega_n$ 对动态性能指标的影响（高频）；
3. 填空题：写出动态性能指标的公式。

---

### 3. 二阶系统的稳态性能（结合稳态误差模块）⭐⭐⭐⭐

**稳态误差定义**：系统在典型输入作用下，输出稳态值与参考输入的偏差，反映系统的稳态精度。

**二阶系统的稳态误差计算（终值定理）**：
$$e_{ss} = \\lim_{s\\to 0} s \\cdot E(s) = \\lim_{s\\to 0} \\frac{sR(s)}{1 + G(s)}$$
其中 $E(s) = R(s) - C(s)$（误差信号），$G(s)$ 为二阶系统开环传递函数。

**二阶系统的系统型别（重点）**：
1. **0 型二阶系统**：开环传递函数无积分环节（分母无 $s$），比如 $G(s) = \\frac{K}{s^2 + 2\\zeta\\omega_n s + \\omega_n^2}$；
2. **I 型二阶系统**：开环传递函数有 1 个积分环节（分母有 1 个 $s$），比如 $G(s) = \\frac{K}{s(s^2 + 2\\zeta\\omega_n s + \\omega_n^2)}$；
3. **II 型二阶系统**：开环传递函数有 2 个积分环节（分母有 2 个 $s$），比如 $G(s) = \\frac{K}{s^2(s^2 + 2\\zeta\\omega_n s + \\omega_n^2)}$。

**典型输入下的稳态误差（必须背）**：

| 系统型别 | 单位阶跃输入（$R(s)=1/s$） | 单位斜坡输入（$R(s)=1/s^2$） | 单位抛物线输入（$R(s)=1/s^3$） |
|:---:|:---:|:---:|:---:|
| 0 型二阶系统 | $e_{ss} = \\frac{1}{1+K}$（$K$ 为开环增益） | $e_{ss} = \\infty$（无法跟踪） | $e_{ss} = \\infty$（无法跟踪） |
| I 型二阶系统 | $e_{ss} = 0$（无静差） | $e_{ss} = \\frac{1}{K_v}$（$K_v$=$K$ 为速度误差系数） | $e_{ss} = \\infty$（无法跟踪） |
| II 型二阶系统 | $e_{ss} = 0$（无静差） | $e_{ss} = 0$（无静差） | $e_{ss} = \\frac{1}{K_a}$（$K_a$=$K$ 为加速度误差系数） |

**通俗解读**：
- 0 型二阶系统：能对准固定靶（阶跃输入），但跟不上移动靶（斜坡、抛物线输入）；
- I 型二阶系统：能对准固定靶、跟得上匀速移动靶（斜坡输入），但跟不上加速移动靶（抛物线输入）；
- II 型二阶系统：能对准固定靶、跟得上匀速/加速移动靶（抛物线输入），稳态精度最高。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 二阶系统的系统型别由"开环传递函数的积分环节个数"决定，不是闭环；
2. 积分环节越多，稳态精度越高，但系统稳定性越差（后续校正模块会重点讲）；
3. 终值定理仅适用于系统稳定（极点全在左半平面），若系统不稳定，无法用终值定理计算稳态误差。
</div>

**考点提示**：
1. 计算题：已知二阶系统开环传递函数，求典型输入的稳态误差（必考）；
2. 分析题：增加积分环节对二阶系统稳态性能和动态性能的影响（高频）。

---

### 4. 二阶系统的校正（结合系统校正模块，大题必考）⭐⭐⭐⭐⭐

**核心需求**：二阶系统的原始性能（如超调量大、调节时间长、稳态误差大）不满足要求，需通过添加校正环节，调整 $\\zeta$ 和 $\\omega_n$，使系统性能达标。

**常用校正方式（重点）**：
1. **比例校正（P 校正）**：仅调整开环增益 $K$，不改变系统的极点、零点位置；
   - 作用：增大 $K$，减小稳态误差（提高稳态精度），但会减小 $\\zeta$，增大超调量，降低稳定性；
   - 适用场景：稳态误差大、超调量满足要求的系统。
2. **比例-积分校正（PI 校正）**：添加积分环节，提高系统型别，消除稳态误差；
   - 作用：消除阶跃稳态误差，提高稳态精度，但会降低相位裕度，增大超调量，需配合微分环节使用；
   - 适用场景：稳态误差大、不允许有静差的系统。
3. **比例-微分校正（PD 校正）**：添加微分环节，增大 $\\zeta$，减小超调量；
   - 作用：减小超调量，提高稳定性，加快响应速度，但不改变稳态误差；
   - 适用场景：超调量大、稳态误差满足要求的系统。
4. **比例-积分-微分校正（PID 校正）**：结合 P、I、D 三个环节的优点，兼顾稳态精度和动态性能；
   - 作用：消除稳态误差，减小超调量，加快响应速度，是工程中最常用的校正方式；
   - 适用场景：对稳态精度和动态性能要求都高的系统。

**通俗解读**：
- P 校正："调增益"，让系统更精准，但可能更振荡；
- PI 校正："补积分"，让系统无静差，但可能更振荡；
- PD 校正："加微分"，让系统更稳定，不振荡，但精度不变；
- PID 校正："全能选手"，既精准、又稳定、又快速。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. PI 校正会降低系统稳定性，不能单独用于不稳定系统；
2. PD 校正对高频噪声敏感，若系统存在高频干扰，需谨慎使用；
3. 校正环节需串联在系统前向通道，常考 PI、PD、PID 校正的参数整定。
</div>

**考点提示**：
1. 大题：设计校正环节（PI/PD/PID），使二阶系统性能达标（每年必考，分值 10-15 分）；
2. 分析题：校正前后二阶系统性能对比（高频）。
</div>
"""
    },
    "三、稳定性判据模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 三、稳定性判据模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">必考 · 占比 15%+ · 核心是判稳 + 求稳定范围</p>

---

### 1. 稳定性的基本概念 ⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>专业定义（李雅普诺夫稳定性）</strong>：若系统在初始扰动作用下，其输出能逐渐衰减并回到原平衡状态，则系统稳定；若输出逐渐发散，则系统不稳定；若输出保持等幅振荡，则系统临界稳定。</p>
</div>

**线性定常系统稳定的充要条件（核心）**：系统的所有闭环极点都在 $s$ 平面的左半平面（$\\text{实部} < 0$）。

**稳定性的分类**：
1. 渐近稳定：输出逐渐衰减到平衡状态（闭环极点全在左半平面），工程中最常用；
2. 临界稳定：输出保持等幅振荡（闭环极点在虚轴，且无多重极点），实际中视为不稳定；
3. 不稳定：输出逐渐发散（存在闭环极点在右半平面，或虚轴上有多重极点）。

**通俗解读**：稳定性就是系统"抗干扰能力"——受到干扰后，能回到原来的状态，就是稳定；越来越偏离原来的状态，就是不稳定；一直振荡不偏离，就是临界稳定（实际中不能用）。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 系统稳定的充要条件是"闭环极点全在左半平面"，不是开环极点；
2. 虚轴上有多重极点（比如二重虚极点），系统不稳定（不是临界稳定）；
3. 系统的稳定性与输入信号、初始条件无关，仅由闭环极点位置决定。
</div>

**考点提示**：
1. 选择题：判断系统是否稳定（必考）；
2. 大题：用判据求系统稳定的增益 $K$ 范围（每年必考）。

---

### 2. 劳斯（Routh）稳定判据 ⭐⭐⭐⭐⭐

**核心作用**：不用求解特征方程，直接通过特征方程的系数，判断系统的稳定性，还能确定右半平面极点的个数。

**适用场景**：线性定常系统，特征方程已知（或可求出）。

**核心步骤（必须背，一步都不能错）**：
1. 写出系统的闭环特征方程：$a_n s^n + a_{n-1} s^{n-1} + \\dots + a_1 s + a_0 = 0$（所有系数都不为 0，且同号）；
2. 构造劳斯表（按 $s$ 的降幂排列，共 $n+1$ 行）：

| $s^n$ | $a_n$ | $a_{n-2}$ | $a_{n-4}$ | $\\dots$ |
|-------|-------|----------|----------|--------|
| $s^{n-1}$ | $a_{n-1}$ | $a_{n-3}$ | $a_{n-5}$ | $\\dots$ |
| $s^{n-2}$ | $b_1$ | $b_2$ | $b_3$ | $\\dots$ |
| $s^{n-3}$ | $c_1$ | $c_2$ | $c_3$ | $\\dots$ |
| $\\dots$ | $\\dots$ | $\\dots$ | $\\dots$ | $\\dots$ |
| $s^0$ | $a_0$ | $|$ | $|$ | $|$ |

其中：
- $b_1 = \\frac{a_{n-1}a_{n-2} - a_n a_{n-3}}{a_{n-1}}$，$b_2 = \\frac{a_{n-1}a_{n-4} - a_n a_{n-5}}{a_{n-1}}$，以此类推；
- $c_1 = \\frac{b_1 a_{n-3} - a_{n-1} b_2}{b_1}$，$c_2 = \\frac{b_1 a_{n-5} - a_{n-1} b_3}{b_1}$，以此类推；
- 若某一行的第一个元素为 0，其他元素不为 0，用 $\\varepsilon$（无穷小正数）代替，继续计算。

3. **劳斯判据规则（必背）**：
   - 系统稳定的充要条件：
     ① 特征方程的所有系数同号（均为正或均为负，考研中均为正）；
     ② 劳斯表的第一列元素全部同号（均为正），无零元素。
   - 若不满足上述条件，系统不稳定：
     ① 特征方程系数异号 → 系统不稳定；
     ② 劳斯表第一列元素变号 → 变号次数 = 右半平面闭环极点的个数。

**特殊情况处理（高频）**：
1. **某一行所有元素都为 0（存在对称于原点的极点，如 $\\pm j\\omega$、$\\pm a$）**：
   - 取上一行的元素作为系数，构造辅助方程；
   - 对辅助方程求导，用导数方程的系数代替全零行；
   - 继续构造劳斯表，判断稳定性；
   - 对称极点可通过辅助方程求解。
2. **某一行第一个元素为 0，其他元素不为 0**：
   - 用 $\\varepsilon$（无穷小正数）代替该元素，继续构造劳斯表；
   - 若 $\\varepsilon\\to 0^+$ 时，第一列元素变号，说明存在右半平面极点，系统不稳定。

**通俗解读**：劳斯判据是"偷懒的判稳方法"，不用解方程求极点，只要看系数和劳斯表第一列，就能判断系统稳不稳，还能知道有几个极点在右半平面。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 特征方程的所有系数必须同号且不为 0，否则直接判断系统不稳定；
2. 构造劳斯表时，计算错误（比如 $b_1$、$c_1$ 的计算公式记混）；
3. 遇到全零行时，忘记构造辅助方程，直接判断系统不稳定；
4. 劳斯表第一列变号次数，是"右半平面极点个数"，不是"不稳定极点个数"（虚轴上的极点不算右半平面）。
</div>

**考点提示**：
1. 大题：用劳斯判据求系统稳定的增益 $K$ 范围（每年必考，分值 10 分左右）；
2. 选择题：用劳斯判据判断系统稳定性、求右半平面极点个数（高频）；
3. 填空题：劳斯表的构造、特殊情况处理。

---

### 3. 奈奎斯特（Nyquist）稳定判据 ⭐⭐⭐⭐

**核心作用**：从频域角度（开环频率特性）判断闭环系统的稳定性，还能分析系统的稳定裕度，适用于难以求解特征方程的复杂系统。

**核心原理**：通过开环频率特性曲线（奈奎斯特曲线），判断闭环系统右半平面极点的个数，进而判断稳定性。

**核心公式（必须背）**：
$$Z = P - 2N$$
其中：
- $Z$：闭环系统右半平面的极点数（$Z=0$ 时，系统稳定）；
- $P$：开环系统右半平面的极点数（已知，由开环传递函数确定）；
- $N$：奈奎斯特曲线包围 $(-1,j0)$ 点的圈数（逆时针包围为正，顺时针包围为负）。

**核心步骤（必背）**：
1. 写出系统的开环传递函数 $G(s)H(s)$，确定开环右半平面极点个数 $P$；
2. 绘制开环频率特性的奈奎斯特曲线（$\\omega$ 从 $0\\to\\infty$，再从 $\\infty\\to 0$，关于实轴对称）；
3. 数奈奎斯特曲线包围 $(-1,j0)$ 点的圈数 $N$（逆时针为正，顺时针为负）；
4. 代入公式 $Z = P - 2N$，若 $Z=0$，系统稳定；若 $Z\\neq0$，系统不稳定，$Z$ 为右半平面闭环极点个数。

**特殊情况处理（开环有积分环节，高频）**：
1. 若开环传递函数有 $v$ 个积分环节（$1/s^v$），则奈奎斯特曲线在 $\\omega\\to 0^+$ 时，需补画"半径为无穷大的圆弧"，圆弧的旋转方向为"逆时针旋转 $v\\times 90^\\circ$"（$v=1\\to 90^\\circ$，$v=2\\to 180^\\circ$，以此类推）；
2. 补画圆弧后，再数包围 $(-1,j0)$ 点的圈数 $N$。

**通俗解读**：奈奎斯特判据是"看曲线绕圈"——开环频率特性曲线绕 $(-1,j0)$ 点逆时针转 $N$ 圈，结合开环右半平面极点个数 $P$，就能判断闭环系统稳不稳。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 开环有积分环节时，忘记补画无穷大圆弧，导致 $N$ 计算错误；
2. $N$ 的符号记反：逆时针包围为正，顺时针包围为负；
3. 混淆 $P$ 的含义：$P$ 是"开环右半平面极点数"，不是闭环；
4. 奈奎斯特曲线是"$\\omega$ 从 $0\\to\\infty$，再从 $\\infty\\to 0$"，关于实轴对称，绘制时不能遗漏。
</div>

**考点提示**：
1. 分析题：根据奈奎斯特曲线判断系统稳定性（高频）；
2. 填空题：求临界稳定的增益 $K$（次高频）；
3. 大题：绘制奈奎斯特曲线，判断系统稳定性（可选，分值 10 分左右）。

---

### 4. 对数稳定判据（伯德图判稳）⭐⭐⭐⭐

**核心作用**：通过伯德图（幅频特性 + 相频特性）判断系统稳定性，比奈奎斯特判据更直观，适用于工程实际。

**核心原理**：根据伯德图上的"0dB 穿越频率"和"-180°穿越频率"，判断系统的稳定裕度，进而判断稳定性。

**核心概念（必背）**：
1. **0dB 穿越频率（$\\omega_c$）**：幅频特性曲线穿越 0dB 线（$20\\lg|G(j\\omega)|=0$）时的频率，反映系统的响应速度；
2. **-180°穿越频率（$\\omega_g$）**：相频特性曲线穿越 -180° 线（$\\angle G(j\\omega)=-180^\\circ$）时的频率，反映系统的相位滞后；
3. **相位裕度（$\\gamma$）**：在 $\\omega_c$ 处，使系统达到临界稳定所需的额外相位滞后量，公式：
   $$\\gamma = 180^\\circ + \\angle G(j\\omega_c)$$
4. **增益裕度（$G_m$）**：在 $\\omega_g$ 处，使系统达到临界稳定所需的增益倍数，公式：
   $$G_m = -20\\lg|G(j\\omega_g)| \\quad (dB)$$

**判稳规则（必背）**：
- 系统稳定的充要条件：相位裕度 $\\gamma > 0^\\circ$，且增益裕度 $G_m > 0dB$；
- 系统临界稳定：$\\gamma = 0^\\circ$，或 $G_m = 0dB$；
- 系统不稳定：$\\gamma < 0^\\circ$，或 $G_m < 0dB$。

**工程要求**：$\\gamma \\in [30^\\circ, 60^\\circ]$，$G_m \\in [6dB, 12dB]$，兼顾系统的稳定性和动态性能。

**通俗解读**：对数稳定判据是"看伯德图的两个关键点"——0dB 穿越频率处的相位裕度（安全余量），-180°穿越频率处的增益裕度（安全余量），两个余量都为正，系统才稳定。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 相位裕度是"开环指标"，不是闭环指标，仅用于单位负反馈系统；
2. 相位裕度 $\\gamma > 0^\\circ$ 才稳定，$\\gamma < 0^\\circ$ 不稳定，极易记反；
3. 增益裕度 $G_m > 0dB$ 才稳定，$G_m < 0dB$ 不稳定（$G_m$ 为正，说明还能增大增益，系统仍稳定）；
4. 若伯德图没有 0dB 穿越频率（幅频特性始终在 0dB 以下），则 $\\gamma$ 不存在，系统稳定；若没有 -180°穿越频率（相频特性始终在 -180°以上），则 $G_m$ 不存在，系统稳定。
</div>

**考点提示**：
1. 分析题：根据伯德图求相位裕度、增益裕度，判断系统稳定性（高频）；
2. 填空题：相位裕度、增益裕度的公式（次高频）。

---

### 5. 李雅普诺夫稳定判据（补充考点，极少考）⭐⭐

**核心作用**：适用于非线性系统、时变系统的稳定性判断（劳斯、奈奎斯特判据仅适用于线性定常系统）。

**核心原理（李雅普诺夫第二方法，直接法）**：
1. 构造一个正定的标量函数 $V(x)$（李雅普诺夫函数），满足 $V(x) > 0$（$x\\neq0$），$V(0)=0$；
2. 求 $V(x)$ 对时间的导数 $\\dot{V}(x)$：
   - 若 $\\dot{V}(x)$ 负定（$\\dot{V}(x) < 0$，$x\\neq0$），则系统渐近稳定；
   - 若 $\\dot{V}(x)$ 半负定（$\\dot{V}(x) \\leq 0$），且 $\\dot{V}(x) = 0$ 仅在 $x=0$ 时成立，则系统渐近稳定；
   - 若 $\\dot{V}(x)$ 正定（$\\dot{V}(x) > 0$），则系统不稳定；
   - 若 $\\dot{V}(x)$ 半正定（$\\dot{V}(x) \\geq 0$），且 $\\dot{V}(x) = 0$ 仅在 $x=0$ 时成立，则系统临界稳定。

**通俗解读**：李雅普诺夫判据是"构造一个能量函数"——若系统的能量随时间逐渐衰减（$\\dot{V}(x) < 0$），则系统稳定；若能量逐渐增加（$\\dot{V}(x) > 0$），则系统不稳定。

**易错点**：李雅普诺夫函数的构造没有通用方法，需根据系统特性灵活构造（极少考，仅作概念了解）。

**考点提示**：选择题判断李雅普诺夫判据的基本原理（极少考）。
</div>
"""
    },
    "四、稳态误差模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 四、稳态误差模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">必考 · 占比 10%+ · 核心是计算 + 系统型别</p>

---

### 1. 稳态误差的基本概念 ⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>专业定义</strong>：系统在典型输入信号（阶跃、斜坡、抛物线）作用下，当过渡过程结束后，输出稳态值与参考输入值的偏差，记为 $e_{ss}$，反映系统的稳态精度（精度越高，$e_{ss}$ 越小）。</p>
</div>

**误差信号的定义**：
1. 按输入定义（常用）：$E(s) = R(s) - C(s)$（输入信号减去输出信号）；
2. 按反馈定义：$E(s) = R(s) - B(s)$（输入信号减去反馈信号，$B(s)=G(s)H(s)R(s)$）；
3. 两种定义在单位负反馈系统（$H(s)=1$）中一致，考研默认单位负反馈。

**稳态误差的计算方法**：
1. **终值定理法（最常用）**：
   $$e_{ss} = \\lim_{s\\to 0} s \\cdot E(s) = \\lim_{s\\to 0} \\frac{sR(s)}{1+G(s)H(s)}$$
   适用条件：系统稳定（闭环极点全在左半平面），且 $sE(s)$ 的极点全在左半平面。

2. **误差系数法（仅适用于典型输入）**：
   - 位置误差系数 $K_p$：$K_p = \\lim_{s\\to 0} G(s)H(s)$（对应单位阶跃输入）；
   - 速度误差系数 $K_v$：$K_v = \\lim_{s\\to 0} sG(s)H(s)$（对应单位斜坡输入）；
   - 加速度误差系数 $K_a$：$K_a = \\lim_{s\\to 0} s^2G(s)H(s)$（对应单位抛物线输入）；
   - 稳态误差与误差系数的关系：
     - 单位阶跃输入：$e_{ss} = \\frac{1}{1+K_p}$
     - 单位斜坡输入：$e_{ss} = \\frac{1}{K_v}$
     - 单位抛物线输入：$e_{ss} = \\frac{1}{K_a}$

**通俗解读**：稳态误差就是系统"最后差多少"——比如输入是 10，输出稳态值是 9.8，稳态误差就是 0.2，误差越小，系统越精准。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 终值定理仅适用于系统稳定，若系统不稳定，无法用终值定理计算稳态误差；
2. 误差系数法仅适用于典型输入（阶跃、斜坡、抛物线），非典型输入需用终值定理；
3. 误差信号的两种定义，仅在单位负反馈系统中一致，非单位负反馈需注意区分。
</div>

**考点提示**：
1. 计算题：用终值定理 / 误差系数法求稳态误差（每年必考，分值 5-10 分）；
2. 选择题：判断不同系统型别、不同输入下的稳态误差（高频）；
3. 填空题：误差系数的公式。

---

### 2. 系统型别 ⭐⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<p><strong>专业定义</strong>：系统型别是由开环传递函数中积分环节（$1/s$）的个数决定的，记为 $v$（$v=0,1,2,\\dots$）：</p>
</div>

- $v=0$：0 型系统（开环无积分环节）；
- $v=1$：I 型系统（开环有 1 个积分环节）；
- $v=2$：II 型系统（开环有 2 个积分环节）；
- $v\\geq3$：III 型及以上系统（工程中极少用，考研不考）。

**系统型别的判定方法**：
1. 写出开环传递函数 $G(s)H(s)$，将其化为"尾 1 型"（分母中 $s$ 的最高次项系数为 1）；
2. 数分母中 $s$ 的次数（即积分环节的个数），即为系统型别 $v$。

示例：$G(s)H(s) = \\frac{K}{s(s^2+2s+3)} \\to$ 尾 1 型，分母有 1 个 $s \\to$ I 型系统（$v=1$）。

**典型输入下不同系统型别的稳态误差**：

<div style="font-size: 14px;">
<table style="width: 100%; border-collapse: collapse; margin: 15px 0;">
  <tr style="background: #fef5f7;">
    <th style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">系统型别（$v$）</th>
    <th style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">单位阶跃输入</th>
    <th style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">单位斜坡输入</th>
    <th style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">单位抛物线输入</th>
    <th style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">误差系数关系</th>
  </tr>
  <tr>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">0 型（$v=0$）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\frac{1}{1+K_p}$（有限值）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\infty$（无法跟踪）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\infty$（无法跟踪）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$K_p=K$，$K_v=0$，$K_a=0$</td>
  </tr>
  <tr style="background: #fef9f9;">
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">I 型（$v=1$）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = 0$（无静差）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\frac{1}{K_v}$（有限值）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\infty$（无法跟踪）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$K_p=\\infty$，$K_v=K$，$K_a=0$</td>
  </tr>
  <tr>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">II 型（$v=2$）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = 0$（无静差）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = 0$（无静差）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$e_{ss} = \\frac{1}{K_a}$（有限值）</td>
    <td style="border: 1px solid #f8e0e6; padding: 10px; text-align: center;">$K_p=\\infty$，$K_v=\\infty$，$K_a=K$</td>
  </tr>
</table>
</div>

**核心规律**：
1. 系统型别 $v$ 越高，稳态精度越高，能跟踪的输入类型越多；
2. 系统能跟踪的输入类型：$v\\geq$ 输入的阶数（阶跃输入阶数 0，斜坡 1，抛物线 2）；
   - 0 型系统：能跟踪 0 阶输入（阶跃），不能跟踪 1 阶及以上（斜坡、抛物线）；
   - I 型系统：能跟踪 0 阶、1 阶输入（阶跃、斜坡），不能跟踪 2 阶及以上（抛物线）；
   - II 型系统：能跟踪 0 阶、1 阶、2 阶输入（阶跃、斜坡、抛物线）。

**通俗解读**：
- 0 型系统："只能盯着固定目标"，不能跟移动目标；
- I 型系统："能盯着固定目标，也能跟匀速移动的目标"，不能跟加速移动的目标；
- II 型系统："能盯着固定目标，能跟匀速、加速移动的目标"，稳态精度最高。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 系统型别由"开环传递函数的积分环节个数"决定，不是闭环传递函数；
2. 积分环节越多，稳态精度越高，但系统的稳定性越差（需通过校正平衡）；
3. 0 型系统对阶跃输入有稳态误差，I 型系统对阶跃输入无稳态误差，II 型系统对阶跃、斜坡输入均无稳态误差，极易记混；
4. 误差系数 $K_p$、$K_v$、$K_a$ 仅与系统型别和开环增益 $K$ 有关，与输入无关。
</div>

**考点提示**：
1. 选择题：判断系统型别，或根据系统型别判断稳态误差（必考）；
2. 计算题：已知系统型别和开环增益，求典型输入的稳态误差（必考）；
3. 分析题：增加积分环节（提高系统型别）对稳态误差的影响（高频）。

---

### 3. 稳态误差的改进方法 ⭐⭐⭐⭐

**核心思路**：提高系统型别、增大开环增益，从而减小稳态误差，但需兼顾系统的稳定性。

**常用改进方法**：
1. **增大开环增益 $K$**：
   - 作用：减小稳态误差（对 0 型系统，$K$ 增大，$e_{ss}$ 减小；对 I 型系统，$K$ 增大，斜坡输入的 $e_{ss}$ 减小；对 II 型系统，$K$ 增大，抛物线输入的 $e_{ss}$ 减小）；
   - 缺点：增大 $K$ 会降低系统的稳定性（超调量增大，甚至不稳定）。

2. **提高系统型别（增加积分环节）**：
   - 作用：消除对应输入的稳态误差（比如增加 1 个积分环节，从 0 型→I 型，消除阶跃输入的稳态误差）；
   - 缺点：增加积分环节会降低系统的相位裕度，降低稳定性，甚至导致系统不稳定，需配合微分环节使用。

3. **采用复合控制（前馈 + 反馈）**：
   - 作用：在反馈控制的基础上，增加前馈环节，补偿输入信号的影响，进一步减小稳态误差，甚至消除稳态误差；
   - 适用场景：对稳态精度要求极高的系统（如精密机床、卫星控制系统）。

**通俗解读**：改进稳态误差的核心是"要么增大增益，要么增加积分环节"，但两者都会降低稳定性，所以需要通过校正环节（如 PD、PID）平衡稳态精度和动态稳定性。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. 增大开环增益 $K$ 不能无限制，否则系统会不稳定；
2. 增加积分环节只能消除"对应输入"的稳态误差，不能消除所有输入的稳态误差（比如增加 1 个积分环节，只能消除阶跃输入的稳态误差，不能消除斜坡输入的稳态误差）；
3. 复合控制的前馈环节不影响系统的稳定性，仅影响稳态误差。
</div>

**考点提示**：分析题：如何改进系统的稳态误差，兼顾稳定性和稳态精度（高频）。

---

### 4. 扰动稳态误差 ⭐⭐⭐

**定义**：系统在扰动信号（如负载变化、噪声干扰）作用下，输出稳态值与参考输入的偏差，记为 $e_{ssn}$，反映系统的抗干扰能力。

**计算方法（终值定理法）**：
假设扰动信号为 $N(s)$，则扰动误差信号 $E_n(s) = -\\frac{G_2(s)N(s)}{1+G_1(s)G_2(s)H(s)}$（$G_1$ 为前向通道前半部分，$G_2$ 为前向通道后半部分），稳态扰动误差：
$$e_{ssn} = \\lim_{s\\to 0} s \\cdot E_n(s) = \\lim_{s\\to 0} \\frac{-sG_2(s)N(s)}{1+G_1(s)G_2(s)H(s)}$$

**核心规律**：
1. 扰动稳态误差与扰动信号的位置、系统型别、开环增益有关；
2. 若扰动信号作用在积分环节之后，增加积分环节（提高系统型别），不能减小扰动稳态误差；
3. 若扰动信号作用在积分环节之前，增加积分环节（提高系统型别），可以减小甚至消除扰动稳态误差。

**通俗解读**：扰动稳态误差是"系统受到干扰后，最后差多少"，抗干扰能力越强，$e_{ssn}$ 越小。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：扰动稳态误差的计算，需注意扰动信号的作用位置，不同位置的计算方法不同。
</div>

**考点提示**：计算题：求扰动信号作用下的稳态误差（次高频，分值 5-8 分）。
</div>
"""
    },
    "五、PID控制模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 五、PID控制模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">考研/复试双重点 · 占比 15%+ · 核心是参数作用 + 整定</p>

---

### 1. PID控制器的基本原理 ⭐⭐⭐⭐⭐

<div style="background: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; border-left: 4px solid #f8c4c4;">
<strong>定义</strong>：PID 控制器是一种线性控制器，通过比例（P）、积分（I）、微分（D）三个环节的组合，根据误差信号 $e(t)$，输出控制信号 $u(t)$，实现对系统的控制，是工程中最常用的控制器。
</div>

**时域表达式**：
$$u(t) = K_p e(t) + K_i \\int_0^t e(\\tau) d\\tau + K_d \\frac{de(t)}{dt}$$

**复频域传递函数**：
$$G_c(s) = K_p + \\frac{K_i}{s} + K_d s = \\frac{K_d s^2 + K_p s + K_i}{s}$$

其中：
- $K_p$：比例系数（决定比例环节的作用强度）；
- $K_i$：积分系数（决定积分环节的作用强度）；
- $K_d$：微分系数（决定微分环节的作用强度）。

**各环节的独立作用**：
1. **比例环节（P）**：
   - 传递函数：$G_c(s) = K_p$
   - 作用：即时跟随误差，误差大输出大；提高系统响应速度；减小稳态误差
   - 缺点：不能消除稳态误差；$K_p$ 过大会剧烈振荡、不稳定
   - 通俗理解：现在差多少，立刻补多少，但永远有小尾巴

2. **积分环节（I）**：
   - 传递函数：$G_c(s) = \\frac{K_i}{s}$
   - 作用：消除阶跃稳态误差；提高系统型别（0 型→I 型）
   - 缺点：相位滞后，降低稳定性；响应变慢，易超调
   - 通俗理解：过去累计差多少，慢慢补回来，直到完全没误差

3. **微分环节（D）**：
   - 传递函数：$G_c(s) = K_d s$
   - 作用：预测误差变化趋势；抑制超调；增加阻尼，提高稳定性
   - 缺点：放大高频噪声；不能单独使用
   - 通俗理解：看误差变化快不快，提前刹车，防止冲过头

**组合控制**：
- **PD控制**：$P+D$，提高响应速度，增加阻尼，减小超调，不改变稳态误差
- **PI控制**：$P+I$，消除阶跃稳态误差，提高稳态精度，但稳定性下降
- **PID控制**：$P+I+D$，快（P）、准（I）、稳（D），工程最常用

**PID参数对系统影响**：
- $K_p \\uparrow$：响应变快，稳态误差减小，易振荡
- $K_i \\uparrow$：消除静差更快，稳定性变差
- $K_d \\uparrow$：阻尼变大，超调减小，抗噪声变差

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点</strong>：
1. I 只能消除阶跃误差，不能消除斜坡 / 抛物线；
2. D 不能单独用；
3. PI 会让系统更不稳定，必须谨慎；
4. PID 不改变系统零点极点结构，只是"修正"。
</div>

**考点提示**：
1. 大题：设计 PID 控制器，使系统性能达标（每年必考，分值 10-15 分）；
2. 选择题：PID 参数变化对系统性能的影响（高频）。

---

### 2. PID参数整定方法 ⭐⭐⭐⭐

**工程整定方法（无需模型）**：
1. **试凑法**：
   - 先调 $K_p$，使系统响应快，略有超调；
   - 再加 $K_i$，消除稳态误差；
   - 最后加 $K_d$，抑制超调，提高稳定性。

2. **Ziegler-Nichols 法（齐格勒-尼科尔斯法）**：
   - 先设 $K_i=0$，$K_d=0$，增大 $K_p$ 使系统临界振荡（等幅振荡）；
   - 记录临界增益 $K_{cr}$ 和临界周期 $T_{cr}$；
   - 按表格计算 PID 参数。

**考点提示**：
1. 简答题：PID 参数整定方法（高频）；
2. 填空题：Ziegler-Nichols 法步骤（次高频）。

---

### 3. 常用校正控制器 ⭐⭐⭐

除了 PID 外，常用的校正控制器还有：
1. **超前校正**：提供相位超前，提高相位裕度，提高稳定性，加快响应，不改善稳态误差
2. **滞后校正**：提高低频增益，减小稳态误差，不明显影响动态
3. **滞后-超前校正**：兼顾稳态 + 动态

**考点提示**：选择题：不同校正装置的特点（高频）。
</div>
"""
    },
    "六、根轨迹模块": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 六、根轨迹模块</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">大题必考 · 占比 15%+ · 核心是绘制 + 分析</p>

---

### 1. 根轨迹定义 ⭐⭐⭐⭐

**定义**：开环增益 $K$ 从 $0 \\to \\infty$ 时，闭环极点在 $s$ 平面运动的轨迹。

**核心思想**：通过开环零极点，分析闭环极点随 $K$ 变化的规律，无需求解特征方程。

---

### 2. 根轨迹 9 条绘制规则 ⭐⭐⭐⭐⭐

1. **根轨迹对称于实轴**（因为复数极点共轭成对出现）；
2. **起点**：开环极点（$K=0$）；
3. **终点**：开环零点（无穷远零点也算）；
4. **实轴上根轨迹**：右侧极点个数 + 零点个数 = 奇数；
5. **渐近线中心**：
   $$\\sigma_a = \\frac{\\sum p_i - \\sum z_i}{n - m}$$
6. **渐近线角度**：
   $$\\frac{(2k+1)\\pi}{n - m}$$
7. **分离 / 会合点**：$\\frac{dK}{ds} = 0$；
8. **与虚轴交点**：劳斯判据求临界 $K$；
9. **出射角、入射角**：极点出射角、零点入射角公式。

<div style="background: #fff5e6; padding: 15px; border-radius: 8px; margin: 15px 0;">
<strong>⚠️ 易错点（实轴根轨迹）</strong>：
口诀：右边极零总数为奇数，才有根轨迹（90% 考生这里丢分）。
</div>

---

### 3. 根轨迹用途 ⭐⭐⭐

- 看稳定性；
- 看动态性能；
- 选合适增益 $K$；
- 设计校正装置。

**考点提示**：
1. 大题：绘制根轨迹（每年必考，分值 10-15 分）；
2. 选择题：根轨迹规则判断（高频）。
</div>
"""
    },
    "七、频域分析/伯德图": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 七、频域分析/伯德图</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">每年必考 · 占比 15%+ · 核心是伯德图 + 稳定裕度</p>

---

### 1. 频率特性定义 ⭐⭐⭐

$$G(j\\omega) = |G(j\\omega)| e^{j\\varphi(\\omega)}$$

**物理意义**：线性系统在正弦输入下，稳态输出的幅值比和相位差，随频率 $\\omega$ 变化的特性。

---

### 2. 伯德图组成 ⭐⭐⭐⭐

- **幅频**：$20\\lg|G(j\\omega)|$（dB）；
- **相频**：$\\angle G(j\\omega)$（度）。

---

### 3. 典型环节斜率 ⭐⭐⭐⭐⭐

| 典型环节       | 幅频斜率（dB/dec） |
|----------------|------------------|
| 比例           | 0                |
| 积分           | -20              |
| 一阶惯性       | -20（转折后）    |
| 一阶微分       | +20（转折后）    |
| 二阶振荡       | -40（转折后）    |

---

### 4. 稳定裕度 ⭐⭐⭐⭐⭐

**相位裕度**：
$$\\gamma = 180^\\circ + \\varphi(\\omega_c)$$

**增益裕度**：
$$GM = -20\\lg|G(j\\omega_g)|$$

**判稳**：
- $\\gamma > 0^\\circ$，稳定；
- $\\gamma < 0^\\circ$，不稳定。

**工程理想**：
$$\\gamma \\in [30^\\circ, 60^\\circ]$$

---

### 5. 频域与时域关系 ⭐⭐⭐

- $\\omega_c \\uparrow$ → 响应更快；
- $\\gamma \\uparrow$ → 阻尼 $\\zeta \\uparrow$ → 超调 $\\downarrow$；
- $M_r \\uparrow$ → 超调 $\\uparrow$。

**考点提示**：
1. 大题：绘制伯德图，求稳定裕度（每年必考，分值 10-15 分）；
2. 选择题：稳定裕度判断、频域与时域关系（高频）。
</div>
"""
    },
    "八、系统校正": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 八、系统校正</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">大题必考 · 占比 10%+ · 核心是超前/滞后校正</p>

---

### 1. 超前校正 ⭐⭐⭐⭐

- 提供相位超前；
- 提高 $\\gamma$，提高稳定性；
- 加快响应；
- 不改善稳态误差。

---

### 2. 滞后校正 ⭐⭐⭐⭐

- 提高低频增益；
- 减小稳态误差；
- 不明显影响动态。

---

### 3. 滞后 - 超前校正 ⭐⭐⭐

- 兼顾稳态 + 动态。

---

### 4. 校正步骤 ⭐⭐⭐⭐

1. 绘制原系统伯德图，求 $\\gamma$；
2. 选校正装置，计算参数；
3. 验证校正后 $\\gamma$；
4. 若不满足，重新调整。

**考点提示**：
1. 大题：设计超前/滞后校正装置（每年必考，分值 10 分左右）；
2. 选择题：不同校正装置的特点（高频）。
</div>
"""
    },
    "九、离散控制": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 九、离散控制</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">部分院校考 · 核心是 z 变换</p>

---

### 1. z 变换 ⭐⭐⭐

常用 z 变换对需记牢：
- 单位阶跃：$\\frac{z}{z-1}$；
- 单位斜坡：$\\frac{Tz}{(z-1)^2}$；
- 指数函数：$\\frac{z}{z-e^{-aT}}$。

---

### 2. 离散系统稳定判据 ⭐⭐⭐

- 所有闭环极点在单位圆内（$|z| < 1$）。

---

### 3. 稳态误差 ⭐⭐⭐

类似连续系统，注意是 z 域终值定理。

---

### 4. 离散化方法 ⭐⭐

- 零阶保持；
- 双线性变换（Tustin）。

**考点提示**：
1. 计算题：z 变换、离散系统稳定判断（部分院校考）；
2. 填空题：常用 z 变换对（高频）。
</div>
"""
    },
    "十、非线性系统": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 十、非线性系统</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">部分院校考 · 核心是描述函数法</p>

---

### 1. 常见非线性 ⭐⭐

- 饱和；
- 死区；
- 间隙；
- 继电器。

---

### 2. 描述函数法 ⭐⭐⭐

**基本思想**：把非线性环节等效为"可变增益"，近似线性化分析。

**条件**：系统结构要求：线性部分低通特性好（高频衰减快）。

**自激振荡判断**：
- $N(A)$ 与 $-1/N(A)$ 曲线交点 → 看是稳定极限环还是不稳定。

---

### 3. 相平面法 ⭐⭐

画 $e-\\dot{e}$ 相轨迹，看运动趋势。

**考点提示**：
1. 分析题：用描述函数法判断自激振荡（部分院校考）；
2. 选择题：常见非线性的特点（高频）。
</div>
"""
    },
    "十一、现代控制基础": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 十一、现代控制基础</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">部分院校考 · 核心是状态空间</p>

---

### 1. 状态空间表达式 ⭐⭐⭐

$$
\\begin{cases}
\\dot{x} = Ax + Bu \\\\
y = Cx + Du
\\end{cases}
$$

---

### 2. 能控性 ⭐⭐⭐

$$\\text{rank}[B\\ AB\\ A^2B\\dots] = n$$

---

### 3. 能观性 ⭐⭐⭐

$$\\text{rank}\\begin{bmatrix} C \\\\ CA \\\\ CA^2 \\\\ \\vdots \\end{bmatrix} = n$$

---

### 4. 极点配置 ⭐⭐

用状态反馈任意配置极点（只要系统能控）。

---

### 5. 李雅普诺夫直接法 ⭐⭐

找正定 $V(x)$，看 $\\dot{V}(x)$ 是否负定。

**考点提示**：
1. 计算题：能控性、能观性判断（部分院校考）；
2. 填空题：状态空间表达式（高频）。
</div>
"""
    },
    "十二、考前速记/填空必背": {
        "content": """
<div style="padding: 20px; background: linear-gradient(135deg, #fef9f9 0%, #fff5f5 100%); border-radius: 16px;">
<h2 style="color: #e68a8a; text-align: center; margin-bottom: 30px;">📚 十二、考前速记/填空必背</h2>
<p style="text-align: center; color: #666; font-size: 16px; margin-bottom: 30px;">100% 高频口袋版 · 记牢就能拿分</p>

---

### 1. 记住这些结论（直接写在答题卡上）

1. **稳定**：闭环极点全在左半 $s$ 平面（现代：全在单位圆内）；
2. **稳态误差**：先看型别 $v$，再看输入阶数；
3. **根轨迹实轴段**：右侧极零数之和为奇数；
4. **伯德图积分环节**：-20 dB/dec，相位 $-90^\\circ$；
5. **相位裕度 $\\gamma$**：$180^\\circ + \\varphi(\\omega_c)$，工程要求 $30^\\circ \\sim 60^\\circ$；
6. **PID 作用**：P 快、I 准、D 稳；
7. **超前校正**：提相位、提稳定、提速度；
8. **滞后校正**：提低频、减误差；
9. **离散稳定**：极点在单位圆内；
10. **能控矩阵**：$[B\\ AB\\ A^2B\\dots]$，满秩为能控。

---

### 2. 考试技巧（拉分细节）

- 画根轨迹：务必标箭头、渐近线、分离点；
- 写伯德图：每个转折频率标清楚，斜率标在旁边；
- 计算题：先写公式，再代数值，一步不跳；
- 判稳：优先用劳斯，其次用奈奎斯特；
- 证明题：先写定义，再推结论，逻辑清晰。

---

### 3. 5 分钟就能做对的题

- 判断稳定：看特征方程系数是否全正，再用劳斯判据；
- 求稳态误差：先判稳，再看型别，最后用终值定理；
- 画根轨迹：按 9 条规则，一步步来，不着急；
- 伯德图求裕度：先找 $\\omega_c$，再算 $\\gamma$，最后找 $\\omega_g$ 算 $GM$。

</div>
"""
    }
}

# ---------------------- 专业术语库和易错点提醒（保留兼容） ----------------------
TERMINOLOGY = {
    "典型输入信号": {
        "terms": {
            "阶跃函数": {
                "definition": "最常用的测试输入信号，表示系统在某一时刻突然加上一个恒定值的输入。",
                "formula": "$r(t) = \\begin{cases} 0, & t<0 \\\\ R, & t\\geq 0 \\end{cases}$（单位阶跃：$R=1$）",
                "example": "单位阶跃函数的拉普拉斯变换：$R(s) = \\frac{1}{s}$"
            },
            "斜坡函数": {
                "definition": "输入信号随时间线性增长，用于测试系统跟踪速度跟踪能力。",
                "formula": "$r(t) = \\begin{cases} 0, & t<0 \\\\ Rt, & t\\geq 0 \\end{cases}$（单位斜坡：$R=1$）",
                "example": "单位斜坡函数的拉普拉斯变换：$R(s) = \\frac{1}{s^2}$"
            },
            "抛物线函数": {
                "definition": "输入信号随时间二次增长，用于测试系统加速度跟踪能力。",
                "formula": "$r(t) = \\begin{cases} 0, & t<0 \\\\ \\frac{1}{2}Rt^2, & t\\geq 0 \\end{cases}$",
                "example": "单位抛物线函数的拉普拉斯变换：$R(s) = \\frac{1}{s^3}$"
            },
            "脉冲函数": {
                "definition": "理想的瞬时脉冲，面积为1，用于测试系统的脉冲响应。",
                "formula": "$\\delta(t) = \\begin{cases} \\infty, & t=0 \\\\ 0, & t\\neq 0 \\end{cases}$，且$\\int_{-\\infty}^{+\\infty} \\delta(t)dt = 1$",
                "example": "单位脉冲函数的拉普拉斯变换：$\\Delta(s) = 1$"
            },
            "正弦函数": {
                "definition": "用于频率响应分析的输入信号。",
                "formula": "$r(t) = A\\sin(\\omega t)$",
                "example": "正弦函数的拉普拉斯变换：$R(s) = \\frac{A\\omega}{s^2 + \\omega^2}$"
            }
        },
        "common_mistakes": [
            "💡 技巧：阶跃函数用于测试系统的阶跃响应，斜坡函数用于测试系统的速度跟踪能力",
            "⚠️ 难点：不同系统型别决定了系统能跟踪的输入类型"
        ]
    },
    "劳斯判据": {
        "terms": {
            "劳斯表构造": {
                "definition": "根据特征方程系数构造劳斯阵列，用于判断系统稳定性。",
                "formula": "特征方程：$a_ns^n + a_{n-1}s^{n-1} + \\dots + a_1s + a_0 = 0$",
                "example": "劳斯表第一行：$a_n, a_{n-2}, a_{n-4}, \\dots$，第二行：$a_{n-1}, a_{n-3}, a_{n-5}, \\dots$"
            },
            "劳斯判据": {
                "definition": "系统稳定的充要条件是劳斯表第一列所有元素均为正。",
                "formula": "第一列元素符号改变次数等于右半s平面极点个数",
                "example": "若第一列有0，说明系统不稳定"
            },
            "特殊情况处理": {
                "definition": "劳斯表某行第一个元素为0的处理方法。",
                "formula": "用小正数ε代替0继续计算，或构造辅助多项式求导",
                "example": "辅助多项式：由上一行构造，求导后系数代替全零行"
            },
            "临界稳定判断": {
                "definition": "当劳斯表出现全零行时，系统临界稳定。",
                "formula": "辅助多项式的根就是纯虚根（共轭成对）",
                "example": "辅助多项式次数为偶数"
            }
        },
        "common_mistakes": [
            "❌ 易错点：构造劳斯表时，必须从最高次到最低次排列系数！",
            "⚠️ 难点：特殊情况处理时要记住用ε代替或辅助多项式",
            "💡 技巧：第一列符号改变次数就是右半平面极点个数"
        ]
    },
    "根轨迹": {
        "terms": {
            "开环极点": {
                "definition": "开环传递函数的极点，即分母多项式的根。开环极点的位置决定了根轨迹的起点。",
                "formula": "令分母为0：$D(s) = 0$，解得的根就是开环极点",
                "example": "对于$G(s) = \\frac{K}{s(s+1)(s+2)}$，开环极点为$s=0, -1, -2$"
            },
            "开环零点": {
                "definition": "开环传递函数的零点，即分子多项式的根。开环零点的位置决定了根轨迹的终点。",
                "formula": "令分子为0：$N(s) = 0$，解得的根就是开环零点",
                "example": "对于$G(s) = \\frac{K(s+3)}{s(s+1)(s+2)}$，开环零点为$s=-3$"
            },
            "分离/汇合点": {
                "definition": "根轨迹在实轴上分离（离开实轴）或汇合（回到实轴）的点。在分离/汇合点，根轨迹有重根。",
                "formula": "满足$\\frac{d}{ds}\\left[\\frac{N(s)}{D(s)}\\right] = 0$的实根",
                "example": "二阶系统的分离点通常在两个实极点之间"
            },
            "渐近线": {
                "definition": "当增益K→∞时，根轨迹趋近于的直线。渐近线帮助我们理解根轨迹在远区的走向。",
                "formula": "渐近线角度：$\\varphi = \\frac{180^\\circ (2k+1)}{n-m}$，渐近线中心：$\\sigma = \\frac{\\sum p_i - \\sum z_i}{n-m}$",
                "example": "三阶系统有3条渐近线，角度分别为60°、180°、300°"
            },
            "临界增益": {
                "definition": "系统由稳定变为不稳定（或反之）时的开环增益值。此时根轨迹穿过虚轴。",
                "formula": "令$s=j\\omega$代入特征方程，求解K和ω",
                "example": "当$K=K_{cr}$时，系统临界稳定，有一对纯虚根"
            },
            "主导极点": {
                "definition": "高阶系统中，离虚轴最近的一对共轭复极点（或实极点）对系统动态性能起主导作用，称为主导极点。",
                "formula": "满足：其他极点的实部绝对值是主导极点实部绝对值的3~5倍以上",
                "example": "对于极点$s=-1±j2$和$s=-10$，主导极点是$s=-1±j2$"
            },
            "根轨迹绘制法则": {
                "definition": "180°根轨迹绘制的基本法则。",
                "formula": "幅角条件：$\\angle G(s)H(s) = \\pm 180^\\circ (2k+1)$，幅值条件：$|G(s)H(s)| = 1$",
                "example": "实轴上根轨迹段：右侧零极点个数和为奇数"
            },
            "出射角/入射角": {
                "definition": "根轨迹从开环复极点出发的角度（出射角）或到达开环复零点的角度（入射角）。",
                "formula": "出射角：$\\theta_p = 180^\\circ + \\sum \\angle (z_i-p) - \\sum \\angle (p_j-p)$",
                "example": "入射角：$\\theta_z = 180^\\circ + \\sum \\angle (p_i-z) - \\sum \\angle (z_j-z)$"
            }
        },
        "common_mistakes": [
            "❌ 易错点：分离点计算时容易忘记筛选实根！分离点必须在实轴上的根轨迹段上",
            "⚠️ 难点：判断哪些实轴段是根轨迹段（右侧零极点个数和为奇数）",
            "💡 技巧：渐近线中心只与极点和零点的位置有关，与增益K无关"
        ]
    },
    "奈奎斯特判据": {
        "terms": {
            "奈奎斯特曲线": {
                "definition": "当s沿奈奎斯特围道绕行一周时，$G(j\\omega)H(j\\omega)$在复平面上的轨迹。",
                "formula": "奈奎斯特围道：包围右半s平面（不经过任何开环极点）",
                "example": "ω从0到+∞绘制$G(j\\omega)$，再从-∞到0绘制对称部分"
            },
            "奈奎斯特判据": {
                "definition": "闭环系统稳定的充要条件：奈奎斯特曲线逆时针包围(-1,j0)点的次数等于开环右半极点个数。",
                "formula": "$Z = P - 2N$，其中Z为闭环右极点个数，P为开环右极点个数，N为逆时针包围次数",
                "example": "若P=0，则奈奎斯特曲线不包围(-1,j0)点则系统稳定"
            },
            "稳定裕度": {
                "definition": "衡量系统相对稳定性的指标。",
                "formula": "相位裕度$\\gamma$、增益裕度$h$",
                "example": "相位裕度：截止频率处相位距-180°的裕度"
            }
        },
        "common_mistakes": [
            "❌ 易错点：奈奎斯特判据要考虑开环右极点个数P！",
            "⚠️ 难点：奈奎斯特曲线绘制时要注意对称性",
            "💡 技巧：Z=P-2N，Z=0系统稳定"
        ]
    },
    "伯德图": {
        "terms": {
            "截止频率": {
                "definition": "幅频特性穿过0dB线时的频率，也叫增益穿越频率。截止频率反映了系统的带宽。",
                "formula": "$|G(j\\omega_c)| = 1$（即20lg|G(jω_c)| = 0dB）",
                "example": "截止频率越大，系统响应速度越快"
            },
            "相位裕度": {
                "definition": "在截止频率处，相位距-180°还有多少裕度。相位裕度是衡量系统相对稳定性的重要指标。",
                "formula": "$\\gamma = 180^\\circ + \\angle G(j\\omega_c)$",
                "example": "相位裕度通常要求30°~60°，工程上常用45°左右"
            },
            "增益裕度": {
                "definition": "相位为-180°时，幅值距0dB还有多少裕度（以dB表示）。增益裕度也是衡量系统相对稳定性的指标。",
                "formula": "$GM = -20\\lg|G(j\\omega_{180})|$",
                "example": "增益裕度通常要求>6dB"
            },
            "转折频率": {
                "definition": "伯德图渐近线斜率发生变化的频率，对应系统的极点或零点的频率值。",
                "formula": "对于极点$s=-a$，转折频率为$\\omega = |a|$ rad/s",
                "example": "每个惯性环节在转折频率处使渐近线斜率变化-20dB/dec"
            },
            "伯德图渐近线": {
                "definition": "对数幅频特性的渐近线近似。",
                "formula": "比例环节：0dB/dec；积分环节：-20dB/dec；微分环节：+20dB/dec",
                "example": "惯性环节：转折频率后-20dB/dec"
            }
        },
        "common_mistakes": [
            "❌ 易错点：相位裕度是180°加上截止频率处的相位，不是直接取相位值！",
            "⚠️ 难点：绘制伯德图渐近线时，要注意每个环节的斜率变化（+20/-20dB/dec）",
            "💡 技巧：最小相位系统的相位裕度>0°且增益裕度>0dB时，系统稳定"
        ]
    },
    "二阶系统": {
        "terms": {
            "阻尼比": {
                "definition": "描述系统阻尼程度的参数，决定了系统响应的振荡特性和收敛速度。",
                "formula": "标准形式：$G(s) = \\frac{\\omega_n^2}{s^2 + 2\\zeta\\omega_n s + \\omega_n^2}$，其中$\\zeta$为阻尼比",
                "example": "$\\zeta=0.707$时，超调量约4.3%，综合性能最优"
            },
            "自然频率": {
                "definition": "系统无阻尼时的振荡频率，决定了系统响应的固有速度。",
                "formula": "$\\omega_n = \\sqrt{\\frac{K}{m}}$（机械系统）或由传递函数分母系数计算",
                "example": "自然频率越大，系统响应速度越快"
            },
            "超调量": {
                "definition": "阶跃响应中，输出最大值超过稳态值的百分比。是衡量系统相对稳定性的指标。",
                "formula": "$\\sigma\\% = e^{-\\zeta\\pi/\\sqrt{1-\\zeta^2}} \\times 100\\%$",
                "example": "$\\zeta=0.5$时，超调量约16.3%"
            },
            "调节时间": {
                "definition": "阶跃响应进入并保持在稳态值±2%（或±5%）误差带内所需的最短时间。",
                "formula": "2%误差带：$t_s \\approx \\frac{4}{\\zeta\\omega_n}$，5%误差带：$t_s \\approx \\frac{3}{\\zeta\\omega_n}$",
                "example": "调节时间越短，系统响应越快"
            },
            "稳态增益": {
                "definition": "系统对阶跃输入的稳态输出与输入幅值之比，由传递函数的直流增益决定。",
                "formula": "$K = \\frac{\\text{分子常数项}}{\\text{分母常数项}}$（对于标准二阶系统，K=1）",
                "example": "改变分子会改变稳态增益，但不会改变动态性能指标（ωₙ、ζ、σ%、tₛ等）"
            },
            "欠阻尼/临界阻尼/过阻尼": {
                "definition": "根据阻尼比ζ的取值范围分类。",
                "formula": "欠阻尼：$0<\\zeta<1$；临界阻尼：$\\zeta=1$；过阻尼：$\\zeta>1$",
                "example": "欠阻尼系统有振荡，过阻尼系统无振荡"
            },
            "峰值时间": {
                "definition": "阶跃响应达到第一个峰值所需的时间。",
                "formula": "$t_p = \\frac{\\pi}{\\omega_n\\sqrt{1-\\zeta^2}}$",
                "example": "峰值时间只存在于欠阻尼系统"
            },
            "上升时间": {
                "definition": "阶跃响应从稳态值的10%上升到90%所需的时间。",
                "formula": "$t_r = \\frac{\\pi - \\beta}{\\omega_d}$，其中$\\beta = \\arccos\\zeta$",
                "example": "上升时间反映系统的快速性"
            }
        },
        "common_mistakes": [
            "❌ 易错点：超调量只与阻尼比ζ有关，与自然频率ωₙ无关！",
            "❌ 易错点：改变分子只会改变稳态增益，不会改变ωₙ、ζ、超调量、调节时间等动态性能！",
            "⚠️ 难点：区分欠阻尼（0<ζ<1）、临界阻尼（ζ=1）、过阻尼（ζ>1）三种情况",
            "💡 技巧：工程上常用ζ≈0.707，兼顾超调量和响应速度"
        ]
    },
    "稳态误差": {
        "terms": {
            "终值定理": {
                "definition": "计算稳态误差的重要定理，利用s域分析求时域t→∞的极限。",
                "formula": "$\\lim_{t\\to\\infty} f(t) = \\lim_{s\\to 0} sF(s)$（要求sF(s)的极点都在左半平面）",
                "example": "求阶跃响应的稳态值：$y(\\infty) = \\lim_{s\\to 0} sT(s)\\cdot\\frac{1}{s}$"
            },
            "误差系数": {
                "definition": "位置误差系数Kp、速度误差系数Kv、加速度误差系数Ka，分别对应阶跃、斜坡、抛物线输入。",
                "formula": "$K_p = \\lim_{s\\to 0} G(s)$，$K_v = \\lim_{s\\to 0} sG(s)$，$K_a = \\lim_{s\\to 0} s^2G(s)$",
                "example": "I型系统的Kv等于开环增益K"
            },
            "单位负反馈": {
                "definition": "反馈通道传递函数为1的反馈系统，是最常用的反馈形式。",
                "formula": "闭环传递函数：$T(s) = \\frac{G(s)}{1+G(s)}$",
                "example": "单位负反馈系统中，误差信号e(t) = r(t) - y(t)"
            },
            "系统型别与稳态误差": {
                "definition": "系统型别决定了跟踪不同输入的稳态误差。",
                "formula": "0型：阶跃$e_{ss}=\\frac{R}{1+K_p}$，斜坡$e_{ss}=\\infty$；I型：阶跃$e_{ss}=0$，斜坡$e_{ss}=\\frac{R}{K_v}$",
                "example": "II型：阶跃和斜坡$e_{ss}=0$，抛物线$e_{ss}=\\frac{R}{K_a}$"
            }
        },
        "common_mistakes": [
            "❌ 易错点：应用终值定理前，必须确保sE(s)的极点都在左半平面！",
            "❌ 易错点：0型系统对斜坡输入的稳态误差是无穷大，不是有限值！",
            "⚠️ 难点：区分开环系统型别和闭环系统性能——型别是开环特性",
            "💡 技巧：记住型别/输入类型/稳态误差关系是必考点"
        ]
    },
    "系统校正": {
        "terms": {
            "串联超前校正": {
                "definition": "在系统前向通道中串联超前校正装置，提供超前相位，提高相角裕度，改善动态性能。",
                "formula": "超前校正传递函数：$G_c(s) = \\frac{1+aTs}{1+Ts}$，其中$a>1$",
                "example": "最大超前相位：$\\phi_m = \\arcsin\\frac{a-1}{a+1}$"
            },
            "串联滞后校正": {
                "definition": "在系统前向通道中串联滞后校正装置，提高稳态精度，保持动态性能。",
                "formula": "滞后校正传递函数：$G_c(s) = \\frac{1+Ts}{1+aTs}$，其中$a>1$",
                "example": "滞后校正主要利用其高频衰减特性"
            },
            "超前-滞后校正": {
                "definition": "结合超前和滞后校正的优点，同时改善动态性能和稳态精度。",
                "formula": "超前-滞后校正传递函数：$G_c(s) = \\frac{(1+aT_1s)(1+T_2s)}{(1+T_1s)(1+aT_2s)}$，其中$a>1$",
                "example": "超前部分改善动态，滞后部分提高稳态精度"
            },
            "PID校正": {
                "definition": "比例-积分-微分校正，是最常用的校正方式。",
                "formula": "PID控制器：$G_c(s) = K_p + \\frac{K_i}{s} + K_d s$",
                "example": "P提高响应，I消除稳态误差，D提供超前相位"
            },
            "校正目标": {
                "definition": "系统校正的性能指标要求。",
                "formula": "稳态精度：稳态误差$e_{ss}$；动态性能：超调量$\\sigma\\%$、调节时间$t_s$；稳定性：相角裕度$\\gamma$、增益裕度$h$",
                "example": "通常要求相角裕度30°~60°，增益裕度>6dB"
            }
        },
        "common_mistakes": [
            "❌ 易错点：超前校正会降低系统的低频增益，可能需要补充增益补偿！",
            "⚠️ 难点：选择校正装置时要综合考虑性能要求",
            "💡 技巧：需要快速性要求高用超前校正，稳态精度要求高用滞后校正"
        ]
    },
    "PID控制": {
        "terms": {
            "比例控制P": {
                "definition": "比例控制器，输出与误差成正比。作用是提高响应速度、减小稳态误差，但会增大超调量。",
                "formula": "$u(t) = K_p e(t)$",
                "example": "增大Kp，响应变快，超调变大，稳态误差减小（但不能消除）"
            },
            "积分控制I": {
                "definition": "积分控制器，输出与误差的积分成正比。作用是消除稳态误差，但会降低稳定性，易引起振荡。",
                "formula": "$u(t) = K_i \\int_0^t e(\\tau) d\\tau$",
                "example": "增加积分环节，系统型别升高，稳态误差减小或消除"
            },
            "微分控制D": {
                "definition": "微分控制器，输出与误差的变化率成正比。作用是提供超前相位，增大相位裕度，减小超调量，提高稳定性。",
                "formula": "$u(t) = K_d \\frac{de(t)}{dt}$",
                "example": "微分环节对高频噪声敏感，实际中常用不完全微分"
            },
            "系统型别": {
                "definition": "开环传递函数中位于原点的极点个数（即积分环节的个数），决定了系统跟踪不同输入的能力。",
                "formula": "型别ν = 分母末尾零的个数",
                "example": "0型系统：阶跃输入有稳态误差；I型系统：阶跃输入无误差，斜坡输入有误差；II型系统：阶跃和斜坡输入无误差，抛物线输入有误差"
            },
            "PID参数整定口诀": {
                "definition": "PID参数整定的经验口诀。",
                "formula": "先P后I最后D，由小到大慢慢调",
                "example": "参数整定：先比例，后积分，最后加微分"
            }
        },
        "common_mistakes": [
            "❌ 易错点：积分环节虽然能消除稳态误差，但会降低相位裕度，可能导致系统不稳定！",
            "❌ 易错点：纯P控制（Ki=0,Kd=0）且Kp=1时，校正前后开环传递函数相同，性能指标无变化！",
            "⚠️ 难点：理解系统型别与输入类型的关系——型别越高，能跟踪的输入越复杂",
            "💡 技巧：PID参数整定口诀：先P后I最后D，由小到大慢慢调"
        ]
    },
    "离散控制": {
        "terms": {
            "Z变换": {
                "definition": "离散信号的拉普拉斯变换，用于离散系统分析。",
                "formula": "$Z[f(kT)] = \\sum_{k=0}^{\\infty} f(kT)z^{-k}$",
                "example": "单位阶跃序列Z变换：$\\frac{z}{z-1}$"
            },
            "离散稳定性": {
                "definition": "离散系统稳定的充要条件：极点全部在Z平面单位圆内。",
                "formula": "所有特征根$|z_i| < 1$",
                "example": "若极点在单位圆上，系统临界稳定"
            },
            "离散稳态误差": {
                "definition": "类似连续系统，用Z终值定理计算。",
                "formula": "$\\lim_{k\\to\\infty} e(k) = \\lim_{z\\to 1} (1-z^{-1})E(z)$",
                "example": "要求$(1-z^{-1})E(z)$的极点必须在单位圆内"
            }
        },
        "common_mistakes": [
            "❌ 易错点：离散系统稳定性是看Z平面单位圆内，不是左半平面！",
            "⚠️ 难点：Z变换终值定理应用条件",
            "💡 技巧：离散系统与连续系统类似，只是变换域不同"
        ]
    },
    "非线性系统": {
        "terms": {
            "常见非线性": {
                "definition": "自动控制系统中常见的非线性特性。",
                "formula": "死区、饱和、间隙、摩擦",
                "example": "饱和非线性：输入超过某一范围后输出不再变化"
            },
            "描述函数法": {
                "definition": "用于分析自激振荡的近似方法。",
                "formula": "$N(A) = \\frac{1}{\\pi A} \\int_0^{2\\pi} y(t)\\sin\\omega t d(\\omega t)$",
                "example": "描述函数法适用于低通滤波特性较好的系统"
            },
            "相平面法": {
                "definition": "在相平面（x-ẋ平面）上绘制相轨迹，分析二阶非线性系统。",
                "formula": "相轨迹方程：$\\frac{d\\dot{x}}{dx} = \\frac{f(x,\\dot{x})}{\\dot{x}}$",
                "example": "相平面法可以直观地看到系统的运动轨迹"
            },
            "自激振荡条件": {
                "definition": "非线性系统产生自激振荡的条件。",
                "formula": "$N(A) = -\\frac{1}{G(j\\omega)}$",
                "example": "描述函数曲线与-1/N(A)曲线的交点"
            }
        },
        "common_mistakes": [
            "❌ 易错点：描述函数法是近似方法，有适用条件！",
            "⚠️ 难点：相平面法绘制时要分段处理非线性",
            "💡 技巧：自激振荡分析用描述函数法"
        ]
    },
    "现代控制基础": {
        "terms": {
            "状态空间表达式": {
                "definition": "现代控制理论的数学模型，由状态方程和输出方程组成。",
                "formula": "$\\begin{cases} \\dot{x} = Ax + Bu \\\\ y = Cx + Du \\end{cases}$",
                "example": "A为系统矩阵，B为输入矩阵，C为输出矩阵"
            },
            "能控性": {
                "definition": "系统能否通过输入在有限时间内将状态从任意初态转移到任意终态。",
                "formula": "能控性矩阵：$rank[B\\ AB\\ A^2B\\ \\dots\\ A^{n-1}B] = n$",
                "example": "能控性矩阵满秩则系统能控"
            },
            "能观性": {
                "definition": "系统能否通过输出在有限时间内确定任意初始状态。",
                "formula": "能观性矩阵：$rank\\begin{bmatrix} C \\\\ CA \\\\ \\vdots \\\\ CA^{n-1} \\end{bmatrix} = n$",
                "example": "能观性矩阵满秩则系统能观"
            },
            "极点配置": {
                "definition": "通过状态反馈任意配置闭环极点（要求系统能控）。",
                "formula": "状态反馈：$u = -Kx$",
                "example": "闭环系统矩阵：$A - BK$"
            },
            "状态观测器": {
                "definition": "重构不能直接测量的状态。",
                "formula": "$\\dot{\\hat{x}} = A\\hat{x} + Bu + L(y - C\\hat{x})$",
                "example": "观测器极点配置要求系统能观"
            }
        },
        "common_mistakes": [
            "❌ 易错点：极点配置要求系统能控，观测器要求系统能观！",
            "⚠️ 难点：能控性和能观性矩阵构造",
            "💡 技巧：现代控制理论是状态空间方法"
        ]
    },
    "考研必背口诀": {
        "terms": {
            "稳定性口诀": {
                "definition": "闭环极点全在左半平面 ⇒ 稳定",
                "formula": "实轴特征：右边零极点奇数个 ⇒ 根轨迹",
                "example": "离散稳定：极点在单位圆内"
            },
            "稳态误差口诀": {
                "definition": "系统型别要记牢，输入类型要分清",
                "formula": "0型阶跃有误差，I型斜坡有误差，II型抛物线有误差",
                "example": "型别越高，能跟踪的输入越复杂"
            },
            "PID口诀": {
                "definition": "PID参数整定：先P后I最后D",
                "formula": "P：响应快，I：消误差，D：减超调",
                "example": "由小到大慢慢调"
            },
            "伯德图口诀": {
                "definition": "截止频率看带宽，相位裕度看稳定",
                "formula": "积分环节：-20dB/dec，微分环节：+20dB/dec",
                "example": "最小相位系统：相位裕度>0°稳定"
            },
            "根轨迹口诀": {
                "definition": "起点极点终点零，实轴右侧奇数个",
                "formula": "渐近线角度中心算，分离汇合求导找",
                "example": "180°根轨迹：幅角±180°(2k+1)"
            }
        },
        "common_mistakes": [
            "💡 技巧：这些口诀是考研必背，一定要记牢！",
            "💡 技巧：做题时多应用这些口诀可以快速解题"
        ]
    }
}


def render_term_popup(term_name, term_info, key_suffix=""):
    """
    渲染单个术语的弹出式解释（点击展开/收起）
    
    Args:
        term_name: 术语名称
        term_info: 术语信息字典
        key_suffix: 用于避免key冲突的后缀
    """
    # 使用session_state来管理展开状态
    state_key = f"term_{term_name}_{key_suffix}"
    if state_key not in st.session_state:
        st.session_state[state_key] = False
    
    # 可点击的术语按钮
    col1, col2 = st.columns([1, 20])
    with col1:
        if st.button("📖", key=f"btn_{term_name}_{key_suffix}"):
            st.session_state[state_key] = not st.session_state[state_key]
    
    with col2:
        st.markdown(f"**{term_name}**")
    
    # 如果展开了，显示详细内容
    if st.session_state[state_key]:
        st.markdown("---")
        st.markdown(f"### 📖 {term_name}")
        st.markdown("**📝 定义：**")
        st.markdown(term_info["definition"])
        st.markdown("**📐 公式：**")
        st.latex(term_info["formula"])
        st.info(f"**💡 例子：** {term_info['example']}")
        st.markdown("---")


def render_mistake_popup(mistake_text, key_suffix=""):
    """
    渲染单个易错点的弹出式提醒
    
    Args:
        mistake_text: 易错点文本
        key_suffix: 用于避免key冲突的后缀
    """
    # 使用session_state来管理展开状态
    state_key = f"mistake_{hash(mistake_text)}_{key_suffix}"
    if state_key not in st.session_state:
        st.session_state[state_key] = False
    
    # 确定图标
    if mistake_text.startswith("❌"):
        icon = "❌"
    elif mistake_text.startswith("⚠️"):
        icon = "⚠️"
    elif mistake_text.startswith("💡"):
        icon = "💡"
    else:
        icon = "📝"
    
    # 可点击的按钮
    if st.button(f"{icon} 查看提醒", key=f"btn_mistake_{hash(mistake_text)}_{key_suffix}"):
        st.session_state[state_key] = not st.session_state[state_key]
    
    # 如果展开了，显示内容
    if st.session_state[state_key]:
        if mistake_text.startswith("❌"):
            st.error(mistake_text)
        elif mistake_text.startswith("⚠️"):
            st.warning(mistake_text)
        elif mistake_text.startswith("💡"):
            st.info(mistake_text)
        else:
            st.markdown(mistake_text)


def render_learning_assistant_page():
    """
    渲染完整的学习辅助页面（独立页面）
    """
    st.markdown('<div class="main-header">', unsafe_allow_html=True)
    st.title("📚 控制原理知识点库")
    st.markdown("### 专业术语解释 | 易错点难点提醒")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # 模块选择器
    module_list = list(TERMINOLOGY.keys())
    selected_module = st.selectbox("选择知识点模块：", module_list)
    
    if selected_module in TERMINOLOGY:
        module_data = TERMINOLOGY[selected_module]
        
        st.markdown(f"---")
        st.subheader(f"📖 {selected_module} - 专业术语解释")
        
        # 显示所有术语
        terms = module_data["terms"]
        for term_name, term_info in terms.items():
            with st.expander(f"📖 {term_name}", expanded=False):
                st.markdown(f"### {term_name}")
                st.markdown("**📝 定义：**")
                st.markdown(term_info["definition"])
                st.markdown("**📐 公式：**")
                st.latex(term_info["formula"])
                st.info(f"**💡 例子：** {term_info['example']}")
        
        st.markdown(f"---")
        st.subheader(f"⚠️ {selected_module} - 易错点与难点提醒")
        
        # 显示所有易错点
        mistakes = module_data["common_mistakes"]
        for mistake in mistakes:
            if mistake.startswith("❌"):
                st.error(mistake)
            elif mistake.startswith("⚠️"):
                st.warning(mistake)
            elif mistake.startswith("💡"):
                st.info(mistake)
            else:
                st.markdown(mistake)


def add_learning_helper_sidebar():
    """
    在侧边栏添加学习辅助导航按钮
    """
    st.sidebar.markdown("---")
    if st.sidebar.button("📚 学习辅助", key="learning_helper_btn"):
        st.session_state.page = "learning"


def add_learning_helper(module_name):
    """
    兼容性函数（暂时保留，避免出错
    """
    pass


# ---------------------- 大厂级配置：全局常量 + 缓存配置 ----------------------
# AI调用配置（大厂规范：常量大写，集中管理）
AI_CONFIG = {
    "TEMPERATURE": 0.5,  # 降低温度，提高响应速度和稳定性
    "MAX_TOKENS_FULL": 3000,  # 减少token数提高速度
    "MAX_TOKENS_SPECIAL": 1500,  # 减少token数提高速度
    "STREAM_TIMEOUT": 45,  # 缩短超时时间
    "CACHE_EXPIRE_SECONDS": 7200  # 缓存有效期延长至2小时
}


# 缓存：解决出卷慢的核心（大厂级缓存策略）
import pickle
import os

_paper_cache = {}  # 内存缓存字典
CACHE_DIR = os.path.join(os.getcwd(), ".cache")

# 确保缓存目录存在
if not os.path.exists(CACHE_DIR):
    os.makedirs(CACHE_DIR)

def get_cached_paper(cache_key):
    """获取缓存的试卷（返回：(questions, answers, cache_time)）"""
    global _paper_cache
    
    # 先检查内存缓存
    if cache_key in _paper_cache:
        questions, answers, cache_time = _paper_cache[cache_key]
        if time.time() - cache_time < AI_CONFIG["CACHE_EXPIRE_SECONDS"]:
            return questions, answers, cache_time
        else:
            del _paper_cache[cache_key]
    
    # 再检查磁盘缓存
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                cached_data = pickle.load(f)
            questions, answers, cache_time = cached_data
            if time.time() - cache_time < AI_CONFIG["CACHE_EXPIRE_SECONDS"]:
                # 更新到内存缓存
                _paper_cache[cache_key] = cached_data
                return cached_data
            else:
                # 删除过期缓存
                os.remove(cache_file)
        except:
            # 缓存文件损坏，删除
            if os.path.exists(cache_file):
                os.remove(cache_file)
    return None


def set_cached_paper(cache_key, questions, answers):
    """设置试卷缓存"""
    global _paper_cache
    cached_data = (questions, answers, time.time())
    
    # 更新内存缓存
    _paper_cache[cache_key] = cached_data
    if len(_paper_cache) > 100:
        oldest_key = min(_paper_cache.keys(), key=lambda k: _paper_cache[k][2])
        del _paper_cache[oldest_key]
    
    # 更新磁盘缓存
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.pkl")
    try:
        with open(cache_file, 'wb') as f:
            pickle.dump(cached_data, f)
    except:
        # 磁盘缓存失败，忽略
        pass


# ---------------------- 知识图谱功能（基于胡寿松自动控制原理第五版） ----------------------
KNOWLEDGE_GRAPH = {
    "id": "root",
    "name": "自动控制原理",
    "description": "胡寿松《自动控制原理》第五版核心知识体系",
    "children": [
        {
            "id": "chapter1",
            "name": "自动控制的一般概念",
            "description": "控制系统基本概念、组成、分类与性能要求",
            "children": [
                {
                    "id": "concept1-1",
                    "name": "控制系统组成",
                    "description": "控制对象、控制器、检测装置、执行机构"
                },
                {
                    "id": "concept1-2",
                    "name": "反馈控制原理",
                    "description": "利用偏差消除偏差的闭环控制原理"
                },
                {
                    "id": "concept1-3",
                    "name": "控制系统分类",
                    "description": "开环/闭环、线性/非线性、连续/离散、定常/时变"
                }
            ]
        },
        {
            "id": "chapter2",
            "name": "控制系统的数学模型",
            "description": "微分方程、传递函数、动态结构图、信号流图",
            "children": [
                {
                    "id": "concept2-1",
                    "name": "微分方程建立",
                    "description": "机理法建模，线性化处理"
                },
                {
                    "id": "concept2-2",
                    "name": "传递函数",
                    "description": "零初始条件下输出与输入拉普拉斯变换之比"
                },
                {
                    "id": "concept2-3",
                    "name": "动态结构图",
                    "description": "环节连接、等效变换、梅森公式"
                }
            ]
        },
        {
            "id": "chapter3",
            "name": "线性系统的时域分析法",
            "description": "一阶/二阶系统响应、稳定性分析、稳态误差",
            "children": [
                {
                    "id": "concept3-1",
                    "name": "二阶系统性能",
                    "description": "阻尼比ζ、自然频率ωₙ、超调量σ%、调节时间tₛ"
                },
                {
                    "id": "concept3-2",
                    "name": "劳斯判据",
                    "description": "代数判据判断系统稳定性"
                },
                {
                    "id": "concept3-3",
                    "name": "稳态误差计算",
                    "description": "终值定理、误差系数、系统型别"
                }
            ]
        },
        {
            "id": "chapter4",
            "name": "根轨迹法",
            "description": "根轨迹绘制、根轨迹分析与校正",
            "children": [
                {
                    "id": "concept4-1",
                    "name": "根轨迹绘制法则",
                    "description": "180°根轨迹、0°根轨迹、参数根轨迹"
                },
                {
                    "id": "concept4-2",
                    "name": "根轨迹分析",
                    "description": "主导极点、系统性能估算"
                }
            ]
        },
        {
            "id": "chapter5",
            "name": "线性系统的频域分析法",
            "description": "频率特性、伯德图、奈奎斯特判据、稳定裕度",
            "children": [
                {
                    "id": "concept5-1",
                    "name": "频率特性",
                    "description": "幅频特性、相频特性、极坐标图"
                },
                {
                    "id": "concept5-2",
                    "name": "伯德图",
                    "description": "对数幅频特性、对数相频特性、渐近线"
                },
                {
                    "id": "concept5-3",
                    "name": "奈奎斯特判据",
                    "description": "奈奎斯特曲线、包围次数、稳定判据"
                },
                {
                    "id": "concept5-4",
                    "name": "稳定裕度",
                    "description": "相角裕度γ、幅值裕度h"
                }
            ]
        },
        {
            "id": "chapter6",
            "name": "线性系统的校正方法",
            "description": "串联校正、PID校正、前馈校正、复合校正",
            "children": [
                {
                    "id": "concept6-1",
                    "name": "串联超前校正",
                    "description": "相位超前、提高相角裕度、改善动态性能"
                },
                {
                    "id": "concept6-2",
                    "name": "串联滞后校正",
                    "description": "相位滞后、提高稳态精度、保持动态性能"
                },
                {
                    "id": "concept6-3",
                    "name": "PID校正",
                    "description": "比例P、积分I、微分D、组合校正"
                }
            ]
        }
    ]
}


def render_knowledge_graph_node(node, level=0):
    """递归渲染知识图谱节点"""
    indent = "  " * level
    if "children" in node and node["children"]:
        expander = st.expander(f"{indent}📚 {node['name']}", expanded=(level < 2))
        with expander:
            st.markdown(f"**{node['description']}**")
            st.divider()
            for child in node["children"]:
                render_knowledge_graph_node(child, level + 1)
    else:
        st.markdown(f"{indent}📖 **{node['name']}**")
        st.caption(node["description"])


def get_knowledge_graph_html():
    """生成交互式知识图谱的HTML（使用D3.js）"""
    return """
    <style>
        .graph-container {
            width: 100%;
            height: 600px;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            overflow: hidden;
        }
        .node circle {
            fill: #667eea;
            stroke: #fff;
            stroke-width: 2px;
            cursor: pointer;
        }
        .node text {
            font-family: 'SimSun', 'Microsoft YaHei', sans-serif;
            font-size: 12px;
            pointer-events: none;
        }
        .link {
            stroke: #999;
            stroke-opacity: 0.6;
        }
        .tooltip {
            position: absolute;
            padding: 8px 12px;
            background: rgba(0,0,0,0.8);
            color: white;
            border-radius: 4px;
            font-size: 12px;
            pointer-events: none;
            z-index: 1000;
        }
    </style>
    <div class="graph-container" id="knowledge-graph"></div>
    <div class="tooltip" id="graph-tooltip" style="display: none;"></div>
    <script src="https://d3js.org/d3.v7.min.js"></script>
    <script>
        const data = {
            "name": "自动控制原理",
            "children": [
                {"name": "自动控制的一般概念", "description": "控制系统基本概念、组成、分类与性能要求"},
                {"name": "控制系统的数学模型", "description": "微分方程、传递函数、动态结构图"},
                {"name": "线性系统的时域分析法", "description": "一阶/二阶系统响应、稳定性分析、稳态误差"},
                {"name": "根轨迹法", "description": "根轨迹绘制、根轨迹分析与校正"},
                {"name": "线性系统的频域分析法", "description": "频率特性、伯德图、奈奎斯特判据"},
                {"name": "线性系统的校正方法", "description": "串联校正、PID校正、复合校正"}
            ]
        };
        
        const width = document.getElementById('knowledge-graph').clientWidth;
        const height = 600;
        
        const svg = d3.select("#knowledge-graph")
            .append("svg")
            .attr("width", width)
            .attr("height", height);
        
        const g = svg.append("g");
        
        const zoom = d3.zoom()
            .scaleExtent([0.3, 3])
            .on("zoom", (event) => {
                g.attr("transform", event.transform);
            });
        
        svg.call(zoom);
        
        const simulation = d3.forceSimulation()
            .force("link", d3.forceLink().id(d => d.name).distance(100))
            .force("charge", d3.forceManyBody().strength(-400))
            .force("center", d3.forceCenter(width / 2, height / 2))
            .force("collision", d3.forceCollide().radius(40));
        
        const nodes = [data].concat(data.children);
        const links = data.children.map(child => ({source: data.name, target: child.name}));
        
        const link = g.append("g")
            .selectAll("line")
            .data(links)
            .enter().append("line")
            .attr("class", "link");
        
        const node = g.append("g")
            .selectAll("g")
            .data(nodes)
            .enter().append("g")
            .attr("class", "node")
            .call(d3.drag()
                .on("start", dragstarted)
                .on("drag", dragged)
                .on("end", dragended));
        
        node.append("circle")
            .attr("r", d => d.name === "自动控制原理" ? 30 : 20)
            .style("fill", d => d.name === "自动控制原理" ? "#e74c3c" : "#3498db");
        
        node.append("text")
            .attr("dy", d => d.name === "自动控制原理" ? 45 : 35)
            .style("text-anchor", "middle")
            .text(d => d.name);
        
        const tooltip = d3.select("#graph-tooltip");
        
        node.on("mouseover", (event, d) => {
            if (d.description) {
                tooltip.style("display", "block")
                    .html(`<strong>${d.name}</strong><br>${d.description}`)
                    .style("left", (event.pageX + 10) + "px")
                    .style("top", (event.pageY - 10) + "px");
            }
        })
        .on("mouseout", () => {
            tooltip.style("display", "none");
        });
        
        simulation.nodes(nodes).on("tick", ticked);
        simulation.force("link").links(links);
        
        function ticked() {
            link
                .attr("x1", d => d.source.x)
                .attr("y1", d => d.source.y)
                .attr("x2", d => d.target.x)
                .attr("y2", d => d.target.y);
            
            node.attr("transform", d => `translate(${d.x},${d.y})`);
        }
        
        function dragstarted(event, d) {
            if (!event.active) simulation.alphaTarget(0.3).restart();
            d.fx = d.x;
            d.fy = d.y;
        }
        
        function dragged(event, d) {
            d.fx = event.x;
            d.fy = event.y;
        }
        
        function dragended(event, d) {
            if (!event.active) simulation.alphaTarget(0);
            d.fx = null;
            d.fy = null;
        }
    </script>
    """


# ---------------------- 火山方舟 SDK 初始化（大厂级单例模式） ----------------------
# 全局客户端，避免重复初始化（单例模式）
_ARK_CLIENT = None
_MODEL_ID = None
_INIT_LOCK = False  # 初始化锁，避免并发初始化


def init_ark_client_once():
    """全局只初始化一次火山方舟客户端（大厂级：加锁防并发，容错增强）"""
    global _ARK_CLIENT, _MODEL_ID, _INIT_LOCK

    # 加锁防并发初始化（大厂必备：避免多线程重复创建客户端）
    if _INIT_LOCK:
        return _ARK_CLIENT, _MODEL_ID
    _INIT_LOCK = True

    try:
        # 先返回已初始化的客户端（避免重复初始化）
        if _ARK_CLIENT and _MODEL_ID:
            _INIT_LOCK = False
            return _ARK_CLIENT, _MODEL_ID

        # 严格读取 secrets.toml 中的键名：VOLC_API_KEY、VOLC_MODEL_ID
        api_key = st.secrets.get("volc_ark", {}).get("VOLC_API_KEY")
        model_id = st.secrets.get("volc_ark", {}).get("VOLC_MODEL_ID")

        # 容错：大厂级参数校验
        if not api_key or not model_id:
            st.error("❌ 密钥配置错误：secrets.toml 中缺少 VOLC_API_KEY 或 VOLC_MODEL_ID")
            _INIT_LOCK = False
            return None, None

        # 官网新版SDK初始化（固定配置，无需修改）
        _ARK_CLIENT = Ark(
            base_url="https://ark.cn-beijing.volces.com/api/v3",
            api_key=api_key
        )
        _MODEL_ID = model_id

        _INIT_LOCK = False
        return _ARK_CLIENT, _MODEL_ID

    except Exception as e:
        st.error(f"❌ 客户端初始化失败：{str(e)}")
        _INIT_LOCK = False
        return None, None


# ---------------------- 基础工具函数（大厂级鲁棒性增强） ----------------------
def safe_parse_coeff(input_str, default):
    """安全解析系数，容错所有输入错误（大厂级：增强异常捕获）"""
    if not isinstance(input_str, str) or input_str.strip() == "":
        return default
    try:
        # 处理各种分隔符（逗号、空格、分号）
        input_str = input_str.replace("，", ",").replace(";", ",").replace(" ", "")
        coeffs = [float(x) for x in input_str.split(",") if x.strip()]
        # 确保至少有一个非零系数
        if all(abs(c) < 1e-6 for c in coeffs):
            return default
        return coeffs
    except (ValueError, TypeError, IndexError):
        # 大厂级：捕获所有可能的异常，避免程序崩溃
        st.warning(f"⚠️ 系数解析失败，使用默认值：{default}")
        return default


def poly_to_latex(coeffs):
    """将系数转换为美观的LaTeX多项式字符串（从高次到低次）
    修复点：
    1. 解决系数为1/-1时的显示问题
    2. 修复LaTeX语法错误（如s^1显示为s）
    3. 增强空值/零值容错
    """
    if not isinstance(coeffs, (list, np.ndarray)) or len(coeffs) == 0:
        return "0"

    terms = []
    # 从最高次到最低次遍历（coeffs[0]是最高次系数，coeffs[-1]是常数项）
    for power_idx, coeff in enumerate(coeffs):
        if abs(coeff) < 1e-6:
            continue

        # 处理系数显示（大厂级：精细化格式）
        coeff_abs = abs(coeff)
        if coeff_abs == 1 and power_idx != len(coeffs) - 1:  # 非常数项系数为1/-1
            coeff_str = ""
        else:
            coeff_str = f"{coeff_abs:.2f}".rstrip('0').rstrip('.') if coeff_abs != int(
                coeff_abs) else f"{int(coeff_abs)}"

        # 处理幂次显示（power是从0开始的索引，实际幂次是 len(coeffs)-1 - power）
        actual_power = len(coeffs) - 1 - power_idx
        if actual_power == 0:
            term = coeff_str if coeff_str else "1"
        elif actual_power == 1:
            term = f"{coeff_str}s" if coeff_str else "s"
        else:
            term = f"{coeff_str}s^{actual_power}" if coeff_str else f"s^{actual_power}"

        # 处理符号（大厂级：避免首项前的+号）
        if terms:
            if coeff > 0:
                terms.append(" + ")
            else:
                terms.append(" - ")
        elif coeff < 0:
            terms.append("-")

        terms.append(term)

    # 处理空结果
    latex_str = "".join(terms) if terms else "0"
    # 最终容错：确保返回合法的LaTeX字符串
    return latex_str.replace("  ", " ").strip()


# ---------------------- 系统性能计算函数（鲁棒性增强） ----------------------
# 二阶系统分析函数已移至下方，使用num, den, K参数版本





def calculate_pid_performance(G_closed):
    """计算PID校正后系统的性能指标（鲁棒性增强，修复NamedSignal类型问题）"""
    try:
        if not isinstance(G_closed, ctrl.TransferFunction):
            return None

        # 获取阶跃响应数据（大厂级：控制时间范围，避免计算量过大）
        t = np.linspace(0, 10, 1000)
        _, y = step(G_closed, T=t)

        # 核心修复：把NamedSignal转成numpy数组
        y = np.array(y).flatten()

        # 容错：空响应数据
        if len(y) == 0:
            return None

        # 计算超调量（容错：稳态值为0/无穷）
        y_ss = y[-1] if np.isfinite(y[-1]) and y[-1] != 0 else 1.0
        y_max = np.max(y)
        overshoot = ((y_max - y_ss) / y_ss) * 100 if y_ss > 0 else 0

        # 计算调节时间（2%误差带，大厂级：反向遍历更高效）
        settling_time = np.inf
        threshold = 0.02 * y_ss
        for i in range(len(y) - 1, -1, -1):
            if abs(y[i] - y_ss) > threshold:
                settling_time = t[i] if i < len(t) else np.inf
                break

        # 计算上升时间
        rise_time = np.inf
        for i in range(len(y)):
            if y[i] >= y_ss:
                rise_time = t[i]
                break

        return {
            "overshoot": round(overshoot, 2),
            "settling_time": round(settling_time, 2) if np.isfinite(settling_time) else np.inf,
            "rise_time": round(rise_time, 2) if np.isfinite(rise_time) else np.inf,
            "steady_state_value": round(y_ss, 2)
        }
    except Exception as e:
        st.warning(f"⚠️ PID性能指标计算失败：{str(e)[:50]}")
        return None

# ---------------------- 题海实战工具函数（核心优化：缓存+LaTeX强制格式） ----------------------
def generate_sample_questions(q_type, q_num=1):
    """生成基础练习题（改为调用AI生成，移除本地固定题库）"""
    # 直接调用AI生成对应题型的题目，不再使用本地题库
    questions, answers = generate_ai_exam_paper(
        mode="专项练习",
        scope=q_type,
        num=q_num
    )
    # 容错处理：如果AI生成失败，补充默认题目（仅兜底）
    while len(questions) < q_num:
        questions.append(f"{len(questions) + 1}. 综合分析题：{q_type}考研核心考点综合应用。")
        answers.append("答题要点：1. 系统建模；2. 稳定性分析；3. 性能指标计算；4. 校正方案设计。")
    return questions, answers


def safe_escape_latex(text):
    """
    安全转义LaTeX中的特殊字符（大厂级：精准转义，不破坏公式）
    核心修复：只转义Python格式化冲突的%，保留LaTeX语法的{}
    """
    if not isinstance(text, str):
        return ""
    # 仅转义%（避免Python格式化），不转义LaTeX的{}
    text = text.replace("%", "\\%")
    # 修复AI可能生成的错误LaTeX格式（如{{K}} → {K}）
    text = text.replace("{{", "{").replace("}}", "}")
    return text


def safe_latex_render(latex_expr):
    """
    安全的LaTeX渲染函数
    - 直接使用Streamlit的latex渲染
    """
    if not isinstance(latex_expr, str):
        return
    st.latex(latex_expr)


def ensure_latex_format(text):
    """清理代码块和复杂格式，只保留纯文本和LaTeX公式"""
    import re
    if not text:
        return text
    
    # 第一步：清理所有代码块（```...```）
    text = re.sub(r'```[\s\S]*?```', '', text)
    
    # 第二步：清理行内代码（`...`）
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # 第三步：清理复杂的LaTeX命令
    # 清理 \text{...}
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    # 清理 \mathrm{...}
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    # 清理 \mathbf{...}
    text = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', text)
    # 清理 \mathit{...}
    text = re.sub(r'\\mathit\{([^}]*)\}', r'\1', text)
    # 清理 \mbox{...}
    text = re.sub(r'\\mbox\{([^}]*)\}', r'\1', text)
    # 清理 \code{...} 或 \texttt{...}
    text = re.sub(r'\\code\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    
    # 第四步：清理Markdown标题符号，但保留内容
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    
    # 第五步：清理多余的星号（**粗体**）和下划线（_斜体_）
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # 第六步：修复格式问题
    if "$" in text:
        text = text.replace("$$", "$")  # 双$转单$
        text = text.replace("{{", "{").replace("}}", "}")  # 修复多余的括号
    
    # 第七步：清理多余的空行
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    
    return text


def parse_specific_exercises(content: str) -> tuple[list, list]:
    """解析专项练习模式下AI返回的题目（只解析题目，无参考答案）"""
    questions = []
    answers = []
    if not content:
        return questions, answers

    lines = [l.strip() for l in content.split("\n") if l.strip()]

    current_q = ""

    for line in lines:
        # 题目识别：只要包含数字+点号的行都视为新题目开始
        if any(f"{i}." in line for i in range(1, 100)):
            # 先处理上一题
            if current_q:
                questions.append(safe_escape_latex(current_q))
                answers.append("")  # 答案为空
            # 初始化新题目
            current_q = line
        elif current_q:
            # 当前题目继续添加内容（除非遇到参考答案标记）
            if "参考答案" not in line:
                current_q += "\n" + line

    # 处理最后一题
    if current_q:
        questions.append(safe_escape_latex(current_q))
        answers.append("")

    # 容错处理：如果没有解析到题目，尝试使用整个内容作为题目
    if not questions:
        questions.append(safe_escape_latex(content[:1000]))  # 限制长度
        answers.append("")

    # 批量校验题目的LaTeX格式
    questions = [ensure_latex_format(q) for q in questions]
    answers = [ensure_latex_format(a) for a in answers]

    return questions, answers


# ---------------------- AI智能出卷核心函数（大厂级优化：缓存+异步+LaTeX强制） ----------------------
def generate_ai_exam_paper(mode: str = "专项练习", scope: str = "根轨迹分析", num: int = None) -> tuple[list, list]:
    """
    AI智能出卷（核心功能）
    大厂级优化：
    1. 缓存：相同参数直接返回缓存结果（秒出）
    2. 强制LaTeX：prompt严格要求AI输出LaTeX格式
    3. 超时控制：避免无限等待
    4. 批量生成：一次调用生成所有题目，而非循环调用
    """
    # 核心修复：初始化K，避免未定义引用
    K = 0.0

    # 步骤1：生成缓存key（大厂级：唯一标识一套试卷）
    cache_key = hashlib.md5(f"{mode}_{scope}_{num}".encode()).hexdigest()

    # 步骤2：先查缓存（核心解决出卷慢）
    cached_data = get_cached_paper(cache_key)
    if cached_data:
        questions, answers, cache_time = cached_data
        st.success("✅ 从缓存加载试卷（秒出）")
        return questions, answers

    # 步骤3：初始化AI客户端
    client, model_id = init_ark_client_once()
    if not client or not model_id:
        return ["❌ AI功能不可用：请检查火山方舟密钥配置"], [""]

    # 统一定义LaTeX示例（大厂级：标准化prompt，避免AI理解偏差）
    latex_examples = r"""
    【绝对重要】所有数学公式、变量、符号、字母必须用简单的LaTeX格式$...$包裹，无任何例外！
    【禁止使用】禁止使用任何复杂LaTeX命令：\text{}、\mathrm{}、\mathbf{}、\mathit{}等
    示例：
    - 传递函数：$G(s) = \frac{K}{s(Ts+1)}$（正确），禁止写成G(s)=K/(s(Ts+1))（错误）
    - 阻尼比：$\zeta = 0.707$（正确），禁止写成ζ=0.707（错误）
    - 自然频率：$\omega_n = 10$ rad/s（正确），禁止写成ωₙ=10 rad/s（错误）
    - 超调量：$\sigma\% = e^{-\zeta\pi/\sqrt{1-\zeta^2}} \times 100\%$（正确）
    - 变量：$K$、$s$、$t$、$G$、$C$、$T$、$n$、$m$、$p$、$z$（正确），禁止直接写K、s、t（错误）
    - 下标：$K_p$、$K_i$、$K_d$、$t_s$、$t_p$、$\omega_c$、$\omega_d$（正确）
    - 百分比：$50\%$（正确），禁止写50%（错误）
    - 度数：$45^\circ$（正确），禁止写45°（错误）
    - 禁止：$K_{\text{min}}$（错误），正确写法：$K_{min}$
    - 禁止：$e_{\text{ss}}$（错误），正确写法：$e_{ss}$
    """

    # 构建精准的出题提示词（大厂级：指令调优，强制格式）
    system_prompt = r"""你是控制原理教育专家，精通自动控制原理核心知识点。
    【绝对必须遵守的输出规则】：
    1. 只有完整的数学公式（如方程、等式、分式、求和、积分等）才用LaTeX格式$...$包裹！
    2. 【重要】单独的字母、数字、符号、百分比、度数（如 K、s、t、50%、45° 等）都直接用纯文本，不要用LaTeX包裹！
    3. 【禁止使用】禁止使用任何复杂LaTeX命令：\text{}、\mathrm{}、\mathbf{}、\mathit{}、\mbox{}等
    4. 下标直接用下划线：K_min、e_ss，不要用$K_{\text{min}}$
    5. 单个字母变量用纯文本：K、s、t、G、C、T、n、m、p、z
    6. 下标符号用纯文本：Kp、Ki、Kd、ts、tp、ωc、ωd
    7. 百分比和度数用纯文本：50%、45°
    8. 题目原创，参数随机，考察角度不同，符合学习和考试要求的难度。
    9. 只输出题目，不要输出【参考答案】！只需要题目内容即可！
    10. 禁止使用{{}}包裹变量（如{{K}}），直接写K。
    11. 输出格式严格按照示例，无多余内容。
    12. 【绝对禁止】禁止使用任何代码块格式！包括```...```和`...`，所有内容都用纯文本和LaTeX公式表示。
    13. 【绝对禁止】不要用任何形式的代码来表示数学推导，全部用自然语言表示。
    14. 【重要】输出必须让不懂编程的学生能看懂！"""

    if mode == "整卷模拟测试":
        # 整卷模式：明确要求AI生成完整的控制原理学习测试卷，不包含参考答案
        user_prompt = f"""
        生成一套完整的控制原理学习测试卷，包含：
        - 选择题（5道）、填空题（5道）、计算题（4道）、分析题（2道）
        - 只有完整的数学公式用LaTeX格式的$...$包裹
        - 单独的字母、数字、符号用纯文本表示
        - 难度适中，包含基础题、中等题、拔高题
        - 只输出题目，不要输出参考答案！

        输出格式：
        ### 控制原理学习测试卷
        一、选择题（每题5分，共25分）
        1. 题目内容
        2. 题目内容
        ...
        二、填空题（每题5分，共25分）
        1. 题目内容
        2. 题目内容
        ...
        三、计算题（每题15分，共60分）
        1. 题目内容
        2. 题目内容
        ...
        四、分析题（每题20分，共40分）
        1. 题目内容
        2. 题目内容
        """
        max_tokens = AI_CONFIG["MAX_TOKENS_FULL"]
    else:
        # 专项练习模式，不包含参考答案
        user_prompt = "生成" + str(num) + "道" + scope + "的专项练习题，要求：\n"
        user_prompt += "- 每道题考察不同角度，原创不重复\n"
        user_prompt += "- 只有完整的数学公式用LaTeX格式的$...$包裹\n"
        user_prompt += "- 单独的字母、数字、符号用纯文本表示\n"
        user_prompt += "- 难度中等偏上，符合控制原理学习要求\n"
        user_prompt += "- 只输出题目，不要输出参考答案！\n\n"
        user_prompt += "输出格式示例：\n"
        user_prompt += "1. 已知开环传递函数$G(s) = K/(s(s + 1)(s + 2))$，求根轨迹的渐近线角度与中心，并判断闭环稳定的K范围。\n"
        user_prompt += "2. 已知二阶系统$G(s) = 1/(s^2 + 2s + 1)$，超调量过大，给出PID参数调整建议并说明理由。"
        max_tokens = AI_CONFIG["MAX_TOKENS_SPECIAL"]

    try:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        # 调用AI生成试卷（大厂级：流式处理+进度提示）
        start_time = time.time()
        
        # 创建一个占位符来显示生成进度
        status_placeholder = st.empty()
        full_content = ""
        
        # 启用流式处理
        response = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=AI_CONFIG["TEMPERATURE"],  # 使用配置的温度参数
            max_tokens=max_tokens,
            stream=True,
            timeout=AI_CONFIG["STREAM_TIMEOUT"]  # 超时控制
        )
        
        # 实时处理流式响应
        for chunk in response:
            if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                content_chunk = chunk.choices[0].delta.content
                full_content += content_chunk
                # 只显示进度提示，不显示题目内容
                status_placeholder.info("AI正在生成题目...")
        
        end_time = time.time()
        status_placeholder.success(f"✅ AI出卷耗时：{end_time - start_time:.2f}秒")
        
        if not full_content:
            return ["❌ 试卷生成失败：AI无有效响应"], [""]

        # 使用完整生成的内容
        content = full_content

        # 整卷模式/专项模式解析
        if mode == "华东理工大学816考研整卷（AI全真模拟）":
            questions, answers = parse_full_exam_paper(content)
        else:
            questions, answers = parse_specific_exercises(content)

        # 步骤4：缓存结果（解决下次出卷慢）
        set_cached_paper(cache_key, questions, answers)

        return questions, answers

    except TimeoutError:
        return [f"❌ 试卷生成超时：AI响应超过{AI_CONFIG['STREAM_TIMEOUT']}秒，请重试"], [""]
    except Exception as e:
        error_msg = str(e)
        if "429" in error_msg or "SetLimitExceeded" in error_msg:
            return ["❌ 试卷生成出错：API调用频率超过限制（429错误）\n\n解决方法：\n1. 稍后再试（建议等待1-2分钟）\n2. 减少题目数量\n3. 检查火山方舟API密钥配置"], [""]
        elif "API key" in error_msg or "authentication" in error_msg:
            return ["❌ 试卷生成出错：API密钥配置错误\n\n解决方法：\n1. 检查secrets.toml中的VOLC_API_KEY配置\n2. 确保API密钥有效"], [""]
        else:
            return [f"❌ 试卷生成出错：{error_msg[:100]}"], [""]

# ---------------------- 辅助解析函数（修复转义逻辑+LaTeX校验） ----------------------
def parse_full_exam_paper(content: str) -> tuple[list, list]:
    """解析整卷模式下AI返回的完整试卷（适配AI动态生成的格式+LaTeX校验）"""
    questions = []
    answers = []

    if not content:
        return questions, answers

    # 解析题目部分（兼容AI生成的格式）
    parts = content.split("参考答案及评分标准")
    if len(parts) < 2:
        # 降级解析：直接按行解析
        lines = [l.strip() for l in content.split("\n") if l.strip()]
        for line in lines:
            if line.startswith(tuple(f"{i}." for i in range(1, 100))):
                if "【参考答案】" in line:
                    q_part, a_part = line.split("【参考答案】")
                    questions.append(safe_escape_latex(q_part))
                    answers.append(safe_escape_latex(a_part.replace("：", "")))
                else:
                    questions.append(safe_escape_latex(line))
        return questions, ["参考答案暂未生成"] * len(questions)

    # 解析题目部分
    question_part = parts[0]
    q_lines = [l.strip() for l in question_part.split("\n") if l.strip()]
    for line in q_lines:
        if line.startswith(tuple(f"{i}." for i in range(1, 100))) or \
                (("选择题" in line or "填空题" in line or "计算题" in line) and any(c.isdigit() for c in line)):
            questions.append(safe_escape_latex(line))

    # 解析答案部分
    answer_part = parts[1]
    a_lines = [l.strip() for l in answer_part.split("\n") if l.strip()]
    for line in a_lines:
        if line.startswith(tuple(f"{i}." for i in range(1, 100))):
            answers.append(safe_escape_latex(line))

    # 容错处理：保证题目和答案数量匹配
    max_len = max(len(questions), len(answers))
    while len(questions) < max_len:
        questions.append(f"{len(questions) + 1}. 题目生成失败，请重试")
    while len(answers) < max_len:
        # 不添加固定的"参考答案暂未生成"，而是使用空字符串
        answers.append("")

    # LaTeX格式校验 - 使用统一的函数
    questions = [ensure_latex_format(q) for q in questions]
    answers = [ensure_latex_format(a) for a in answers]

    return questions, answers


# ---------------------- 格式化计时函数（智能单位） ----------------------
def format_elapsed_time(seconds):
    if seconds < 60:
        return f"{seconds:.1f} 秒"
    elif seconds < 3600:
        return f"{seconds / 60:.1f} 分钟"
    elif seconds < 86400:
        return f"{seconds / 3600:.1f} 小时"
    else:
        return f"{seconds / 86400:.1f} 天"


# ---------------------- AI流式输出函数（大厂级：超时+LaTeX强制） ----------------------
def get_volc_ai_answer_stream(user_question: str):
    """AI答疑流式输出（修复DOM操作报错+超时控制+LaTeX强制）"""
    if not user_question.strip():
        yield "请输入有效的控制原理考研问题！"
        return

    client, model_id = init_ark_client_once()
    if not client or not model_id:
        yield "AI功能不可用：客户端初始化失败"
        return

    try:
        # 强制LaTeX格式的system prompt（大厂级：精准指令，改用原始字符串）
        system_prompt = r"""你是控制工程领域的高级教授，拥有丰富的教学和科研经验，精通自动控制原理、控制系统分析与设计等核心课程。
        
        【绝对必须遵守的回答要求】：
        1. 专业定位：你是控制工程领域的权威专家，回答要体现专业性和权威性。
        2. 【重要】只有完整的数学公式（如方程、等式、分式、求和、积分等）才用LaTeX格式$...$包裹！
        3. 【重要】单独的字母、数字、符号、百分比、度数（如 K、s、t、50%、45° 等）都直接用纯文本，不要用LaTeX包裹！
        4. 【禁止使用】禁止使用任何复杂LaTeX命令：\text{}、\mathrm{}、\mathbf{}、\mathit{}、\mbox{}等
        5. 下标直接用下划线：K_min、e_ss，不要用$K_{\text{min}}$
        6. 单个字母变量用纯文本：K、s、t、G、C、T、n、m、p、z
        7. 下标符号用纯文本：Kp、Ki、Kd、ts、tp、ωc、ωd
        8. 百分比和度数用纯文本：50%、45°
        9. 清晰表达：步骤贴合学术规范，语言简洁易懂，优先给出核心结论。
        10. 全面准确：回答要全面、准确，涵盖问题的各个方面。
        11. 禁止使用{{}}包裹变量（如{{K}}），直接写K。
        12. 【绝对禁止】禁止使用任何代码块格式！包括```...```和`...`，所有内容都用纯文本表示。
        13. 【绝对禁止】不要用任何形式的代码来表示数学推导，全部用自然语言表示。
        14. 【重要】输出必须让不懂编程的学生能看懂！
        
        正确示例：
        - 阻尼比：ζ=0.707
        - 自然频率：ωn=10 rad/s
        - 传递函数：$G(s)=\frac{1}{s^2+2s+1}$（完整公式才用$包裹）
        - 根轨迹渐近线：$\sigma = \frac{\sum poles - \sum zeros}{n - m}$（完整公式才用$包裹）
        - 变量：K、s、t、G（纯文本，不用$包裹）
        - 百分比：25%（纯文本，不用$包裹）
        - 度数：60°（纯文本，不用$包裹）
        - 禁止：K_{\text{min}}（错误），正确写法：K_min
        - 禁止：e_{\text{ss}}（错误），正确写法：e_ss"""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_question}
        ]

        # 流式调用（大厂级：超时控制+性能优化）
        response = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=0.7,
            max_tokens=2000,
            stream=True,
            timeout=AI_CONFIG["STREAM_TIMEOUT"]
        )

        # 流式返回（修复：空chunk容错）
        for chunk in response:
            if chunk and chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                content = chunk.choices[0].delta.content
                # 实时修复LaTeX格式（AI漏加$的情况）
                # 检查是否包含LaTeX命令或数学符号
                latex_commands = ["\\frac", "\\sum", "\\int", "\\sqrt", "\\omega", "\\zeta", "\\sigma", "\\omega_n"]
                math_symbols = ["=", "/", "^", "_", "ω", "ζ", "π", "∞", "∑", "∏", "√", "∫"]
                
                # 检查是否需要添加$符号
                has_latex = any(cmd in content for cmd in latex_commands)
                has_math = any(sym in content for sym in math_symbols)
                
                if (has_latex or has_math) and "$" not in content:
                    # 确保内容被正确包裹在$符号中
                    content = f"${content}$"
                yield content
        return

    except TimeoutError:
        yield "AI回答超时：请重试或简化问题！"
    except Exception as e:
        yield f"AI回答失败：{str(e)[:50]}"
        return


# ---------------------- 考研答疑 AI 函数 ----------------------
def get_volc_ai_answer(user_question: str) -> str:
    """兼容原有调用的非流式版本"""
    full_answer = ""
    for chunk in get_volc_ai_answer_stream(user_question):
        full_answer += chunk
    return full_answer


# ---------------------- 考研出题 AI 函数（大厂级：缓存+批量+LaTeX） ----------------------
def generate_volc_exam_questions(exam_type: str = "华东理工816", q_num: int = 5) -> tuple[list, list]:
    """生成高质量考研题目（AI动态生成，去重+举一反三+缓存优化）"""
    # 缓存key
    cache_key = hashlib.md5(f"volc_{exam_type}_{q_num}".encode()).hexdigest()
    # 查缓存
    cached_data = st.session_state.get(f"cache_{cache_key}")
    if cached_data:
        questions, answers, cache_time = cached_data
        if time.time() - cache_time < AI_CONFIG["CACHE_EXPIRE_SECONDS"]:
            st.success("✅ 从缓存加载题目（秒出）")
            return questions, answers

    client, model_id = init_ark_client_once()
    if not client or not model_id:
        return ["AI功能不可用：客户端初始化失败"], [""]

    # 优化提示词，强制LaTeX格式（大厂级：标准化，改用原始字符串+手动替换变量）
    prompt = r"""
    你是控制原理考研命题专家，拥有10年华东理工816考研辅导经验，生成{q_num}道高质量题目，要求：
    1. 题目原创不重复，每道题考察不同核心考点，能引导学生举一反三。
    2. 所有公式、变量必须用LaTeX格式的$...$包裹，禁止使用{{{{}}}}包裹变量。
    3. 难度符合考研真题，参考答案包含详细步骤和答题技巧。
    4. 输出格式：
    1. 题干内容（LaTeX规范）
    【参考答案】：详细解题步骤
    2. 题干内容（LaTeX规范）
    【参考答案】：详细解题步骤
    ...
    """.format(q_num=q_num, exam_type=exam_type)

    try:
        messages = [
            {
                "role": "system",
                "content": r"你是控制工程考研命题教授，擅长出高质量、有区分度的题目，所有公式必须用LaTeX格式的$...$包裹。"
            },
            {
                "role": "user",
                "content": prompt
            }
        ]

        # 批量生成（避免循环调用）
        response = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=0.8,
            max_tokens=AI_CONFIG["MAX_TOKENS_SPECIAL"],
            stream=False,
            timeout=AI_CONFIG["STREAM_TIMEOUT"]
        )

        if not response or not response.choices:
            return ["题目生成失败：无响应"], [""]

        content = response.choices[0].message.content

        # 解析生成的题目和答案
        questions = []
        answers = []
        lines = [l.strip() for l in content.split("\n") if l.strip()]

        current_q = ""
        current_a = ""
        in_answer = False

        for line in lines:
            if line.startswith(tuple(f"{i}." for i in range(1, q_num + 1))):
                if current_q and current_a:
                    questions.append(safe_escape_latex(current_q))
                    answers.append(safe_escape_latex(current_a))
                current_q = line
                in_answer = False
            elif line.startswith("【参考答案】"):
                current_a = line.replace("【参考答案】：", "").replace("【参考答案】", "")
                in_answer = True
            elif in_answer:
                current_a += "\n" + line

        # 添加最后一道题
        if current_q and current_a:
            questions.append(safe_escape_latex(current_q))
            answers.append(safe_escape_latex(current_a))

        # 补充不足的题目
        while len(questions) < q_num:
            questions.append(f"{len(questions) + 1}. 综合分析题：{exam_type}考研核心考点综合应用")
            answers.append("解题要点：1. 系统建模；2. 稳定性分析；3. 性能指标计算；4. 校正方案设计；5. 参数优化。")

        # LaTeX格式校验 - 简化版，避免错误包裹
        def ensure_latex_format(text):
            if not text:
                return text
            # 检查是否已经包含$符号
            if "$" in text:
                return text
            # 不自动包裹，避免错误
            return text

        questions = [ensure_latex_format(q) for q in questions]
        answers = [ensure_latex_format(a) for a in answers]

        # 缓存结果
        set_cached_paper(cache_key, questions, answers)

        return questions[:q_num], answers[:q_num]

    except TimeoutError:
        return ["题目生成超时：请重试"], [""]
    except Exception as e:
        return [f"题目生成出错：{str(e)[:50]}"], [""]


# ---------------------- 导出 PDF/Word 函数（修复缩进错误+鲁棒性） ----------------------
def export_to_pdf(questions, answers, filename="control_exam.pdf"):
    """导出题目和答案为 PDF（大厂级：容错+中文字体+路径校验）"""
    if not EXPORT_AVAILABLE:
        st.error("导出功能不可用：reportlab库未安装")
        return None

    # 容错：空题目/答案
    if not questions or not answers:
        st.error("导出失败：无题目/答案可导出")
        return None

    try:
        # 路径校验（大厂级：避免路径错误）
        filename = os.path.join(os.getcwd(), filename)
        # 注册中文字体（容错：字体不存在）
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        except:
            st.warning("⚠️ 中文字体注册失败，PDF可能显示乱码")

        c = canvas.Canvas(filename, pagesize=A4)
        c.setFont("STSong-Light", 12)
        width, height = A4
        y = height - 50

        c.drawString(50, y, "控制原理学习资料")
        y -= 30

        # 批量写入（容错：题目和答案数量不匹配）
        for i, (q, a) in enumerate(zip(questions, answers), 1):
            # 容错：超长文本换行（简化版）
            if len(q) > 80:
                q_lines = [q[i:i + 80] for i in range(0, len(q), 80)]
                for line in q_lines:
                    c.drawString(50, y, f"{i}. {line}")
                    y -= 15
            else:
                c.drawString(50, y, f"{i}. {q}")
            y -= 20

            if y < 50:
                c.showPage()
                c.setFont("STSong-Light", 12)
                y = height - 50

            # 答案处理
            if len(a) > 80:
                a_lines = [a[i:i + 80] for i in range(0, len(a), 80)]
                c.drawString(70, y, "答案：")
                y -= 15
                for line in a_lines:
                    c.drawString(80, y, line)
                    y -= 15
            else:
                c.drawString(70, y, f"答案：{a}")
            y -= 30

            if y < 50:
                c.showPage()
                c.setFont("STSong-Light", 12)
                y = height - 50

        c.save()
        st.success(f"✅ PDF导出成功：{filename}")
        return filename
    except Exception as e:
        st.error(f"PDF导出失败：{str(e)[:100]}")
        return None


def export_to_word(questions, answers, filename="control_exam.docx", is_full_exam=False):
    """导出题目和答案为高质量 Word（满足格式要求）
    
    Args:
        questions: 题目列表
        answers: 答案列表
        filename: 保存的文件名
        is_full_exam: 是否为整卷模式（整卷模式下答案统一放在最后）
    
    Returns:
        文件路径或None
    """
    if not EXPORT_AVAILABLE:
        st.error("导出功能不可用：python-docx库未安装")
        return None

    if not questions or not answers:
        st.error("导出失败：无题目/答案可导出")
        return None

    try:
        filename = os.path.join(os.getcwd(), filename)
        
        # 检查是否有 Spire.Doc 库
        spire_available = False
        try:
            spire_available = 'SPIRE_AVAILABLE' in globals() and SPIRE_AVAILABLE
        except NameError:
            pass
        
        if spire_available:
            doc = Document()
            section = doc.AddSection()
            section.PageSetup.Margins.All = 72.0
            
            # 添加试卷标题
            title_para = section.AddParagraph()
            title_para.AppendText("控制原理学习资料")
            title_para.Format.HorizontalAlignment = HorizontalAlignment.Center
            
            section.AddParagraph()
            
            # 辅助函数：设置段落格式
            def set_spire_paragraph_format(paragraph):
                # 设置行距为1.5倍（12磅 * 1.5 = 18磅）
                paragraph.Format.LineSpacing = 18
                # 设置段前间距
                paragraph.Format.SpaceBefore = 12
                # 设置段后间距
                paragraph.Format.SpaceAfter = 12
            
            # 辅助函数：处理文本和公式
            def add_text_with_formulas(paragraph, text):
                import re
                # 提取LaTeX公式
                latex_pattern = r'\$(.*?)\$'
                parts = re.split(latex_pattern, text)
                
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        # 普通文本
                        if part:
                            text_range = paragraph.AppendText(part)
                            # 设置中文字体为宋体，英文字体为Times New Roman
                            text_range.CharacterFormat.FontName = "Times New Roman"
                            text_range.CharacterFormat.FontNameAscii = "Times New Roman"
                            text_range.CharacterFormat.FontNameEastAsia = "宋体"
                            # 设置字号为小四（12磅）
                            text_range.CharacterFormat.FontSize = 12
                    else:
                        # LaTeX公式
                        if part:
                            math_obj = OfficeMath(doc)
                            math_obj.FromLatexMathCode(part)
                            paragraph.Items.Add(math_obj)
            
            # 添加题目
            section.AddParagraph().AppendText("一、题目")
            
            for i, (q, a) in enumerate(zip(questions, answers), 1):
                # 添加题目（带大纲级别）
                question_heading = section.AddParagraph()
                question_heading.AppendText(f"题目 {i}")
                
                # 添加题目内容
                q_paragraph = section.AddParagraph()
                set_spire_paragraph_format(q_paragraph)
                add_text_with_formulas(q_paragraph, q)
                
                # 如果不是整卷模式，答案紧跟题目
                if not is_full_exam:
                    answer_heading = section.AddParagraph()
                    answer_heading.AppendText(f"题目 {i} 参考答案")
                    
                    a_paragraph = section.AddParagraph()
                    set_spire_paragraph_format(a_paragraph)
                    add_text_with_formulas(a_paragraph, a)
                
                section.AddParagraph()
            
            # 如果是整卷模式，答案统一放在最后
            if is_full_exam:
                section.AddParagraph().AppendBreak(BreakType.PageBreak)
                section.AddParagraph().AppendText("二、参考答案")
                
                for i, a in enumerate(answers, 1):
                    answer_heading = section.AddParagraph()
                    answer_heading.AppendText(f"题目 {i} 参考答案")
                    
                    a_paragraph = section.AddParagraph()
                    set_spire_paragraph_format(a_paragraph)
                    add_text_with_formulas(a_paragraph, a)
                    section.AddParagraph()
            
            doc.SaveToFile(filename, FileFormat.Docx2016)
            doc.Dispose()
        else:
            # 回退到使用 python-docx 库
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = docx.Document()
            
            # 设置页面边距
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            
            # 添加试卷标题
            title = doc.add_heading("控制原理学习资料", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 设置标题格式
            for run in title.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(22)
                run.font.bold = True
            
            doc.add_paragraph()
            
            # 辅助函数：设置段落格式（中文小四、英文Times New Roman、1.5倍行距、首行缩进）
            def set_paragraph_format(paragraph):
                # 段落格式
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.5  # 1.5倍行距
                paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两格
                paragraph_format.space_before = Pt(6)  # 段前间距
                paragraph_format.space_after = Pt(6)  # 段后间距
                
                # 字体设置
                for run in paragraph.runs:
                    # 对每个字符分别设置字体
                    new_runs = []
                    current_text = ""
                    current_is_chinese = None
                    
                    for char in run.text:
                        # 判断字符类型
                        is_chinese = '\u4e00' <= char <= '\u9fff'
                        
                        if current_is_chinese is None:
                            current_is_chinese = is_chinese
                            current_text = char
                        elif current_is_chinese == is_chinese:
                            current_text += char
                        else:
                            # 添加新的run
                            new_run = paragraph.add_run(current_text)
                            if current_is_chinese:
                                new_run.font.name = '宋体'
                                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            else:
                                new_run.font.name = 'Times New Roman'
                                new_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                                new_run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            new_run.font.size = Pt(12)  # 小四
                            new_runs.append(new_run)
                            
                            current_is_chinese = is_chinese
                            current_text = char
                    
                    # 添加最后一段文本
                    if current_text:
                        new_run = paragraph.add_run(current_text)
                        if current_is_chinese:
                            new_run.font.name = '宋体'
                            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        else:
                            new_run.font.name = 'Times New Roman'
                            new_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                            new_run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        new_run.font.size = Pt(12)
                        new_runs.append(new_run)
                    
                    # 移除原始run
                    # 安全地移除原始run元素
                    p = paragraph._element
                    # 创建一个副本以避免在迭代时修改列表
                    original_runs = list(paragraph.runs)
                    # 只保留最后len(new_runs)个run（新创建的）
                    if len(original_runs) > len(new_runs):
                        for run in original_runs[:-len(new_runs)]:
                            try:
                                p.remove(run._element)
                            except:
                                # 忽略移除失败的情况
                                pass
            
            # 添加题目
            doc.add_heading("一、题目", level=1)
            
            for i, (q, a) in enumerate(zip(questions, answers), 1):
                # 添加题目（带大纲级别）
                question_heading = doc.add_heading(f"题目 {i}", level=2)
                for run in question_heading.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(14)
                    run.font.bold = True
                
                # 添加题目内容
                q_paragraph = doc.add_paragraph(q)
                set_paragraph_format(q_paragraph)
                
                # 如果不是整卷模式，答案紧跟题目
                if not is_full_exam:
                    answer_heading = doc.add_heading(f"题目 {i} 参考答案", level=2)
                    for run in answer_heading.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(14)
                        run.font.bold = True
                    
                    a_paragraph = doc.add_paragraph(a)
                    set_paragraph_format(a_paragraph)
                
                doc.add_paragraph()
            
            # 如果是整卷模式，答案统一放在最后
            if is_full_exam:
                doc.add_page_break()
                doc.add_heading("二、参考答案", level=1)
                
                for i, a in enumerate(answers, 1):
                    answer_heading = doc.add_heading(f"题目 {i} 参考答案", level=2)
                    for run in answer_heading.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(14)
                        run.font.bold = True
                    
                    a_paragraph = doc.add_paragraph(a)
                    set_paragraph_format(a_paragraph)
                    doc.add_paragraph()
            
            doc.save(filename)
        
        st.success(f"✅ Word导出成功：{filename}")
        return filename
    except Exception as e:
        st.error(f"Word导出失败：{str(e)[:100]}")
        return None


def get_word_download_data(questions, answers, is_full_exam=False):
    """生成 Word 文件的二进制数据，用于 Streamlit 下载
    
    Args:
        questions: 题目列表
        answers: 答案列表
        is_full_exam: 是否为整卷模式
    
    Returns:
        bytes: Word 文件的二进制数据
    """
    # 重新检查依赖，确保在部署环境中也能正确检测
    try:
        import docx
        from io import BytesIO
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import parse_xml
        import re
        docx_available = True
    except ImportError:
        docx_available = False
        st.error("导出功能不可用：python-docx库未安装")
        return None

    # 检查latex2mathml依赖
    try:
        from latex2mathml.converter import convert as latex_to_mathml
        latex2mathml_available = True
    except ImportError:
        latex2mathml_available = False
        st.warning("latex2mathml库未安装，公式将以文本形式显示")

    if not questions or not answers:
        st.error("导出失败：无题目/答案可导出")
        return None

    try:
        # 辅助函数：把带有LaTeX的文本添加到Word段落中（简化版本）
        def add_text_with_latex(paragraph, text):
            """把带有$...$ LaTeX公式的文本添加到Word段落中，公式直接保留为LaTeX格式"""
            if not text:
                return
            
            # 使用全局的ensure_latex_format函数清理内容
            text = ensure_latex_format(text)
            
            # 直接添加文本，保持LaTeX格式原样，让用户看到完整内容
            paragraph.add_run(text)
        
        # 直接使用python-docx库，避免依赖spire.doc
        # 在内存中创建文档
        doc = docx.Document()
        
        # 设置页面边距
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        
        # 添加试卷标题
        title = doc.add_heading("控制原理学习资料", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 设置标题格式
        for run in title.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(22)
            run.font.bold = True
        
        doc.add_paragraph()
        
        # 辅助函数：设置段落格式
        def set_paragraph_format(paragraph):
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5
            paragraph_format.first_line_indent = Cm(0.74)
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            
            # 设置所有文本的字体
            for run in paragraph.runs:
                try:
                    run.font.size = Pt(12)
                    # 设置中文字体
                    try:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    except:
                        pass
                    # 设置英文字体
                    try:
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    except:
                        pass
                except:
                    pass
        
        # 只添加题目，不添加参考答案
        doc.add_heading("一、题目", level=1)
        
        for i, q in enumerate(questions, 1):
            question_heading = doc.add_heading(f"题目 {i}", level=2)
            for run in question_heading.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(14)
                run.font.bold = True
            
            # 添加题目文本
            q_paragraph = doc.add_paragraph()
            add_text_with_latex(q_paragraph, q)
            set_paragraph_format(q_paragraph)
            
            doc.add_paragraph()
        
        # 保存到内存，避免使用临时文件
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError as e:
        st.error(f"导出功能依赖缺失：{str(e)}")
        return None
    except Exception as e:
        st.error(f"生成Word下载数据失败：{str(e)}")
        return None


# ---------------------- 绘图与分析函数（大厂级：性能+容错+LaTeX） ----------------------


# 补充缺失的calculate_second_order_params函数
def calculate_second_order_params(num, den, K):
    """计算二阶系统性能指标（考研核心考点）"""
    try:
        # 保存用户输入的原始系数
        original_num = num.copy()
        original_den = den.copy()
        
        # 创建传递函数用于计算
        G = ctrl.TransferFunction(num, den)

        # 检查是否为二阶系统
        if len(original_den) != 3:  # 严格判断二阶系统
            st.warning("⚠️ 传递函数非二阶系统，无法计算二阶性能指标（请输入分母为2次的传递函数）")
            return None

        # 大厂级：参数合法性校验
        if original_den[0] <= 0 or original_den[2] <= 0:
            st.warning("⚠️ 二阶系统分母系数不合法（需正数）")
            return None

        omega_n = np.sqrt(original_den[2] / original_den[0])
        zeta = original_den[1] / (2 * np.sqrt(original_den[0] * original_den[2]))
        
        # 计算稳态增益（分子常数项/分母常数项）
        steady_gain = original_num[-1] / original_den[-1]

        # 计算性能指标（容错：避免除零/根号负数）
        if zeta < 1 and zeta >= 0:
            overshoot = np.exp(-zeta * np.pi / np.sqrt(1 - zeta ** 2)) * 100
            peak_time = np.pi / (omega_n * np.sqrt(1 - zeta ** 2)) if omega_n > 0 else np.inf
        else:
            overshoot = 0
            peak_time = np.inf

        settling_time = 4 / (zeta * omega_n) if (zeta > 0 and omega_n > 0) else np.inf  # 2%误差带

        # 显示传递函数（使用用户输入的原始系数）
        st.success("✅ 传递函数解析成功：")
        num_latex = poly_to_latex(original_num)
        den_latex = poly_to_latex(original_den)
        safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))
        
        # 先绘制阶跃响应图
        try:
            configure_matplotlib_for_chinese()
            fig, ax = plt.subplots(figsize=(10, 4))
            
            t = np.linspace(0, 10 / omega_n, 1000)
            y, _ = ctrl.step_response(G, T=t)
            
            ax.plot(t, y, linewidth=2, color='#8da1b9')
            ax.set_title("二阶系统阶跃响应", fontsize=18)
            ax.set_xlabel("时间 (s)", fontsize=16)
            ax.set_ylabel("输出", fontsize=16)
            ax.grid(True, alpha=0.7)
            ax.axhline(y=steady_gain, color='red', linestyle='--', alpha=0.7, label=f"稳态值: {steady_gain:.2f}")
            ax.legend(fontsize=14)
            
            try:
                plt.tight_layout()
            except:
                pass
            st.pyplot(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ 阶跃响应绘制失败：{str(e)[:50]}")

        # 学习辅助：移到侧边栏
        add_learning_helper("二阶系统")

        # 再显示分析结果
        st.markdown('<div class="result-card">', unsafe_allow_html=True)
        st.subheader("📊 二阶系统性能分析")
        st.markdown("#### 1. 系统参数")
        st.markdown(f"自然频率 $\\omega_n$: {omega_n:.2f} rad/s")
        st.markdown(f"阻尼比 $\\zeta$: {zeta:.2f}")
        st.markdown(f"稳态增益 $K$: {steady_gain:.2f}")
        st.info("**说明：** 改变分子只会改变稳态增益 $K$，不会改变 $\\omega_n$、$\\zeta$、超调量、调节时间、峰值时间等动态性能指标。")
        
        st.markdown("#### 2. 动态性能指标")
        st.markdown(f"超调量 $\\sigma\\%$: {overshoot:.2f}%")
        settling_time_str = f"{settling_time:.2f}" if np.isfinite(settling_time) else "无穷大"
        st.markdown(f"调节时间 $t_s$（2%误差带）: {settling_time_str} s")
        peak_time_str = f"{peak_time:.2f}" if np.isfinite(peak_time) else "无"
        peak_time_unit = " s" if np.isfinite(peak_time) else ""
        st.markdown(f"峰值时间 $t_p$: {peak_time_str}{peak_time_unit}")
        
        # 稳定性判断
        st.markdown("#### 3. 稳定性分析")
        if zeta > 1:
            st.success("✅ 过阻尼系统，无超调，响应较慢")
        elif zeta == 1:
            st.success("✅ 临界阻尼系统，无超调，响应速度适中")
        elif 0 < zeta < 1:
            st.success("✅ 欠阻尼系统，有超调，响应较快")
        else:
            st.warning("⚠️ 负阻尼系统，不稳定")
        
        # 4. 阻尼特性分析
        st.markdown("#### 4. 阻尼特性分析")
        st.info(r"**阻尼特性：** 分析欠阻尼/临界阻尼/过阻尼工况，解释 $\zeta$ 对振荡与收敛速度的影响。")
        
        if zeta > 1:
            st.markdown(r"• **过阻尼（$\zeta > 1$）**：系统无振荡，响应单调收敛。两个实极点分别为 $s_{1,2} = -\zeta\omega_n \pm \omega_n\sqrt{\zeta^2-1}$。")
        elif zeta == 1:
            st.markdown(r"• **临界阻尼（$\zeta = 1$）**：系统无振荡，响应最快无超调。二重实极点 $s_{1,2} = -\omega_n$。")
        elif 0 < zeta < 1:
            st.markdown(r"• **欠阻尼（$0 < \zeta < 1$）**：系统有振荡，衰减振荡收敛。复极点 $s_{1,2} = -\zeta\omega_n \pm j\omega_n\sqrt{1-\zeta^2}$。")
            st.markdown(r"  - 阻尼振荡频率：$\omega_d = \omega_n\sqrt{1-\zeta^2}$")
        else:
            st.markdown(r"• **负阻尼（$\zeta < 0$）**：系统不稳定，响应发散。")
        
        # 5. 参数敏感性分析
        st.markdown("#### 5. 参数敏感性分析")
        st.info(r"**参数敏感性：** 分析 $\zeta$、$\omega_n$ 变化对系统响应的影响，指导参数调整。")
        
        st.markdown(r"• **自然频率 $\omega_n$ 的影响**：")
        st.markdown(r"  - $\omega_n$ 越大，系统响应速度越快（调节时间 $t_s \propto 1/\omega_n$）")
        st.markdown(r"  - $\omega_n$ 不影响超调量 $\sigma\%$（仅由 $\zeta$ 决定）")
        
        st.markdown(r"• **阻尼比 $\zeta$ 的影响**：")
        st.markdown(r"  - $\zeta$ 越大，超调量 $\sigma\%$ 越小（$\zeta \uparrow \Rightarrow \sigma\% \downarrow$）")
        st.markdown(r"  - $\zeta$ 过大（$\zeta > 0.707$），响应变慢（调节时间 $t_s \uparrow$）")
        st.markdown(r"  - 工程最优 $\zeta \approx 0.707$：兼顾超调量和响应速度")
        
        st.markdown('</div>', unsafe_allow_html=True)

        return {
            "omega_n": round(omega_n, 2),
            "zeta": round(zeta, 2),
            "steady_gain": round(steady_gain, 2),
            "overshoot": round(overshoot, 2),
            "settling_time": round(settling_time, 2) if np.isfinite(settling_time) else np.inf,
            "peak_time": round(peak_time, 2) if np.isfinite(peak_time) else np.inf
        }
    except Exception as e:
        st.warning(f"⚠️ 二阶系统参数计算失败：{str(e)[:50]}")
        return None


# ---------------------- 以下是修改后的绘图函数 ----------------------
def plot_root_locus(num, den, K):
    """绘制根轨迹图并分析关键参数（详细标注+动态性能+稳定范围）"""
    if not CONTROL_AVAILABLE:
        st.error("control库未安装，无法绘制根轨迹")
        return

    # 创建传递函数
    try:
        G = ctrl.TransferFunction(num, den)
    except Exception as e:
        st.error(f"传递函数创建失败：{str(e)}")
        return

    # 修复LaTeX格式错误：确保传递函数正确渲染
    st.success("✅ 传递函数解析成功：")
    num_latex = poly_to_latex(num)
    den_latex = poly_to_latex(den)
    safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))

    # 配置Matplotlib以支持中文
    configure_matplotlib_for_chinese()
    fig, ax = plt.subplots(figsize=(12, 7))
    
    # 给右边留出空间放标注
    plt.subplots_adjust(right=0.75)

    # 存储计算结果
    poles = None
    zeros = None
    break_points = []
    asymptote_sigma = None
    asymptote_angles = []
    critical_gain = None
    jw_intersections = []

    # 计算关键参数
    try:
        # 计算开环极点和零点 - 使用更安全的方式
        poles = np.array(G.poles())
        zeros = np.array(G.zeros())
        
        # 计算分离/汇合点 - 用安全的方法
        break_points = []
        try:
            num_np = np.array(num, dtype=float)
            den_np = np.array(den, dtype=float)
            
            num_deriv = np.polyder(num_np)
            den_deriv = np.polyder(den_np)
            
            # 安全计算：term1 = num' * den, term2 = num * den'
            term1 = np.polymul(num_deriv, den_np)
            term2 = np.polymul(num_np, den_deriv)
            
            # 安全地相减：先确保长度一致
            len1 = len(term1)
            len2 = len(term2)
            if len1 > len2:
                term2 = np.pad(term2, (len1 - len2, 0), mode='constant')
            elif len2 > len1:
                term1 = np.pad(term1, (len2 - len1, 0), mode='constant')
            
            d_num = term1 - term2
            
            # 求根
            all_roots = np.roots(d_num)
            
            # 筛选实根且在根轨迹实轴段上的
            break_points = []
            for r in all_roots:
                if np.isreal(r) and np.isfinite(r):
                    # 检查是否在实轴上的根轨迹段（简化版：只要是实根就保留）
                    break_points.append(float(r.real))
            
            # 去重并排序
            break_points = sorted(list(set(break_points)))
        except Exception as e:
            break_points = []
        
        # 计算渐近线 - 完全简化，不依赖复杂计算
        n = len(den) - 1
        m = len(num) - 1
        if n > m:
            asymptote_angles = [(180 * (2 * k + 1)) / (n - m) for k in range(n - m)]
            try:
                if len(zeros) > 0:
                    asymptote_sigma = (np.sum(np.real(poles)) - np.sum(np.real(zeros))) / (n - m)
                else:
                    asymptote_sigma = np.sum(np.real(poles)) / (n - m)
            except:
                asymptote_sigma = None
        
        # 计算与虚轴交点和临界增益
        try:
            gm, pm, wg, wp = ctrl.margin(G)
            if np.isfinite(wg):
                jw_intersections = [0 + 1j * wg, 0 - 1j * wg]
                critical_gain = gm if gm != float('inf') else None
        except:
            pass
    except Exception as e:
        st.warning(f"⚠️ 关键参数计算部分失败：{str(e)[:30]}")

    # 绘制根轨迹
    try:
        # 兼容不同版本的control库
        try:
            rlist, klist = ctrl.root_locus(G, ax=ax, grid=True, plot=True)
        except:
            rlist, klist = ctrl.root_locus(G, grid=True, plot=True)
    except Exception as e:
        st.warning(f"⚠️ 根轨迹绘制失败：{str(e)[:50]}")
        return

    # 添加详细标注 - 放在右边空白区域，不挡住图
    try:
        # 1. 绘制虚轴（稳定性边界）
        ax.axvline(x=0, color='#ff6b6b', linestyle='--', linewidth=2.5, alpha=0.8, label='虚轴（稳定边界）')
        
        # 收集所有需要显示的标注信息
        annotations = []
        
        # 2. 开环极点标注
        if poles is not None:
            for i, p in enumerate(poles):
                ax.plot(p.real, p.imag, 'rx', markersize=14, markeredgewidth=3, zorder=5, label='开环极点' if i == 0 else "")
                if abs(p.imag) < 1e-6:
                    annotations.append((f'$p_{{{i+1}}}$: {p.real:.2f}', '#ffebee'))
                else:
                    if p.imag > 0:
                        annotations.append((f'$p_{{{i+1}}}$: {p.real:.2f}+j{abs(p.imag):.2f}', '#ffebee'))
                    else:
                        pass  # 共轭极点只标注一个
        
        # 3. 开环零点标注
        if zeros is not None and len(zeros) > 0:
            for i, z in enumerate(zeros):
                ax.plot(z.real, z.imag, 'go', markersize=12, markeredgewidth=3, markerfacecolor='none', zorder=5, label='开环零点' if i == 0 else "")
                if abs(z.imag) < 1e-6:
                    annotations.append((f'$z_{{{i+1}}}$: {z.real:.2f}', '#e8f5e9'))
                else:
                    if z.imag > 0:
                        annotations.append((f'$z_{{{i+1}}}$: {z.real:.2f}+j{abs(z.imag):.2f}', '#e8f5e9'))
                    else:
                        pass  # 共轭零点只标注一个
        
        # 4. 渐近线中心
        if asymptote_sigma is not None:
            ax.plot(asymptote_sigma, 0, 'm*', markersize=14, zorder=4, label='渐近线中心')
            annotations.append((f'$\\sigma$: {asymptote_sigma:.2f}', '#f3e5f5'))
            
            # 绘制渐近线
            xlim = ax.get_xlim()
            ylim = ax.get_ylim()
            max_range = max(abs(xlim[1] - xlim[0]), abs(ylim[1] - ylim[0])) / 2
            
            for angle in asymptote_angles:
                angle_rad = np.radians(angle)
                dx = max_range * np.cos(angle_rad)
                dy = max_range * np.sin(angle_rad)
                ax.plot([asymptote_sigma - dx, asymptote_sigma + dx], 
                       [-dy, dy], 'm--', linewidth=2, alpha=0.6)
        
        # 5. 分离/汇合点
        if break_points:
            for i, bp in enumerate(break_points):
                ax.plot(bp, 0, 'yD', markersize=10, zorder=4, label='分离/汇合点' if i == 0 else "")
                annotations.append((f'分离点: {bp:.2f}', '#fff9c4'))
        
        # 6. 虚轴交点和临界增益
        if jw_intersections and critical_gain:
            for jw in jw_intersections:
                if jw.imag > 0:
                    ax.plot(0, jw.imag, 'c^', markersize=12, zorder=4, label='虚轴交点')
                    annotations.append((f'$K_{{cr}}$: {critical_gain:.2f}', '#e0f7fa'))
        
        # 在右侧分散显示标注 - 一部分在右上角，一部分在右下角，更美观
        xlim = ax.get_xlim()
        ylim = ax.get_ylim()
        
        # 计算右边的位置：在坐标轴范围的最右边外面
        right_x = xlim[1] + (xlim[1] - xlim[0]) * 0.15
        
        # 把标注分成两部分：上半部分和下半部分
        num_annotations = len(annotations)
        split_idx = num_annotations // 2
        
        # 上半部分 - 右上角（从上往下排列）
        upper_y_start = ylim[1] - (ylim[1] - ylim[0]) * 0.1
        upper_y_step = -(ylim[1] - ylim[0]) * 0.08
        
        for i, (text, color) in enumerate(annotations[:split_idx]):
            ax.text(right_x, upper_y_start + i * upper_y_step, text, 
                    fontsize=10, ha='left', va='center',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor=color, alpha=0.9))
        
        # 下半部分 - 右下角（从下往上排列）
        if split_idx < num_annotations:
            lower_y_start = ylim[0] + (ylim[1] - ylim[0]) * 0.1
            lower_y_step = (ylim[1] - ylim[0]) * 0.08
            
            for i, (text, color) in enumerate(annotations[split_idx:]):
                ax.text(right_x, lower_y_start + i * lower_y_step, text, 
                        fontsize=10, ha='left', va='center',
                        bbox=dict(boxstyle='round,pad=0.3', facecolor=color, alpha=0.9))
        
        # 设置网格
        ax.grid(True, which='both', linestyle='-', alpha=0.4)
        
        # 图例放在中间右边空白区域
        ax.legend(fontsize=11, bbox_to_anchor=(1.02, 0.5), loc='center left')
        
    except Exception as e:
        st.warning(f"⚠️ 标注绘制部分失败：{str(e)[:30]}")

    plt.title("根轨迹图", fontsize=20, fontweight='bold', pad=20)
    plt.xlabel("实轴 (Re)", fontsize=16, labelpad=15)
    plt.ylabel("虚轴 (Im)", fontsize=16, labelpad=15)
    try:
        plt.tight_layout()
    except:
        pass
    st.pyplot(fig, use_container_width=True)

    # 学习辅助：移到侧边栏
    add_learning_helper("根轨迹")

    # 关键参数分析
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.subheader("📊 根轨迹关键参数分析")

    # 判断系统阶数 - 准确显示
    system_order = len(den) - 1
    if system_order == 1:
        st.info("提示：这是一个一阶系统")
    elif system_order == 2:
        st.info("提示：这是一个二阶系统")
    elif system_order == 3:
        st.info("提示：这是一个三阶系统")
    elif system_order == 4:
        st.info("提示：这是一个四阶系统")
    else:
        st.info(f"提示：这是一个{system_order}阶系统")

    # 1. 开环极点位置 + 稳定性影响分析
    st.markdown("#### 1. 开环极点与稳定性影响")
    try:
        if poles is not None:
            pole_text = []
            for i, p in enumerate(poles):
                if abs(p.imag) < 1e-6:
                    pole_text.append(f"极点{i + 1}：${p.real:.2f}$（实极点）")
                else:
                    pole_text.append(f"极点{i + 1}：${p.real:.2f} ± j{abs(p.imag):.2f}$（复极点）")
            for text in pole_text:
                st.write(text)

            # 新增：判断增益变化对稳定性的影响
            st.markdown("**📈 增益变化对稳定性的影响：**")
            if critical_gain and np.isfinite(critical_gain):
                st.success(f"✅ 稳定增益范围：$0 < K < {critical_gain:.2f}$")
                st.warning(f"⚠️ 当 $K = {critical_gain:.2f}$ 时，系统临界稳定")
                st.error(f"❌ 当 $K > {critical_gain:.2f}$ 时，系统不稳定")
            else:
                # 检查是否所有极点都在左半平面
                if all(np.real(poles) < -1e-6):
                    st.success("✅ 系统对所有 $K > 0$ 都稳定（所有开环极点在左半平面）")
                elif any(np.real(poles) > 1e-6):
                    st.error("❌ 存在右半平面开环极点，系统可能不稳定")
                else:
                    st.info("ℹ️ 存在虚轴极点，需进一步分析")
                    
        # 原有的开环稳定性判断
        if all(np.real(poles) < -1e-6):
            st.info("开环状态：所有开环极点均在左半平面，系统开环稳定")
        elif any(np.real(poles) > 1e-6):
            st.warning("开环状态：存在右半平面极点，系统开环不稳定")
        else:
            st.info("开环状态：存在虚轴极点，系统开环临界稳定")
    except Exception as e:
        st.write(f"开环极点计算：暂无法获取（{str(e)[:50]}）")

    # 2. 分离/汇合点
    st.markdown("#### 2. 分离/汇合点")
    try:
        if break_points:
            break_points_latex = ", ".join([f"{p:.2f}" for p in break_points])
            st.markdown(f"**分离/汇合点位置：** ${break_points_latex}$")
            st.info("📝 考点：分离点是根轨迹从实轴分离的位置，汇合点是回到实轴的位置")
        else:
            st.write("无分离/汇合点（该传递函数根轨迹无实轴分离/汇合）")
    except:
        st.write("暂无法计算分离/汇合点（通常因传递函数结构导致）")

    # 3. 渐近线参数
    st.markdown("#### 3. 渐近线参数")
    try:
        if asymptote_sigma is not None and len(asymptote_angles) > 0:
            angles_latex = ", ".join([f"{a:.0f}°" for a in asymptote_angles])
            st.markdown(f"**渐近线角度：** ${angles_latex}$")
            st.markdown(f"**渐近线中心（实轴截距）：** ${asymptote_sigma:.2f}$")
            st.info(r"📝 答题模板：渐近线角度 $\varphi = \frac{180^\circ (2k+1)}{n-m}$，中心 $\sigma = \frac{\sum p_i - \sum z_i}{n-m}$")
        else:
            st.write("无渐近线（极点个数 ≤ 零点个数）")
    except:
        st.write("暂无法计算渐近线参数")

    # 4. 动态性能估算（基于主导极点）
    st.markdown("#### 4. 动态性能估算（基于主导极点）")
    try:
        if poles is not None and len(poles) >= 2:
            # 寻找主导极点（最靠近虚轴的复极点对）
            complex_poles = [p for p in poles if abs(p.imag) > 1e-6]
            if len(complex_poles) >= 2:
                # 找到实部最大（最靠近虚轴）的复极点
                dominant_pole = max(complex_poles, key=lambda p: p.real)
                zeta = -dominant_pole.real / np.sqrt(dominant_pole.real**2 + dominant_pole.imag**2)
                omega_n = np.sqrt(dominant_pole.real**2 + dominant_pole.imag**2)
                
                if 0 < zeta < 1:
                    overshoot = np.exp(-zeta * np.pi / np.sqrt(1 - zeta**2)) * 100
                    settling_time = 4 / (zeta * omega_n) if zeta > 0 else float('inf')
                    
                    st.markdown(f"**主导极点：** ${dominant_pole.real:.2f} ± j{abs(dominant_pole.imag):.2f}$")
                    st.markdown(f"**阻尼比 ζ：** ${zeta:.3f}$")
                    st.markdown(f"**自然频率 ωₙ：** ${omega_n:.2f}$ rad/s")
                    st.markdown(f"**超调量 σ%：** ${overshoot:.2f}\\%$")
                    if np.isfinite(settling_time):
                        st.markdown(f"**调节时间 tₛ（2%误差带）：** ${settling_time:.2f}$ s")
                    
                    st.info("**什么是主导极点？** 高阶系统中，离虚轴最近的一对共轭复极点（或实极点）对系统动态性能起主导作用，称为主导极点。")
                    st.info("📝 性能估算说明：基于主导极点近似，适用于高阶系统有一对主导复极点的情况")
                else:
                    st.info("主导极点不是欠阻尼复极点，无法估算动态性能")
            else:
                st.write("无复极点对，无法估算动态性能")
    except Exception as e:
        st.write(f"动态性能估算：暂无法计算（{str(e)[:30]}）")

    st.markdown('</div>', unsafe_allow_html=True)


def plot_bode_diagram(num, den, K):
    """绘制伯德图并分析频域参数（详细标注+稳定判断+频域性能+校正建议）"""
    if not CONTROL_AVAILABLE:
        st.error("control库未安装，无法绘制伯德图")
        return

    # 创建传递函数
    try:
        G = ctrl.TransferFunction(num, den)
    except Exception as e:
        st.error(f"传递函数创建失败：{str(e)}")
        return

    st.success("✅ 传递函数解析成功：")
    num_latex = poly_to_latex(num)
    den_latex = poly_to_latex(den)
    safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))

    # 配置Matplotlib以支持中文
    configure_matplotlib_for_chinese()
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(13, 9), sharex=True)

    # 存储计算结果
    mag = None
    phase = None
    omega = None
    gm = None
    pm = None
    wg = None
    wp = None
    gm_dB = None

    try:
        mag, phase, omega = ctrl.bode(G, dB=True, Hz=False, deg=True, plot=False)
        gm, pm, wg, wp = ctrl.margin(G)
        gm_dB = 20 * np.log10(gm) if (gm > 0 and not np.isinf(gm) and not np.isnan(gm)) else np.inf
    except Exception as e:
        st.warning(f"⚠️ 伯德图数据计算失败：{str(e)[:50]}")
        return

    # 绘制幅频特性
    ax1.semilogx(omega, 20 * np.log10(mag), linewidth=2.5, color='#1e88e5', label='幅频特性')
    ax1.set_title("幅频特性 (dB)", fontsize=20, fontweight='bold', pad=20)
    ax1.set_ylabel("幅值 (dB)", fontsize=16, labelpad=15)
    ax1.grid(True, which='both', linestyle='-', alpha=0.4)
    ax1.grid(True, which='minor', linestyle=':', alpha=0.2)
    
    # 标注0dB线
    ax1.axhline(y=0, color='#ff6b6b', linestyle='--', linewidth=2.5, alpha=0.8, label='0 dB 线')
    
    # 标注截止频率
    if np.isfinite(wp) and not np.isnan(wp):
        ax1.axvline(x=wp, color='#f9a825', linestyle='--', linewidth=2, alpha=0.8, label=f'截止频率 $\\omega_c$={wp:.2f} rad/s')
        # 找到截止频率处的幅值
        mag_at_wc = np.interp(wp, omega, 20 * np.log10(mag))
        ax1.plot(wp, mag_at_wc, 'yo', markersize=12, markeredgewidth=2, zorder=5)
        # 智能选择标注位置
        ylim = ax1.get_ylim()
        if mag_at_wc > (ylim[0] + ylim[1]) / 2:
            ax1.text(wp * 1.1, mag_at_wc - 5, f'$\\omega_c$={wp:.2f}', 
                    fontsize=12, ha='left', va='top',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#fff9c4', alpha=0.9))
        else:
            ax1.text(wp * 1.1, mag_at_wc + 5, f'$\\omega_c$={wp:.2f}', 
                    fontsize=12, ha='left', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#fff9c4', alpha=0.9))

    # 标注转折频率和斜率（简化版）
    try:
        # 计算转折频率
        poles = np.roots(den)
        zeros = np.roots(num) if len(num) > 1 else []
        corner_freqs = []
        for p in poles:
            if abs(p.real) > 1e-6 and abs(p.imag) < 1e-6:
                corner_freqs.append(abs(p.real))
            elif abs(p.imag) > 1e-6:
                corner_freqs.append(np.sqrt(p.real**2 + p.imag**2))
        for z in zeros:
            if abs(z.real) > 1e-6 and abs(z.imag) < 1e-6:
                corner_freqs.append(abs(z.real))
            elif abs(z.imag) > 1e-6:
                corner_freqs.append(np.sqrt(z.real**2 + z.imag**2))
        
        corner_freqs = sorted(list(set([f for f in corner_freqs if f > 1e-6 and np.isfinite(f)])))
        for i, cf in enumerate(corner_freqs):
            if cf > omega[0] and cf < omega[-1]:
                ax1.axvline(x=cf, color='#90a4ae', linestyle=':', linewidth=1.5, alpha=0.6, 
                           label=f'转折频率{i+1}={cf:.2f}' if i == 0 else "")
                ax1.text(cf, ax1.get_ylim()[1] * 0.9, f'{cf:.1f}', 
                        fontsize=10, ha='center', va='top',
                        bbox=dict(boxstyle='round,pad=0.2', facecolor='#eceff1', alpha=0.9))
    except:
        pass

    ax1.legend(fontsize=11, loc='best')

    # 绘制相频特性
    ax2.semilogx(omega, phase, linewidth=2.5, color='#43a047', label='相频特性')
    ax2.set_title("相频特性 (°)", fontsize=20, fontweight='bold', pad=20)
    ax2.set_xlabel("频率 (rad/s)", fontsize=16, labelpad=15)
    ax2.set_ylabel("相位 (°)", fontsize=16, labelpad=15)
    ax2.grid(True, which='both', linestyle='-', alpha=0.4)
    ax2.grid(True, which='minor', linestyle=':', alpha=0.2)
    
    # 标注-180°线
    ax2.axhline(y=-180, color='#ff6b6b', linestyle='--', linewidth=2.5, alpha=0.8, label='-180° 线')
    
    # 标注相位穿越频率
    if np.isfinite(wg) and not np.isnan(wg):
        ax2.axvline(x=wg, color='#7e57c2', linestyle='--', linewidth=2, alpha=0.8, 
                   label=f'相位穿越频率 $\\omega_{{180}}$={wg:.2f} rad/s')
        # 找到相位穿越频率处的相位
        phase_at_wg = np.interp(wg, omega, phase)
        ax2.plot(wg, phase_at_wg, 'mo', markersize=12, markeredgewidth=2, zorder=5)
        # 智能选择标注位置
        if phase_at_wg > -180:
            ax2.text(wg * 1.1, phase_at_wg - 10, f'$\\omega_{{180}}$={wg:.2f}', 
                    fontsize=12, ha='left', va='top',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#ede7f6', alpha=0.9))
        else:
            ax2.text(wg * 1.1, phase_at_wg + 10, f'$\\omega_{{180}}$={wg:.2f}', 
                    fontsize=12, ha='left', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#ede7f6', alpha=0.9))

    # 在截止频率处也标注相位
    if np.isfinite(wp) and not np.isnan(wp):
        phase_at_wp = np.interp(wp, omega, phase)
        ax2.plot(wp, phase_at_wp, 'yo', markersize=10, markeredgewidth=2, zorder=4)

    ax2.legend(fontsize=11, loc='best')

    try:
        plt.tight_layout()
    except:
        pass
    st.pyplot(fig, use_container_width=True)

    # 学习辅助：移到侧边栏
    add_learning_helper("伯德图")

    # 频域参数分析
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.subheader("📊 伯德图关键参数分析")

    # 1. 截止频率、相位裕度 + 稳定性判断
    st.markdown("#### 1. 截止频率、相位裕度与闭环稳定性")
    try:
        if np.isfinite(wp) and not np.isnan(wp):
            st.markdown(f"**截止频率 $\\omega_c$（增益穿越频率）：** ${wp:.2f}$ rad/s")
            st.markdown(f"**相位裕度 $PM$：** ${pm:.2f}^\\circ$")

            if pm > 45:
                st.success("✅ 相位裕度>45°：系统动态性能良好，超调量小")
            elif 30 <= pm <= 45:
                st.info("ℹ️ 相位裕度30°~45°：系统性能一般，需关注稳定性")
            else:
                st.warning("⚠️ 相位裕度<30°：系统振荡严重，建议校正")
        else:
            st.write("截止频率：暂无法计算（系统无增益穿越0dB的频率）")

        st.markdown("**🎯 闭环系统稳定性判断：**")
        is_stable = True
        stability_msg = []
        
        if np.isfinite(pm) and pm <= 0:
            is_stable = False
            stability_msg.append("❌ 相位裕度≤0°，系统不稳定")
        if np.isfinite(gm_dB) and gm_dB <= 0:
            is_stable = False
            stability_msg.append("❌ 增益裕度≤0dB，系统不稳定")
        
        if is_stable:
            if np.isfinite(pm) and pm > 0 and (not np.isfinite(gm_dB) or gm_dB > 0):
                st.success("✅ 闭环系统稳定（相位裕度>0°且增益裕度>0dB）")
        else:
            for msg in stability_msg:
                st.error(msg)
                
    except Exception as e:
        st.write(f"截止频率与相位裕度计算：暂无法获取（{str(e)[:50]}）")

    # 2. 增益裕度与相位穿越频率
    st.markdown("#### 2. 增益裕度与相位穿越频率")
    try:
        if np.isfinite(wg) and not np.isinf(gm_dB) and not np.isnan(gm_dB):
            st.markdown(f"**相位穿越频率 $\\omega_{{180}}$：** ${wg:.2f}$ rad/s")
            st.markdown(f"**增益裕度 $GM$：** ${gm_dB:.2f}$ dB")

            if gm_dB > 6:
                st.success("✅ 增益裕度>6dB：系统稳定裕度充足，抗干扰能力强")
            elif gm_dB >= 0:
                st.info("ℹ️ 增益裕度0~6dB：系统稳定裕度不足，需优化")
            else:
                st.error("❌ 增益裕度<0dB：系统不稳定，必须校正")
        else:
            st.write("增益裕度：暂无法计算（系统无相位穿越-180°的频率）")
    except Exception as e:
        st.write(f"增益裕度计算：暂无法获取（{str(e)[:50]}）")

    # 3. 新增：频域性能分析
    st.markdown("#### 3. 频域性能分析")
    try:
        if np.isfinite(wp) and not np.isnan(wp):
            # 估算带宽（截止频率近似为带宽）
            bandwidth = wp
            st.markdown(f"**系统带宽 $\\omega_b$：** ≈ ${bandwidth:.2f}$ rad/s")
            st.info("📝 带宽说明：带宽近似等于截止频率，反映系统响应速度，带宽越大响应越快")
            
            # 估算谐振峰值（简化版）
            if np.isfinite(pm) and 0 < pm < 90:
                zeta_est = pm / 100  # 简化估算
                if 0 < zeta_est < 0.707:
                    Mr_est = 1 / (2 * zeta_est * np.sqrt(1 - zeta_est**2))
                    st.markdown(f"**谐振峰值 $M_r$（估算）：** ≈ ${Mr_est:.2f}$")
                    st.info("📝 谐振峰值说明：反映系统相对稳定性，Mr越大超调量越大")
            
            # 时域指标估算
            if np.isfinite(pm) and 0 < pm < 90:
                zeta_from_pm = pm / 100  # 简化关系
                if 0 < zeta_from_pm < 1:
                    overshoot_est = np.exp(-zeta_from_pm * np.pi / np.sqrt(1 - zeta_from_pm**2)) * 100
                    st.markdown(f"**超调量 $\\sigma\\%$（估算）：** ≈ ${overshoot_est:.2f}\\%$")
                    
    except Exception as e:
        st.write(f"频域性能分析：暂无法计算（{str(e)[:30]}）")

    # 4. 新增：校正设计建议
    st.markdown("#### 4. 校正设计建议")
    try:
        suggestions = []
        
        if np.isfinite(pm) and pm < 30:
            suggestions.append("⚠️ 相位裕度过小，建议采用**超前校正**提高相角裕度")
        elif np.isfinite(pm) and 30 <= pm <= 45:
            suggestions.append("ℹ️ 相位裕度一般，可考虑**超前-滞后校正**综合优化")
        
        if np.isfinite(gm_dB) and gm_dB < 6:
            suggestions.append("⚠️ 增益裕度不足，建议采用**滞后校正**提高低频增益")
        
        if np.isfinite(wp) and wp < 1:
            suggestions.append("📈 截止频率过低，响应慢，建议**超前校正**增大带宽")
        
        if not suggestions:
            suggestions.append("✅ 系统频域性能良好，无需校正")
        
        for suggestion in suggestions:
            st.write(suggestion)
        
        st.info("📝 校正说明：超前校正提高相角裕度和带宽；滞后校正提高稳态精度和增益裕度")
        
    except Exception as e:
        st.write(f"校正设计建议：暂无法提供（{str(e)[:30]}）")

    st.markdown('</div>', unsafe_allow_html=True)


def plot_second_order_step(G, num, den):
    """绘制二阶系统阶跃响应并分析性能指标（容错+LaTeX）"""
    if not CONTROL_AVAILABLE:
        st.error("control库未安装，无法绘制阶跃响应")
        return

    st.success("✅ 传递函数解析成功：")
    num_latex = poly_to_latex(num)
    den_latex = poly_to_latex(den)
    safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))

    # 修复：传递正确的参数给 calculate_second_order_params 函数
    params = calculate_second_order_params(num, den, 1.0)
    if params:
        # 配置Matplotlib以支持中文
        configure_matplotlib_for_chinese()
        fig, ax = plt.subplots(figsize=(10, 4))

        t = np.linspace(0, 10 / params["omega_n"], 1000)
        y, _ = ctrl.step_response(G, T=t)

        ax.plot(t, y, linewidth=2, color='#8da1b9')
        ax.set_title("二阶系统阶跃响应曲线", fontsize=18)
        ax.set_xlabel("时间 (s)", fontsize=16)
        ax.set_ylabel("输出", fontsize=16)
        ax.grid(True, alpha=0.7)
        ax.axhline(y=1, color='red', linestyle='--', alpha=0.7, label="稳态值")
        ax.legend(fontsize=14)

        try:
            plt.tight_layout()
        except:
            pass
        # 修复：移除key参数，避免FigureCanvasAgg.print_png错误
        st.pyplot(fig, use_container_width=True)

        # 显示性能指标（修复峰值时间显示问题）
        st.markdown('<div class="result-card">', unsafe_allow_html=True)
        st.subheader("📊 核心性能指标（考研答题直接用）")
        st.markdown(f"#### 1. 系统参数")
        st.write(f"自然频率 $\\omega_n$：${params['omega_n']}$ rad/s")
        st.write(f"阻尼比 $\\zeta$：${params['zeta']}$")

        st.markdown(f"#### 2. 动态性能指标")
        st.write(f"超调量 $\\sigma\\%$：${params['overshoot']}\\%$")
        st.write(f"调节时间 $t_s$（2%误差带）：${params['settling_time']}$ s")

        # 修复峰值时间显示
        if params['zeta'] >= 1:
            peak_time_str = "无峰值（临界/过阻尼）"
        else:
            peak_time_str = f"{params['peak_time']} s"
        st.write(f"峰值时间 $t_p$：{peak_time_str}")
        

        # 性能判断
        st.markdown(f"#### 3. 性能评价（考研评分点）")
        if params['zeta'] > 0.707:
            st.success("✅ 阻尼比>0.707：系统无超调，响应平稳（过阻尼）")
        elif 0.4 < params['zeta'] <= 0.707:
            st.info("ℹ️ 阻尼比0.4~0.707：系统超调小，响应快（欠阻尼，性能最佳）")
        elif params['zeta'] <= 0.4:
            st.warning("⚠️ 阻尼比<0.4：系统超调大，振荡严重（欠阻尼）")
        else:
            st.info("ℹ️ 阻尼比=1：临界阻尼，无超调但响应较慢")

        # 考研答题模板（LaTeX格式）
        template = r"📝 考研答题模板：已知$\zeta=" + str(params['zeta']) + r"$，$\omega_n=" + str(params['omega_n']) + r"$，则超调量$\sigma\%=" + str(params['overshoot']) + r"\%$，调节时间$t_s=" + str(params['settling_time']) + r"$s"
        st.info(template)
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.warning("⚠️ 当前传递函数非二阶系统，无法计算二阶性能指标（请输入分母为2次的传递函数）")


def plot_pid_step_response(G_plant, num, den, K, Kp, Ki, Kd, input_type_str):
    """绘制PID校正前后响应对比（支持多种输入类型：阶跃/斜坡/抛物线，优化版）"""
    if not CONTROL_AVAILABLE:
        st.error("control库未安装，无法绘制PID校正响应")
        return

    # 保存原始系数
    original_num = num.copy()
    original_den = den.copy()

    # 创建传递函数
    try:
        G_plant = ctrl.TransferFunction(original_num, original_den)
    except Exception as e:
        st.error(f"传递函数创建失败：{str(e)}")
        return

    st.success("✅ 被控对象传递函数解析成功：")
    num_latex = poly_to_latex(original_num)
    den_latex = poly_to_latex(original_den)
    safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))
    st.markdown(f"**开环增益 K：** ${K:.2f}$")

    # 解析输入类型
    if "单位阶跃" in input_type_str:
        input_type = "step"
        input_title = "单位阶跃输入 1(t)"
    elif "单位斜坡" in input_type_str:
        input_type = "ramp"
        input_title = "单位斜坡输入 r(t)=t"
    else:
        input_type = "parabola"
        input_title = "单位抛物线输入 r(t)=0.5t²"
    
    st.markdown(f"**输入类型：** {input_title}")

    # 自动识别系统型别
    def get_system_type(den_coeffs):
        """计算系统型别（分母末尾0的个数）"""
        system_type = 0
        den_reversed = den_coeffs[::-1]
        for coeff in den_reversed:
            if abs(coeff) < 1e-6:
                system_type += 1
            else:
                break
        return system_type

    # 计算校正前和校正后的系统型别
    # 校正前：原始系统
    type_original = get_system_type(original_den)
    
    # 校正后：带PID的系统（Ki!=0时会增加一个积分环节）
    type_pid = type_original
    if abs(Ki) > 1e-6:
        type_pid += 1

    st.markdown(f"**校正前系统型别：** {type_original}型")
    st.markdown(f"**校正后系统型别：** {type_pid}型")

    # 构建PID控制器 - 标准形式：C(s) = Kp + Ki/s + Kd*s = (Kd*s^2 + Kp*s + Ki)/s
    try:
        C = ctrl.TransferFunction([Kd, Kp, Ki], [1, 0])
        
        if Ki == 0 and Kd == 0:
            st.markdown(f"PID控制器（P控制）：$C(s) = {Kp:.1f}$")
        elif Ki == 0:
            st.markdown(f"PID控制器（PD控制）：$C(s) = {Kd:.1f}s + {Kp:.1f}$")
        elif Kd == 0:
            st.markdown(f"PID控制器（PI控制）：$C(s) = {Kp:.1f} + \\frac{{{Ki:.1f}}}{{s}}$")
        else:
            st.markdown(f"PID控制器：$C(s) = {Kd:.1f}s + {Kp:.1f} + \\frac{{{Ki:.1f}}}{{s}}$")
        
        st.markdown(f"**PID参数：** $K_p={Kp:.2f}$, $K_i={Ki:.2f}$, $K_d={Kd:.2f}$")
    except Exception as e:
        st.warning(f"⚠️ PID控制器构建失败：{str(e)[:50]}")
        return

    # 构建开环和闭环系统（关键：乘以开环增益 K）
    try:
        # 开环系统：校正前 = K * G_plant，校正后 = K * C * G_plant
        open_loop_original = K * G_plant
        open_loop_pid = K * C * G_plant
        
        # 闭环系统：单位负反馈
        closed_loop_original = ctrl.feedback(open_loop_original, 1)
        closed_loop_pid = ctrl.feedback(open_loop_pid, 1)
        
        st.markdown("📝 开环传递函数：$G_{open}(s) = K \\cdot C(s) \\cdot G(s)$（单位负反馈）")
        st.markdown("📝 闭环系统：$T(s) = \\frac{G_{open}(s)}{1 + G_{open}(s)}$")
    except Exception as e:
        st.warning(f"⚠️ 闭环系统构建失败：{str(e)[:50]}")
        return

    # 配置Matplotlib以支持中文
    configure_matplotlib_for_chinese()
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 9), sharex=True)

    # 时间范围
    t_end = 20
    t = np.linspace(0, t_end, 1000)
    
    # 生成输入信号
    if input_type == "step":
        r = np.ones_like(t)
    elif input_type == "ramp":
        r = t
    else:
        r = 0.5 * t**2
    
    # 计算响应 - 通用方法：使用状态方程模拟
    def simulate_response(sys, t, r):
        """模拟闭环系统对任意输入的响应"""
        try:
            # 将系统转换为状态方程
            sys_ss = ctrl.ss(sys)
            A, B, C, D = sys_ss.A, sys_ss.B, sys_ss.C, sys_ss.D
            
            # 模拟响应
            x = np.zeros((A.shape[0], len(t)))
            y = np.zeros(len(t))
            
            dt = t[1] - t[0]
            for i in range(1, len(t)):
                # 前向欧拉积分
                dx = A @ x[:, i-1] + B @ [r[i-1]]
                x[:, i] = x[:, i-1] + dx * dt
                y[i] = C @ x[:, i] + D @ [r[i]]
            
            return y
        except:
            try:
                # 备选方案：使用control库的forced_response
                _, y = ctrl.forced_response(sys, T=t, U=r)
                return np.array(y).flatten()
            except:
                return None
    
    y_original = None
    y_pid = None
    
    try:
        y_original = simulate_response(closed_loop_original, t, r)
        if y_original is not None:
            ax1.plot(t, y_original, '--', linewidth=1.5, color='#9e9e9e', label='校正前（单位负反馈）')
            ax2.plot(t, r - y_original, '--', linewidth=1.5, color='#9e9e9e', label='校正前误差')
    except Exception as e:
        st.warning(f"⚠️ 校正前响应计算失败：{str(e)[:50]}")

    try:
        y_pid = simulate_response(closed_loop_pid, t, r)
        if y_pid is not None:
            ax1.plot(t, y_pid, '-', linewidth=3.5, color='#f06292', label=f'PID校正后')
            ax2.plot(t, r - y_pid, '-', linewidth=3.5, color='#f06292', label='PID校正后误差')
    except Exception as e:
        st.warning(f"⚠️ 校正后响应计算失败：{str(e)[:50]}")

    # 绘制输入信号 - 用明显的深黄色
    ax1.plot(t, r, ':', linewidth=2.5, color='#ffc107', alpha=0.95, label='输入信号')

    # 在误差图上标注稳态误差数值
    if y_original is not None:
        e_org = r[-1] - y_original[-1]
        ax2.text(t_end * 0.85, e_org * 1.05, f'  {e_org:.4f}', 
                fontsize=10, color='#9e9e9e', fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.9))
    
    if y_pid is not None:
        e_pid = r[-1] - y_pid[-1]
        ax2.text(t_end * 0.85, e_pid * 0.95, f'  {e_pid:.4f}', 
                fontsize=10, color='#f06292', fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.9))

    # 第一个子图：系统响应
    ax1.set_title(f"PID校正前后响应对比 ({input_title})", fontsize=18, fontweight='bold')
    ax1.set_ylabel("系统输出", fontsize=14)
    ax1.grid(True, alpha=0.7)
    ax1.legend(fontsize=11, loc='best')

    # 第二个子图：误差
    ax2.set_title("误差信号对比", fontsize=14, fontweight='bold')
    ax2.set_xlabel("时间 (s)", fontsize=16)
    ax2.set_ylabel("误差 e(t)", fontsize=14)
    ax2.grid(True, alpha=0.7)
    ax2.legend(fontsize=11, loc='best')

    try:
        plt.tight_layout()
    except:
        pass
    st.pyplot(fig, use_container_width=True)

    # 学习辅助：移到侧边栏
    add_learning_helper("PID控制")

    # 计算性能指标
    def calculate_performance(y, r, t):
        """计算性能指标（含峰值时间）"""
        if y is None or len(y) == 0:
            return None
        
        e = r - y
        y_final = y[-1]
        r_final = r[-1]
        
        if not np.isfinite(y_final) or abs(y_final) < 1e-6:
            y_final = r_final
        
        steady_state_error = e[-1]
        
        # 超调量和峰值时间（仅对阶跃输入有意义）
        overshoot = 0
        peak_time = np.inf
        if input_type == "step":
            max_y = np.max(y)
            overshoot = ((max_y - y_final) / y_final) * 100 if y_final > 0 else 0
            # 计算峰值时间：找到第一个最大值的时间
            if max_y > y_final * 1.001:  # 只有当有超调时才计算
                peak_idx = np.argmax(y)
                peak_time = t[peak_idx]
        
        return {
            'overshoot': round(float(overshoot), 2),
            'peak_time': round(float(peak_time), 2) if np.isfinite(peak_time) else np.inf,
            'steady_state_error': round(float(steady_state_error), 4),
            'steady_state_value': round(float(y_final), 4)
        }

    perf_original = calculate_performance(y_original, r, t) if y_original is not None else None
    perf_pid = calculate_performance(y_pid, r, t) if y_pid is not None else None

    # 计算相位裕度（仅开环传递函数，与输入无关）
    pm_original = None
    pm_pid = None
    try:
        gm_org, pm_org, wg_org, wp_org = ctrl.margin(open_loop_original)
        pm_original = pm_org if np.isfinite(pm_org) else None
    except:
        pass
    
    try:
        gm_pid, pm_pid_val, wg_pid, wp_pid = ctrl.margin(open_loop_pid)
        pm_pid = pm_pid_val if np.isfinite(pm_pid_val) else None
    except:
        pass

    # 性能指标对比表
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.subheader("📊 PID校正分析")
    
    st.markdown("#### 1. 性能指标对比表")
    table_data = []
    table_data.append(["指标", "校正前", "校正后", "改善效果"])
    
    # 超调量（仅阶跃输入显示）
    if input_type == "step":
        ovs_org = perf_original['overshoot'] if perf_original else "—"
        ovs_pid = perf_pid['overshoot'] if perf_pid else "—"
        ovs_improve = "✅ 减小" if (perf_original and perf_pid and perf_pid['overshoot'] < perf_original['overshoot']) else ("⚠️ 增大" if (perf_original and perf_pid and perf_pid['overshoot'] > perf_original['overshoot']) else "—")
        table_data.append(["超调量", f"{ovs_org}%" if ovs_org != "—" else "—", f"{ovs_pid}%" if ovs_pid != "—" else "—", ovs_improve])
        
        # 峰值时间（仅阶跃输入显示）
        tp_org = perf_original['peak_time'] if perf_original else "—"
        tp_pid = perf_pid['peak_time'] if perf_pid else "—"
        tp_improve = "✅ 缩短" if (perf_original and perf_pid and np.isfinite(tp_org) and np.isfinite(tp_pid) and tp_pid < tp_org) else ("⚠️ 变长" if (perf_original and perf_pid and np.isfinite(tp_org) and np.isfinite(tp_pid) and tp_pid > tp_org) else "—")
        table_data.append(["峰值时间 $t_p$", f"{tp_org} s" if tp_org != "—" and np.isfinite(tp_org) else "—", f"{tp_pid} s" if tp_pid != "—" and np.isfinite(tp_pid) else "—", tp_improve])
    
    # 稳态误差
    ess_org = perf_original['steady_state_error'] if perf_original else "—"
    ess_pid = perf_pid['steady_state_error'] if perf_pid else "—"
    ess_improve = "✅ 减小" if (perf_original and perf_pid and abs(ess_pid) < abs(ess_org)) else ("⚠️ 增大" if (perf_original and perf_pid and abs(ess_pid) > abs(ess_org)) else "—")
    table_data.append(["稳态误差", f"{ess_org:.4f}" if ess_org != "—" else "—", f"{ess_pid:.4f}" if ess_pid != "—" else "—", ess_improve])
    
    # 相位裕度
    pm_org_str = f"{pm_original:.1f}°" if pm_original is not None else "—"
    pm_pid_str = f"{pm_pid:.1f}°" if pm_pid is not None else "—"
    pm_improve = "✅ 增大" if (pm_original is not None and pm_pid is not None and pm_pid > pm_original) else ("⚠️ 减小" if (pm_original is not None and pm_pid is not None and pm_pid < pm_original) else "—")
    table_data.append(["相位裕度", pm_org_str, pm_pid_str, pm_improve])
    
    # 用markdown表格显示
    table_md = "| " + " | ".join(table_data[0]) + " |\n"
    table_md += "| " + " | ".join(["---"] * len(table_data[0])) + " |\n"
    for row in table_data[1:]:
        table_md += "| " + " | ".join(str(x) for x in row) + " |\n"
    st.markdown(table_md)
    
    # 输入类型相关说明 + 系统型别针对性结论
    if input_type == "step":
        st.info("**说明**：单位阶跃输入下，超调量、峰值时间、稳态误差有意义。")
        if type_original == 0:
            st.markdown(f"- **校正前结论**：{type_original}型系统，阶跃输入**有稳态误差**")
        else:
            st.markdown(f"- **校正前结论**：{type_original}型系统，阶跃输入**无稳态误差**")
        
        if type_pid == 0:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，阶跃输入**有稳态误差**")
        else:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，阶跃输入**无稳态误差**")
    
    elif input_type == "ramp":
        st.info("**说明**：单位斜坡输入下，需要I型或更高型系统才能跟踪无稳态误差。")
        if type_original == 0:
            st.markdown(f"- **校正前结论**：{type_original}型系统，斜坡输入**无法跟踪（误差无穷大）**")
        elif type_original == 1:
            st.markdown(f"- **校正前结论**：{type_original}型系统，斜坡输入**有稳态误差**")
        else:
            st.markdown(f"- **校正前结论**：{type_original}型系统，斜坡输入**无稳态误差**")
        
        if type_pid == 0:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，斜坡输入**无法跟踪（误差无穷大）**")
        elif type_pid == 1:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，斜坡输入**有稳态误差**")
        else:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，斜坡输入**无稳态误差**")
    
    else:
        st.info("**说明**：单位抛物线输入下，需要II型或更高型系统才能跟踪无稳态误差。")
        if type_original <= 1:
            st.markdown(f"- **校正前结论**：{type_original}型系统，抛物线输入**无法跟踪（误差无穷大）**")
        elif type_original == 2:
            st.markdown(f"- **校正前结论**：{type_original}型系统，抛物线输入**有稳态误差**")
        else:
            st.markdown(f"- **校正前结论**：{type_original}型系统，抛物线输入**无稳态误差**")
        
        if type_pid <= 1:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，抛物线输入**无法跟踪（误差无穷大）**")
        elif type_pid == 2:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，抛物线输入**有稳态误差**")
        else:
            st.markdown(f"- **校正后结论**：{type_pid}型系统，抛物线输入**无稳态误差**")
    
    if perf_original:
        st.markdown(f"- 校正前：实际输出 ≈ {perf_original['steady_state_value']:.4f} → 稳态误差 ≈ {perf_original['steady_state_error']:.4f}")
    if perf_pid:
        st.markdown(f"- 校正后：实际输出 ≈ {perf_pid['steady_state_value']:.4f} → 稳态误差 ≈ {perf_pid['steady_state_error']:.4f}")
    
    # 特殊情况提示：纯P控制且Kp=1
    if abs(Kp - 1) < 1e-6 and abs(Ki) < 1e-6 and abs(Kd) < 1e-6:
        st.warning("**提示**：当前PID控制器为纯P控制且$K_p=1$，校正前后开环传递函数完全一致，因此性能指标无变化。若需观察校正效果，请调整$K_p$、$K_i$或$K_d$。")
    
    st.markdown("#### 2. 参数作用分析")
    st.markdown("**参数作用：** 解释 $K_p$（提高响应速度）、$K_i$（消除稳态误差）、$K_d$（抑制振荡）各自的贡献。")
    
    st.markdown("• **比例环节 $K_p$**：增大$K_p$提高系统响应速度，减小稳态误差，但会增大超调量，可能降低稳定性。")
    st.markdown("• **积分环节 $K_i$**：消除稳态误差（对阶跃输入，I型及以上系统稳态误差为0），但会降低系统相位裕度，易引起振荡。")
    st.markdown("• **微分环节 $K_d$**：提供超前相位，增大相位裕度，减小超调量，提高系统稳定性，但对高频噪声敏感。")
    
    st.markdown('</div>', unsafe_allow_html=True)


def calculate_steady_error(num, den, K, input_type, A):
    """计算稳态误差并显示结果（修复文字错误，扩展输入类型+LaTeX）"""
    if not CONTROL_AVAILABLE:
        st.error("control库未安装，无法计算稳态误差")
        return

    # 初始化所有变量，避免未定义引用（核心修复）
    ess = None
    ess_text = ""
    conclusion = ""

    # 创建传递函数
    try:
        G = ctrl.TransferFunction(num, den)
        G = ctrl.minreal(G)
    except Exception as e:
        st.error(f"传递函数创建失败：{str(e)}")
        return

    num_latex = poly_to_latex(num)
    den_latex = poly_to_latex(den)
    # 改用原始字符串避免LaTeX转义错误（核心修复）
    safe_latex_render(r"G(s) = \frac{%s}{%s}" % (num_latex, den_latex))

    # 系统型别判断（容错）
    system_type = 0
    try:
        den_reversed = den[::-1]
        for coeff in den_reversed:
            if abs(coeff) < 1e-6:
                system_type += 1
            else:
                break
    except Exception as e:
        st.warning(f"⚠️ 系统型别判断失败：{str(e)[:50]}")
        system_type = 0

    type_text = f"{system_type}型系统"
    st.success(f"✅ 系统型别识别结果：**{type_text}**")

    # 开环增益计算（容错）
    try:
        s = 1e-6
        G_at_0 = ctrl.evalfr(G, s)
        K = abs((s ** system_type) * G_at_0.real)
        st.info(f"📊 开环增益 $K$ ≈ {K:.2f}")
    except Exception as e:
        K = 0
        st.warning(f"⚠️ 暂无法计算开环增益，将使用默认值 $K=0$（{str(e)[:30]}）")

    # 扩展所有输入类型的稳态误差计算（LaTeX格式，统一用原始字符串）
    if input_type == "阶跃输入 1(t)":
        st.markdown("### 阶跃输入分析")
        try:
            safe_latex_render(r"r(t) = A·1(t) \quad R(s) = \frac{A}{s}")
            safe_latex_render(r"e_{ss} = \frac{A}{1 + K_p} \quad (K_p = \lim_{s \to 0} G(s))")
        except:
            st.markdown(r"$r(t) = A·1(t) \quad R(s) = \frac{A}{s}$")
            st.markdown(r"$e_{ss} = \frac{A}{1 + K_p} \quad (K_p = \lim_{s \to 0} G(s))$")

        if system_type >= 1:
            ess = 0.0
            ess_text = r"$e_{ss} = 0$"
            conclusion = "I型/II型系统对阶跃输入的稳态误差为0（积分环节消除了阶跃输入的稳态误差）"
        else:
            Kp = K
            ess = A / (1 + Kp) if (1 + Kp) != 0 else float('inf')
            ess_text = r"$e_{ss} = \frac{%s}{1 + %.2f} = %.3f$" % (A, Kp, ess)
            conclusion = "0型系统对阶跃输入存在稳态误差，误差大小与开环增益成反比"

    elif input_type == "斜坡输入 t":
        st.markdown("### 斜坡输入分析")
        try:
            safe_latex_render(r"r(t) = A·t \quad R(s) = \frac{A}{s^2}")
            safe_latex_render(r"e_{ss} = \frac{A}{K_v} \quad (K_v = \lim_{s \to 0} s·G(s))")
        except:
            st.markdown(r"$r(t) = A·t \quad R(s) = \frac{A}{s^2}$")
            st.markdown(r"$e_{ss} = \frac{A}{K_v} \quad (K_v = \lim_{s \to 0} s·G(s))$")

        if system_type >= 2:
            ess = 0.0
            ess_text = r"$e_{ss} = 0$"
            conclusion = "II型及以上系统对斜坡输入的稳态误差为0"
        elif system_type == 1:
            Kv = K
            ess = A / Kv if Kv != 0 else float('inf')
            ess_text = r"$e_{ss} = \frac{%s}{%.2f} = %.3f$" % (A, Kv, ess)
            conclusion = "I型系统对斜坡输入存在稳态误差，可通过增大开环增益减小误差"
        else:
            ess = float('inf')
            ess_text = r"$e_{ss} = ∞$"
            conclusion = "0型系统无法跟踪斜坡输入，稳态误差为无穷大"

    elif input_type == "抛物线输入 0.5t²":
        st.markdown("### 抛物线输入分析")
        try:
            safe_latex_render(r"r(t) = A·\frac{1}{2}t^2 \quad R(s) = \frac{A}{s^3}")
            safe_latex_render(r"e_{ss} = \frac{A}{K_a} \quad (K_a = \lim_{s \to 0} s^2·G(s))")
        except:
            st.markdown(r"$r(t) = A·\frac{1}{2}t^2 \quad R(s) = \frac{A}{s^3}$")
            st.markdown(r"$e_{ss} = \frac{A}{K_a} \quad (K_a = \lim_{s \to 0} s^2·G(s))$")

        if system_type >= 3:
            ess = 0.0
            ess_text = r"$e_{ss} = 0$"
            conclusion = "III型及以上系统对抛物线输入的稳态误差为0（实际工程中极少使用）"
        elif system_type == 2:
            Ka = K
            ess = A / Ka if Ka != 0 else float('inf')
            ess_text = r"$e_{ss} = \frac{%s}{%.2f} = %.3f$" % (A, Ka, ess)
            conclusion = "II型系统对抛物线输入存在稳态误差，需更高型别系统才能消除"
        else:
            ess = float('inf')
            ess_text = r"$e_{ss} = ∞$"
            conclusion = "0型/I型系统无法跟踪抛物线输入，稳态误差为无穷大"

    elif input_type.startswith("正弦输入"):
        st.markdown("### 正弦输入分析")
        safe_latex_render(r"r(t) = A·sin(ωt) \quad R(s) = \frac{Aω}{s^2 + ω^2}")
        safe_latex_render(r"e_{ss} = \frac{A}{|1 + G(jω)|}")
        ess = "依赖于频率ω"
        ess_text = r"$e_{ss} = \frac{A}{|1 + G(jω)|}$"
        conclusion = "线性系统对正弦输入的稳态误差与输入频率相关，可通过频率特性分析"

    elif input_type.startswith("余弦输入"):
        st.markdown("### 余弦输入分析")
        safe_latex_render(r"r(t) = A·cos(ωt) \quad R(s) = \frac{As}{s^2 + ω^2}")
        safe_latex_render(r"e_{ss} = \frac{A}{|1 + G(jω)|}")
        ess = "依赖于频率ω"
        ess_text = r"$e_{ss} = \frac{A}{|1 + G(jω)|}$"
        conclusion = "线性系统对余弦输入的稳态误差与输入频率相关，本质同正弦输入"

    elif input_type.startswith("指数输入"):
        st.markdown("### 指数输入分析")
        safe_latex_render(r"r(t) = A·e^{-at} \quad R(s) = \frac{A}{s + a}")
        safe_latex_render(r"e_{ss} = \frac{A|1 + G(-a)|}{|G(-a)|}")
        ess = "依赖于衰减系数a"
        ess_text = r"$e_{ss} = \frac{A|1 + G(-a)|}{|G(-a)|}$"
        conclusion = "指数输入的稳态误差需通过终值定理计算，前提是s=-a在左半平面"

    # 兜底：处理未知输入类型（核心修复）
    else:
        ess_text = r"$e_{ss} = 未知$"
        conclusion = f"暂不支持该输入类型：{input_type}"
        st.warning(f"⚠️ 未识别的输入类型：{input_type}")

    # 显示结果（修复缩进，确保在函数内部）
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.subheader("📊 稳态误差计算结果")
    st.write(f"**输入类型**：{input_type}")
    st.write(f"**系统型别**：{system_type}型系统")
    st.write(f"**开环增益**：$K = {K:.2f}$")
    st.write(f"**稳态误差**：{ess_text}")
    st.success(f"✅ 结论（考研答题直接用）：{conclusion}")
    st.markdown('</div>', unsafe_allow_html=True)

    # 学习辅助：移到侧边栏
    add_learning_helper("稳态误差")

    # 补充考研考点总结（完善功能）
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.markdown("#### 核心考点（必背）")
    st.write("""
    | 系统型别 | 阶跃输入 ($1(t)$) | 斜坡输入 ($t$) | 抛物线输入 ($0.5t^2$) |
    |----------|-------------------|----------------|-----------------------|
    | 0型      | $\\frac{A}{1+K}$   | $∞$            | $∞$                   |
    | I型      | $0$               | $\\frac{A}{K}$  | $∞$                   |
    | II型     | $0$               | $0$            | $\\frac{A}{K}$        |
    | III型    | $0$               | $0$            | $0$                   |
    """)
    st.markdown('</div>', unsafe_allow_html=True)


# ---------------------- 完整初始化：补充缺失的依赖判断 ----------------------
# 最终校验：确保所有依赖标记正确
if __name__ == "__main__":
    # 仅在直接运行时打印依赖状态（调试用）
    print(f"CONTROL_AVAILABLE: {CONTROL_AVAILABLE}")
    print(f"VOLC_AVAILABLE: {VOLC_AVAILABLE}")
    print(f"EXPORT_AVAILABLE: {EXPORT_AVAILABLE}")